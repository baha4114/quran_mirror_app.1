# -*- coding: utf-8 -*-
"""
قطب‌نمای قرآنی — پردازش آینه‌ای (نسخهٔ موبایل / اندروید)
بازنویسی‌شده با Kivy — همهٔ قابلیت‌های نسخهٔ دسکتاپ، بدون وابستگی به PyQt.
روی ویندوز با پایتون اجرا می‌شود و با GitHub Actions به APK تبدیل می‌شود.
"""
import os
import json
import shutil
import zipfile
import threading
import base64
from datetime import datetime

# --- شکل‌دهیِ بومیِ متنِ عربی/فارسی با Pango (فقط اندروید) ---
# اگر ارائه‌دهندهٔ Pango در بیلد موجود باشد، Kivy حروف را «متصل» و راست‌به‌چپ شکل می‌دهد
# و ویرایش/انتخاب/حذف در کادرهای متن کاملاً طبیعی می‌شود. مقدار 'pango,sdl2' یعنی اگر
# Pango نبود، خودکار به sdl2 برمی‌گردد تا برنامه هرگز کرش نکند. باید پیش از importهای kivy تنظیم شود.
if 'ANDROID_ARGUMENT' in os.environ:
    os.environ.setdefault('KIVY_TEXT', 'pango,sdl2')

from kivy.utils import platform as _kivy_platform
# روی اندروید/iOS تنظیماتِ ماوس اعمال نشود تا لمسِ انگشت به‌صورت استاندارد و روان مدیریت شود
# (فعال‌بودنِ ارائه‌دهندهٔ ماوس روی موبایل باعث تداخل با اسکرول و لمسِ دکمه‌ها می‌شود)
if _kivy_platform not in ('android', 'ios'):
    from kivy.config import Config
    Config.set('input', 'mouse', 'mouse,multitouch_on_demand')

from kivy.app import App
from kivy.core.text import LabelBase
from kivy.core.window import Window
from kivy.uix.screenmanager import ScreenManager, Screen, SlideTransition, FadeTransition
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.anchorlayout import AnchorLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.scrollview import ScrollView
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.image import Image
from kivy.uix.widget import Widget
from kivy.uix.popup import Popup
from kivy.uix.spinner import Spinner
from kivy.uix.behaviors import ButtonBehavior
from kivy.graphics import Color, RoundedRectangle, Rectangle
from kivy.animation import Animation
from kivy.clock import Clock
from kivy.metrics import dp
from kivy.utils import platform

import core
import features
import ai_manager
from rtl import rtl, rtl_multiline


def _native_text_shaping():
    """True اگر ارائه‌دهندهٔ متنِ فعالِ Kivy از نوع Pango باشد (شکل‌دهیِ بومیِ عربی/فارسی)."""
    try:
        from kivy.core.text import Label as _CoreTextLabel
        return 'pango' in (getattr(_CoreTextLabel, '__module__', '') or '').lower()
    except Exception:
        return False
from kivy.uix.recycleview import RecycleView
from kivy.uix.recycleboxlayout import RecycleBoxLayout
from kivy.uix.recycleview.views import RecycleDataViewBehavior

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# اگر پوشهٔ assets وجود داشته باشد از آن استفاده می‌کنیم؛ وگرنه فایل‌ها را از همین پوشهٔ اصلی می‌خوانیم
# (روی گیت‌هاب فایل‌ها در ریشه هستند، روی ویندوز داخل assets)
ASSET_DIR = os.path.join(BASE_DIR, 'assets')
if not os.path.isdir(ASSET_DIR):
    ASSET_DIR = BASE_DIR


def asset(name):
    return os.path.join(ASSET_DIR, name)


def _atomic_write_json(path, data, indent=2):
    """ذخیرهٔ امن (اتمیک) JSON.
    اول روی یک فایل موقت نوشته می‌شود و فقط پس از نوشتنِ کاملِ آن، با یک عملیاتِ
    جایگزینیِ اتمیک روی فایل اصلی می‌نشیند. یک نسخهٔ پشتیبان (.bak) از فایلِ سالمِ
    قبلی هم نگه داشته می‌شود تا قطع‌شدنِ ناگهانی وسطِ ذخیره، داده را خراب نکند."""
    import os as _os, json as _json, shutil as _shutil, tempfile as _tempfile
    directory = _os.path.dirname(_os.path.abspath(path)) or '.'
    try:
        if _os.path.exists(path) and _os.path.getsize(path) > 0:
            _shutil.copy2(path, path + '.bak')
    except Exception:
        pass
    fd, tmp = _tempfile.mkstemp(prefix='.tmp_', suffix='.json', dir=directory)
    try:
        with _os.fdopen(fd, 'w', encoding='utf-8') as f:
            _json.dump(data, f, ensure_ascii=False, indent=indent)
            f.flush()
            _os.fsync(f.fileno())
        _os.replace(tmp, path)
    except Exception:
        try:
            _os.remove(tmp)
        except Exception:
            pass
        raise


def normalize_mode(mode_value):
    """یکسان‌سازی نام عملگر با فرم استانداردِ کوتاه (سازگار با نسخهٔ ویندوز)."""
    mapping = {
        "تقارن درجا فقط آیه (سوره ثابت، آیه آینه می‌شود)": "تقارن درجا فقط آیه",
        "تقارن درجا فقط سوره (آیه ثابت، سوره آینه می‌شود)": "تقارن درجا فقط سوره",
        "تقارن درجا کامل (آینه‌ی کامل)": "تقارن درجا کامل",
        "جابجایی کامل (تعویض جای سوره و آیه)": "جابجایی کامل",
    }
    return mapping.get(mode_value, mode_value)


def _normalize_item_modes(item):
    """مبدّلِ خودکار: نام عملگرِ یک کشف (و مقصدهای گروهی‌اش) را به فرم استانداردِ کوتاه تبدیل می‌کند، تا فایل‌های قدیمی یا واردشده از ویندوز دقیقاً هم‌راستا شوند."""
    if not isinstance(item, dict):
        return item
    if 'mode' in item:
        item['mode'] = normalize_mode(item.get('mode', ''))
    tgts = item.get('all_targets')
    if isinstance(tgts, list):
        for t in tgts:
            if isinstance(t, dict) and 'operator' in t:
                t['operator'] = normalize_mode(t.get('operator', ''))
    return item


# ---- ثبت فونت‌ها ----
LabelBase.register(name='arabic', fn_regular=asset('font.ttf'), fn_bold=asset('font.ttf'))
LabelBase.register(name='ui', fn_regular=asset('font.ttf'), fn_bold=asset('font.ttf'))
# فونتِ مخصوصِ جملهٔ «به نام الله برای الله» (فایلِ A Ali.ttf کنارِ main.py)
LabelBase.register(name='besmele', fn_regular=asset('A Ali.ttf'), fn_bold=asset('A Ali.ttf'))
# فونت پیش‌فرض kivy (Roboto) را هم به font.ttf تغییر می‌دهیم تا
# همهٔ ویجت‌ها (منوی کشویی Spinner، عنوان Popup، دکمه‌های ساده، فایل‌یاب) فارسی را درست نشان دهند.
LabelBase.register(name='Roboto', fn_regular=asset('font.ttf'), fn_bold=asset('font.ttf'))

# ---- پالت رنگ ----
C_BG = (0.05, 0.08, 0.14, 1)
C_PANEL = (1, 1, 1, 0.10)
C_PANEL_SOLID = (0.10, 0.14, 0.22, 1)
# نسخه و نشانهٔ بیلد (روی صفحهٔ خانه نشان داده می‌شود) — هر بار کد را عوض کردی این را هم بالا ببر
BUILD_VERSION = '4.1'

def _tag_multiselect(container, tags, current_str, title_text):
    """کنترلِ چندانتخابیِ برچسب‌ها (رفتار شبکه). خروجی: تابعی که رشتهٔ پیوسته برمی‌گرداند."""
    sep = chr(1548) + ' '  # ، و فاصله
    sel = set()
    for _p in str(current_str or '').replace(chr(1548), sep.strip()).split(sep.strip()):
        _p = _p.strip()
        if _p and _p != 'نامشخص':
            sel.add(_p)
    all_tags = list(tags)
    for _t in sel:
        if _t not in all_tags:
            all_tags.append(_t)
    container.add_widget(RLabel(title_text, font_size='14sp', color=C_GOLD,
                               halign='right', size_hint_y=None, height=dp(30)))
    _sv = ScrollView(size_hint_y=None, height=dp(150), bar_width=dp(4))
    _grid = GridLayout(cols=2, size_hint_y=None, spacing=dp(6), padding=(dp(2), dp(2)))
    _grid.bind(minimum_height=_grid.setter('height'))
    _sv.add_widget(_grid)
    _btns = {}
    def _refresh():
        for _tt, _b in _btns.items():
            _on = (_tt in sel)
            _b._bg = list(C_GOLD if _on else (1, 1, 1, 0.10))
            _b.color = (0.05, 0.08, 0.14, 1) if _on else HOME_FG
            _b._state()
    def _toggle(_tt):
        if _tt in sel:
            sel.discard(_tt)
        else:
            sel.add(_tt)
        _refresh()
    for _t in all_tags:
        _cb = PillButton(_t, bg=(1, 1, 1, 0.10), fg=HOME_FG, size_hint_y=None,
                         height=dp(42), font_size='13sp', radius=14)
        _cb.bind(on_release=lambda _inst, _x=_t: _toggle(_x))
        _btns[_t] = _cb
        _grid.add_widget(_cb)
    container.add_widget(_sv)
    _refresh()
    def _get():
        return sep.join([t for t in all_tags if t in sel]) or 'نامشخص'
    return _get

BUILD_TAG = '2026-07-19'

C_GOLD = (0.95, 0.77, 0.36, 1)
C_BLUE = (0.15, 0.55, 0.92, 1)
C_PURPLE = (0.61, 0.28, 0.80, 1)
C_TEAL = (0.16, 0.70, 0.72, 1)    # turquoise for search
C_BURG = (0.45, 0.11, 0.19, 1)    # deep burgundy for the number button
C_GRAPHITE = (0.17, 0.19, 0.23, 1)  # unified refined dark graphite
C_INDIGO = (0.40, 0.36, 0.78, 1)  # calm indigo for AI section
C_NAVY = (0.13, 0.20, 0.42, 1)    # deep navy (sormeh-ei)
C_OLIVE = (0.24, 0.30, 0.13, 1)   # dark murky olive-green
C_ORANGE = (1.0, 0.60, 0.10, 1)
C_GREEN = (0.20, 0.72, 0.45, 1)
C_RED = (0.90, 0.28, 0.28, 1)
C_TEXT = (0.96, 0.97, 1, 1)
C_MUTED = (0.72, 0.78, 0.88, 1)
HOME_BTN = (0.13, 0.13, 0.16, 1)
HOME_FG = (0.97, 0.98, 1, 1)


# --- کمک‌تابع‌های گرافیکی (گرادیان پس‌زمینه و قاب نئون) ---
_BG_GRAD_CACHE = {}


def _bg_gradient():
    tex = _BG_GRAD_CACHE.get('bg')
    if tex is not None:
        return tex
    from kivy.graphics.texture import Texture
    h = 256
    top = (16, 26, 52)
    bot = (5, 8, 16)
    buf = bytearray()
    for i in range(h):
        t = i / (h - 1)
        buf += bytes((
            int(bot[0] + (top[0] - bot[0]) * t),
            int(bot[1] + (top[1] - bot[1]) * t),
            int(bot[2] + (top[2] - bot[2]) * t),
        ))
    tex = Texture.create(size=(1, h), colorfmt='rgb')
    tex.blit_buffer(bytes(buf), colorfmt='rgb', bufferfmt='ubyte')
    tex.wrap = 'clamp_to_edge'
    _BG_GRAD_CACHE['bg'] = tex
    return tex


def _neon_border(widget, color, width=1.6, alpha=0.85):
    from kivy.graphics import Color as _Color, Line as _Line
    with widget.canvas.after:
        _Color(color[0], color[1], color[2], alpha)
        ln = _Line(width=width)

    def _u(*a):
        try:
            ln.rounded_rectangle = (widget.x + dp(1), widget.y + dp(1),
                                    widget.width - dp(2), widget.height - dp(2), dp(14))
        except Exception:
            pass
    widget.bind(pos=_u, size=_u)
    Clock.schedule_once(_u, 0)


def P(text):
    """متن فارسی/عربی آمادهٔ نمایش."""
    return rtl_multiline(text)


# ==================================================================
# ویدجت‌های پایه
# ==================================================================
_MEASURE_CACHE = {}
_CORE_LABELS = {}


def _text_width(s, font_name, font_size):
    """عرض رندرشدهٔ یک رشته با همان فونت/اندازه."""
    key = (font_name, round(float(font_size), 1), s)
    v = _MEASURE_CACHE.get(key)
    if v is not None:
        return v
    try:
        from kivy.core.text import Label as CoreLabel
        ck = (font_name, round(float(font_size), 1))
        cl = _CORE_LABELS.get(ck)
        if cl is None:
            cl = CoreLabel(text='', font_name=font_name, font_size=font_size)
            _CORE_LABELS[ck] = cl
        cl.text = s
        cl.refresh()
        w = cl.content_width
    except Exception:
        w = len(s) * float(font_size) * 0.6
    if len(_MEASURE_CACHE) < 20000:
        _MEASURE_CACHE[key] = w
    return w


class RLabel(Label):
    """لیبل فارسی با شکل‌دهی راست‌به‌چپ و شکستن صحیح خطوط (رفع مشکل آینه‌ای)."""
    def __init__(self, text='', arabic=False, **kw):
        kw.setdefault('font_name', 'arabic' if arabic else 'ui')
        kw.setdefault('color', C_TEXT)
        kw.setdefault('halign', 'right')
        kw.setdefault('valign', 'middle')
        kw.setdefault('markup', False)
        self._fit_single = bool(kw.pop('fit_single', False))
        self._syncing = False
        self._base_fs = None
        self._raw = '' if text is None else str(text)
        super().__init__(**kw)
        self.bind(size=self._sync, font_size=self._sync)
        self._sync()

    def _wrap_para(self, para, max_w):
        words = para.split(' ')
        lines = []
        cur = ''
        for wd in words:
            trial = wd if not cur else cur + ' ' + wd
            if not cur or _text_width(rtl(trial), self.font_name, self.font_size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = wd
        if cur:
            lines.append(cur)
        return lines

    def _sync(self, *a):
        if getattr(self, '_syncing', False):
            return
        w = self.width
        if getattr(self, '_fit_single', False):
            self._fit_single_line(w)
            return
        self.text_size = (w, None)
        raw = self._raw
        if not raw:
            self.text = ''
            return
        try:
            if w and w > 8:
                max_w = max(1.0, w - dp(6))
                out = []
                for para in raw.split('\n'):
                    if para == '':
                        out.append('')
                        continue
                    for ln in self._wrap_para(para, max_w):
                        out.append(rtl(ln))
                self.text = '\n'.join(out)
            else:
                self.text = P(raw)
        except Exception:
            self.text = P(raw)

    def _fit_single_line(self, w):
        raw = self._raw or ''
        disp = rtl(raw) if raw else ''
        self.text_size = (None, None)
        if not disp:
            self.text = ''
            return
        if w and w > 8:
            if self._base_fs is None:
                self._base_fs = self.font_size
            avail = max(1.0, w - dp(10))
            fs = self._base_fs
            guard = 0
            while fs > dp(9) and _text_width(disp, self.font_name, fs) > avail and guard < 60:
                fs -= dp(1)
                guard += 1
            if abs(fs - self.font_size) > 0.5:
                self._syncing = True
                try:
                    self.font_size = fs
                finally:
                    self._syncing = False
        self.text = disp

    def set_text(self, text):
        self._raw = '' if text is None else str(text)
        self._sync()


class RoundBox(BoxLayout):
    """جعبهٔ گوشه‌گرد با پس‌زمینه، سایهٔ نرمِ شناور و لبهٔ شیشه‌ای."""
    def __init__(self, bg=C_PANEL_SOLID, radius=18, border=None, shadow=None, **kw):
        super().__init__(**kw)
        from kivy.graphics import Line as _Line
        self._bg = bg
        self._radius = radius
        self._border = border
        # سایه فقط برای کارت‌های تقریباً مات (تا از پشتِ پنل‌های خیلی شفاف تیره دیده نشود)
        _opaque = (len(bg) < 4) or (bg[3] >= 0.5)
        self._shadow = _opaque if shadow is None else shadow
        # لبهٔ روشنِ شیشه‌ای فقط برای پنل‌های نیمه‌شفاف (حسِ glassmorphism)
        self._glass = (not _opaque) and (shadow is None)
        with self.canvas.before:
            # ۱) سایهٔ نرمِ چندلایه (هرچه دورتر، بزرگ‌تر و محوتر) → حسِ شناوربودن
            self._shadows = []
            if self._shadow:
                for k in range(4):
                    Color(0, 0, 0, 0.08)
                    sr = RoundedRectangle(radius=[radius + dp(2) * k])
                    self._shadows.append((sr, k))
            # ۲) قابِ دورِ کارت
            if border:
                self._bcol = Color(*border)
                self._brect = RoundedRectangle(radius=[radius])
            # ۳) پس‌زمینه
            self._col = Color(*bg)
            self._rect = RoundedRectangle(radius=[radius])
            # ۴) لبهٔ روشنِ شیشه‌ای (بازتابِ ن��������ر روی شیشه)
            self._rim = None
            if self._glass:
                Color(1, 1, 1, 0.10)
                self._rim = _Line(width=1.2)
            # top rim-light for every card (light-from-above glass highlight)
            Color(1, 1, 1, 0.20)
            self._toprim = _Line(width=1.3)
        self.bind(pos=self._upd, size=self._upd)

    def _upd(self, *a):
        self._rect.pos = self.pos
        self._rect.size = self.size
        for sr, k in getattr(self, '_shadows', []):
            ex = dp(1) + dp(2.5) * k
            dy = dp(2) + dp(2) * k
            sr.pos = (self.x - ex, self.y - ex - dy)
            sr.size = (self.width + 2 * ex, self.height + 2 * ex)
        if self._border:
            self._brect.pos = (self.x - dp(1.5), self.y - dp(1.5))
            self._brect.size = (self.width + dp(3), self.height + dp(3))
        if getattr(self, '_rim', None) is not None:
            try:
                self._rim.rounded_rectangle = (self.x + dp(1), self.y + dp(1),
                                               self.width - dp(2), self.height - dp(2), self._radius)
            except Exception:
                pass
        if getattr(self, '_toprim', None) is not None:
            _r = self._radius
            self._toprim.points = [self.x + _r + dp(2), self.y + self.height - dp(1.5),
                                   self.x + self.width - _r - dp(2), self.y + self.height - dp(1.5)]

    def set_bg(self, color):
        self._col.rgba = color


class ClickCard(ButtonBehavior, RoundBox):
    """کارت شیشه‌ایِ کلیک‌پذیر (برای ردیف‌های فهرست نتایج)."""
    pass


class PillButton(Button):
    """دکمهٔ گوشه‌گرد رنگی با انیمیشن فشردن."""
    def __init__(self, text='', bg=C_BLUE, fg=(1, 1, 1, 1), radius=16, font_size='16sp', **kw):
        super().__init__(**kw)
        self.background_normal = ''
        self.background_color = (0, 0, 0, 0)
        self.font_name = 'ui'
        self.color = fg
        self.font_size = font_size
        self.text = P(text)
        self._bg = list(bg)
        self._radius = radius
        self._press_ins = 0
        self._grad = False
        with self.canvas.before:
            self._col = Color(1, 1, 1, 1)
            self._rect = RoundedRectangle(radius=[radius])
        self._refresh_bg()
        self.bind(pos=self._upd, size=self._upd, state=self._state)

    def _refresh_bg(self):
        bg = self._bg
        opaque = (len(bg) < 4) or (bg[3] >= 0.5)
        if opaque:
            self._grad = True
            top, bot = _grad_pair(bg)
            self._rect.texture = _get_grad_tex(top, bot)
            self._col.rgba = (1, 1, 1, 1)
        else:
            self._grad = False
            self._rect.texture = None
            self._col.rgba = list(bg)

    def _apply_rect(self):
        ins = getattr(self, '_press_ins', 0)
        self._rect.pos = (self.x + ins, self.y + ins)
        self._rect.size = (max(0, self.width - 2 * ins), max(0, self.height - 2 * ins))

    def _upd(self, *a):
        self._apply_rect()

    def _state(self, *a):
        # بازخوردِ لمسی: دکمه هنگامِ فشار کمی روشن‌تر و کمی فرورفته می‌شود
        self._refresh_bg()
        if self.state == 'down':
            if self._grad:
                self._col.rgba = (0.82, 0.82, 0.82, 1)
            else:
                self._col.rgba = [min(1, c * 1.25) for c in self._bg[:3]] + [self._bg[3]]
            self._press_ins = dp(2.5)
        else:
            self._press_ins = 0
        self._apply_rect()

    def set_text(self, text):
        self.text = P(text)


def _html_to_lines(raw_html):
    """تبدیل HTML به متن خوانا برای نمایش داخل خودِ برنامه (بدون نیاز به مرورگر)."""
    from html.parser import HTMLParser
    import html as _htmlmod

    class _Extract(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts = []
            self._skip = 0

        def handle_starttag(self, tag, attrs):
            if tag in ('script', 'style'):
                self._skip += 1
            elif tag == 'li':
                self.parts.append('\n• ')
            elif tag == 'br':
                self.parts.append('\n')
            elif tag in ('p', 'div', 'tr', 'h1', 'h2', 'h3', 'h4', 'ul', 'ol', 'table'):
                self.parts.append('\n')
            elif tag in ('td', 'th'):
                self.parts.append('   ')

        def handle_endtag(self, tag):
            if tag in ('script', 'style'):
                self._skip = max(0, self._skip - 1)
            elif tag in ('p', 'div', 'tr', 'h1', 'h2', 'h3', 'h4', 'li'):
                self.parts.append('\n')

        def handle_data(self, data):
            if not self._skip:
                self.parts.append(data)

    parser = _Extract()
    try:
        parser.feed(raw_html)
    except Exception:
        pass
    text = _htmlmod.unescape(''.join(parser.parts))
    out = []
    blank = False
    for ln in text.split('\n'):
        ln = ln.strip()
        if ln:
            out.append(ln)
            blank = False
        elif out and not blank:
            out.append('')
            blank = True
    return '\n'.join(out).strip()


def show_html_in_app(raw_html):
    """گزارش HTML را داخل خودِ برنامه (پنجرهٔ اسکرول‌شونده) نمایش می‌دهد."""
    text = _html_to_lines(raw_html)
    if not text.strip():
        toast('محتوایی برای نمایش یافت نشد.', 'گزارش')
        return
    # متن را به قطعه‌های کوچک می‌شکنیم و هر قطعه را در یک برچسبِ جدا می‌گذاریم.
    # دلیل: روی اندروید یک برچسبِ بسیار بلند، یک بافتِ (texture) گرافیکیِ غول‌آسا می‌سازد
    # که از حداکثرِ مجازِ GPU بزرگ‌تر می‌شود و «صفحهٔ سیاه/خالی» می‌دهد (روی ویندوز مشکلی ندارد).
    _lines = text.split('\n')
    _chunks = []
    _buf, _buf_len = [], 0
    for _ln in _lines:
        _buf.append(_ln)
        _buf_len += len(_ln) + 1
        if len(_buf) >= 25 or _buf_len >= 1200:
            _chunks.append('\n'.join(_buf))
            _buf, _buf_len = [], 0
    if _buf:
        _chunks.append('\n'.join(_buf))
    if not _chunks:
        _chunks = [text]

    root = BoxLayout(orientation='vertical', spacing=dp(8), padding=dp(8))
    sv = ScrollView(do_scroll_x=False, bar_width=dp(8), scroll_type=['bars', 'content'])
    box = BoxLayout(orientation='vertical', size_hint_y=None, padding=dp(12), spacing=dp(4))
    box.bind(minimum_height=box.setter('height'))
    # پس‌زمینهٔ تیرهٔ صریح تا متنِ روشن حتماً دیده شود
    from kivy.graphics import Color as _Clr, Rectangle as _Rect
    with box.canvas.before:
        _Clr(0.06, 0.09, 0.16, 1)
        _bg = _Rect(pos=box.pos, size=box.size)
    box.bind(pos=lambda _i, _v: setattr(_bg, 'pos', _v),
             size=lambda _i, _v: setattr(_bg, 'size', _v))
    for _ch in _chunks:
        _lbl = RLabel(_ch, font_size='15sp', halign='right', color=(1, 1, 1, 1), size_hint_y=None)
        _lbl.bind(texture_size=lambda _i, _v: setattr(_i, 'height', _v[1] + dp(8)))
        box.add_widget(_lbl)
    sv.add_widget(box)
    root.add_widget(sv)
    pop = Popup(title=P('گزارش'), content=root, size_hint=(0.96, 0.92),
                title_font='ui', title_align='center', separator_color=C_GOLD)
    close = PillButton('بستن', bg=C_RED, size_hint_y=None, height=dp(46))
    close.bind(on_release=lambda *a: pop.dismiss())
    root.add_widget(close)
    pop.open()


class _KbFocusMixin:
    """رفع مشکل کیبورد اندروید: آوردنِ خودکارِ باکس به بالای کیبورد."""

    def __init__(self, **kw):
        super().__init__(**kw)
        self.bind(focus=self._kb_on_focus)

    def _kb_on_focus(self, _inst, val):
        # وقتی باکس با لمسِ انگشت فوکوس می‌گیرد، خودکار بالای کیبورد می‌آید
        # (فقط روی موبایل؛ روی ویندوز/دسکتاپ هیچ کاری لازم نیست و باعثِ جلوگیری از هنگ)
        if False:  # غیرفعال شد: جابه‌جاییِ صفحه هنگامِ فوکوس روی اندروید باعثِ هنگ و سیاه‌شدنِ صفحه می‌شد
            Clock.schedule_once(self._kb_ensure_visible, 0.35)
            Clock.schedule_once(self._kb_ensure_visible, 0.6)

    def _kb_ensure_visible(self, _dt):
        try:
            w = self.parent
            sv = None
            while w is not None:
                if isinstance(w, ScrollView):
                    sv = w
                    break
                w = w.parent
            if sv is None:
                return
            content = sv.children[0] if sv.children else None
            if content is None:
                return
            ch = content.height
            vh = sv.height
            if ch <= vh:
                return  # محتوا کوتاه‌تر از پنجره است؛ چیزی برای اسکرول نیست
            # ۱) گوشهٔ پایینِ ویجت در مختصاتِ پنجره (با دادنِ 0,0 دیگر دوبار جمع نمی‌شود)
            wx, wy = self.to_window(0, 0)
            # ۲) تبدیل به مختصاتِ محتوای داخلِ اسکرول
            cx, cy = content.to_local(wx, wy)
            # ۳) پایینِ باکس حدود ۵۵٪ ارتفاعِ صفحه بالاتر از لبهٔ پایینِ ویوپورت بایستد (بالای کیبورد)
            gap = vh * 0.55
            target_viewport_bottom = cy - gap
            # ۴) تبدیل به نسبتِ اسکرولِ Kivy (بین ۰ و ۱)
            s = target_viewport_bottom / float(ch - vh)
            s = max(0.0, min(1.0, s))
            Animation(scroll_y=s, d=0.25, t='out_quad').start(sv)
        except Exception:
            pass
    # توابع on_touch_down و on_touch_up قبلی حذف شدند (تداخل و کرش با کیبوردِ سیستم).


class PlainInput(_KbFocusMixin, TextInput):
    pass


class PersianTextInput(_KbFocusMixin, TextInput):
    # فیلد متنی راست‌به‌چپ که هنگام تایپ، فارسی را درست (بدون آینه‌ای) نشان می‌دهد
    def __init__(self, on_change=None, **kw):
        kw.setdefault('font_name', 'ui')
        kw.setdefault('halign', 'right')
        kw.setdefault('multiline', False)
        super().__init__(**kw)
        self._logical = ''
        self._guard = False
        self.on_change = on_change
        # با تغییرِ عرضِ باکسِ چندخطی، شکستِ خطوطِ فارسی دوباره محاسبه می‌شود
        if self.multiline:
            self.bind(width=lambda *a: self._render())

    def _wrap_para(self, para, max_w):
        """شکستنِ یک بند به خطوطِ جداگانه بر اساسِ عرضِ در دسترس (در فضای منطقی)."""
        words = para.split(' ')
        lines = []
        cur = ''
        for wd in words:
            trial = wd if not cur else cur + ' ' + wd
            if not cur or _text_width(rtl(trial), self.font_name, self.font_size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = wd
        if cur:
            lines.append(cur)
        return lines

    def _display_multiline(self):
        """متنِ چندخطی را بندبند و کلمه‌به‌کلمه پیش‌شکست می‌دهد و هر خط را جداگانه
        راست‌به‌چپ می‌کند تا نه آینه‌ای شود و نه بخشی از خطِ اول به خطِ بعد بپرد."""
        logical = self._logical or ''
        if not logical:
            return ''
        try:
            pad = self.padding
            pl = pad[0] if isinstance(pad, (list, tuple)) and len(pad) >= 1 else dp(6)
            pr = pad[2] if isinstance(pad, (list, tuple)) and len(pad) >= 3 else pl
        except Exception:
            pl = pr = dp(6)
        avail = self.width - pl - pr - dp(6)
        out = []
        for para in logical.split('\n'):
            if para == '':
                out.append('')
                continue
            if avail and avail > 8:
                for ln in self._wrap_para(para, avail):
                    out.append(rtl(ln))
            else:
                out.append(rtl(para))
        return '\n'.join(out)

    def _render(self):
        self._guard = True
        try:
            if self.multiline:
                self.text = self._display_multiline()
                try:
                    self.cursor = self.get_cursor_from_index(len(self.text))
                except Exception:
                    pass
            else:
                self.text = rtl(self._logical) if self._logical else ''
                self.cursor = (len(self.text), 0)
        finally:
            self._guard = False
        if self.on_change:
            self.on_change(self._logical)

    def insert_text(self, substring, from_undo=False):
        if self._guard or from_undo:
            return super().insert_text(substring, from_undo=from_undo)
        # رفع کرش: اگر کیبوردِ اندروید متنِ خالی فرستاد، رندر و جستجو را بی‌دلیل اجرا نکن
        if not substring:
            return None
        self._logical += substring
        self._render()
        return None

    def do_backspace(self, from_undo=False, mode='bkspc'):
        if self._guard or from_undo:
            return super().do_backspace(from_undo=from_undo, mode=mode)
        if self._logical:
            self._logical = self._logical[:-1]
            self._render()
        return None

    @property
    def query(self):
        return self._logical

    def set_logical(self, value):
        self._logical = value or ''
        self._render()

    def clear_logical(self):
        self._logical = ''
        self._guard = True
        try:
            self.text = ''
        finally:
            self._guard = False


class PersianEditor(_KbFocusMixin, TextInput):
    """ویرایشگرِ متنِ فارسی/عربی با پشتیبانیِ کاملِ مکان‌نما، انتخاب و حذفِ درجا.

    مشکلِ قبلی: چون متنِ نمایشی «شکل‌دهی‌شده» بود، تایپ/حذف همیشه از انتهای متن انجام می‌شد.
    راهکار:
      • هنگام ویرایش (فوکوس دارد): متنِ خام نشان داده می‌شود و ویرایشِ بومیِ Kivy
        (کلیک در هر نقطه، انتخاب/Select، حذفِ همان‌جا، Backspace/Delete، کپی/چسباندن) کاملاً کار می‌کند.
      • بدون فوکوس: متن با اتصالِ حروف و راست‌به‌چپ «شکل‌دهی» می‌شود تا زیبا و درست دیده شود.
    API سازگار با نسخهٔ قبلی: query / set_logical / clear_logical.
    """

    def __init__(self, on_change=None, **kw):
        kw.setdefault('font_name', 'ui')
        kw.setdefault('halign', 'right')
        kw.setdefault('multiline', True)
        super().__init__(**kw)
        try:
            self.base_direction = 'rtl'  # اگر نسخهٔ Kivy پشتیبانی کند، ترتیبِ راست‌به‌چپ هم رعایت می‌شود
        except Exception:
            pass
        self._logical = kw.get('text', '') or ''
        self._shaped = False
        self._guard = False
        self.on_change = on_change
        self._native = _native_text_shaping()
        if self._native:
            # حالتِ بومی (Pango): متنِ نمایشی = متنِ منطقی؛ شکل‌دهی خودکار و مکان‌نما/انتخاب/حذف بومیِ Kivy
            if self._logical and not self.text:
                self.text = self._logical
            self.bind(text=self._on_text_native)
        else:
            self.bind(focus=self._on_focus_dir)
            if self.multiline:
                self.bind(width=lambda *a: (None if self.focus else self._render_shaped()))
            self._render_shaped()

    def _emit(self):
        self._update_preview()
        if self.on_change:
            try:
                self.on_change(self.query)
            except Exception:
                pass

    def _on_text_native(self, *a):
        self._logical = self.text
        self._emit()

    def _wrap_para(self, para, max_w):
        words = para.split(' ')
        lines = []
        cur = ''
        for wd in words:
            trial = wd if not cur else cur + ' ' + wd
            if not cur or _text_width(rtl(trial), self.font_name, self.font_size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = wd
        if cur:
            lines.append(cur)
        return lines

    def _shaped_display(self):
        raw = self._logical or ''
        if not raw:
            return ''
        if not self.multiline:
            return rtl(raw)
        try:
            pad = self.padding
            pl = pad[0] if isinstance(pad, (list, tuple)) and len(pad) >= 1 else dp(6)
            pr = pad[2] if isinstance(pad, (list, tuple)) and len(pad) >= 3 else pl
        except Exception:
            pl = pr = dp(6)
        avail = self.width - pl - pr - dp(6)
        out = []
        for para in raw.split(chr(10)):
            if para == '':
                out.append('')
                continue
            if avail and avail > 8:
                for ln in self._wrap_para(para, avail):
                    out.append(rtl(ln))
            else:
                out.append(rtl(para))
        return chr(10).join(out)

    def _render_shaped(self):
        self._guard = True
        try:
            self.text = self._shaped_display()
            self._shaped = True
        finally:
            self._guard = False

    def _render_raw(self):
        self._guard = True
        try:
            self.text = self._logical or ''
            self._shaped = False
        finally:
            self._guard = False

    def _on_focus_dir(self, _inst, val):
        if val:
            self._render_raw()
        else:
            self._logical = self.text
            self._render_shaped()
            self._emit()

    def insert_text(self, substring, from_undo=False):
        r = super().insert_text(substring, from_undo=from_undo)
        if getattr(self, '_native', False):
            return r
        if not self._guard and not self._shaped:
            self._logical = self.text
            self._emit()
        return r

    def do_backspace(self, from_undo=False, mode='bkspc'):
        r = super().do_backspace(from_undo=from_undo, mode=mode)
        if getattr(self, '_native', False):
            return r
        if not self._guard and not self._shaped:
            self._logical = self.text
            self._emit()
        return r

    @property
    def query(self):
        if getattr(self, '_native', False):
            return self.text
        return self.text if not self._shaped else self._logical

    def set_logical(self, value):
        self._logical = value or ''
        self._update_preview()
        if getattr(self, '_native', False):
            if self.text != self._logical:
                self.text = self._logical
            return
        if self.focus:
            self._render_raw()
        else:
            self._render_shaped()

    def clear_logical(self):
        self._logical = ''
        self._update_preview()
        if getattr(self, '_native', False):
            self.text = ''
            return
        self._guard = True
        try:
            self.text = ''
            self._shaped = False
        finally:
            self._guard = False

    def attach_preview(self, label):
        """اتصالِ برچسبِ پیش‌نمایشِ فقط‌خواندنی که متنِ درستِ شکل‌گرفته را زنده نشان می‌دهد."""
        self._preview = label
        self._update_preview()

    def _update_preview(self):
        p = getattr(self, '_preview', None)
        if p is None:
            return
        try:
            p.set_text(self._logical or '')
        except Exception:
            pass


class SeedInput(_KbFocusMixin, TextInput):
    """ورودی عددی بذر که با یک تاچ ساده فوکوس می‌گیرد و کیبورد را بالا می‌آورد
    (رفع مشکل بالا نیامدنِ کیبورد داخل ScrollView)."""
    pass


def toast(message, title='پیام', kind=None):
    # تشخیصِ خودکارِ نوعِ پیام از روی عنوان/متن (تا صدها فراخوانیِ موجود بدونِ تغییر، رنگ و نشانهٔ درست بگیرند)
    if kind is None:
        _t = (title or '') + ' ' + (message or '')
        if 'خطا' in _t:
            kind = 'error'
        elif ('✓' in _t) or ('ذخیره' in _t) or ('ثبت شد' in _t) or ('موفقیت' in _t) \
                or ('اضافه شد' in _t) or ('کپی شد' in _t):
            kind = 'success'
        elif ('یافت نشد' in _t) or ('خالی' in _t) or ('نمانده' in _t):
            kind = 'warn'
        else:
            kind = 'info'
    _styles = {
        'success': (C_GREEN, '✓'),
        'error': (C_RED, '×'),
        'warn': (C_ORANGE, '!'),
        'info': (C_GOLD, '●'),
    }
    accent, icon = _styles.get(kind, (C_GOLD, '●'))
    content = BoxLayout(orientation='vertical', padding=dp(16), spacing=dp(10))
    content.add_widget(RLabel(icon, font_size='36sp', halign='center', color=accent,
                              size_hint_y=None, height=dp(46)))
    content.add_widget(RLabel(message, font_size='16sp', halign='center'))
    p = Popup(title=P(title), content=content, size_hint=(0.85, 0.44),
              title_font='ui', title_align='center', separator_color=accent)
    btn = PillButton('باشه', bg=(C_BLUE if kind == 'info' else accent),
                     size_hint_y=None, height=dp(46))
    btn.bind(on_release=p.dismiss)
    content.add_widget(btn)
    p.open()
    _fade_in(content, 0.18)
    # پیام‌های موفقیت پس از چند لحظه خودکار و نرم محو می‌شوند
    if kind == 'success':
        def _auto(_dt):
            try:
                _a = Animation(opacity=0, d=0.35)
                _a.bind(on_complete=lambda *x: p.dismiss())
                _a.start(content)
            except Exception:
                try:
                    p.dismiss()
                except Exception:
                    pass
        Clock.schedule_once(_auto, 2.2)
    return p


def confirm(message, on_yes, title='تأیید'):
    content = BoxLayout(orientation='vertical', padding=dp(16), spacing=dp(12))
    content.add_widget(RLabel(message, font_size='16sp', halign='center'))
    row = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(10))
    p = Popup(title=P(title), content=content, size_hint=(0.85, 0.42),
              title_font='ui', title_align='center', separator_color=C_GOLD)
    yes = PillButton('بله', bg=C_GREEN)
    no = PillButton('انصراف', bg=C_RED)
    def _yes(*a):
        p.dismiss()
        on_yes()
    yes.bind(on_release=_yes)
    no.bind(on_release=p.dismiss)
    row.add_widget(yes)
    row.add_widget(no)
    content.add_widget(row)
    p.open()
    return p


def _fade_in(widget, d=0.28):
    """ظاهرشدنِ نرمِ یک ویجت (محو ← نمایان) برای ورودِ روانِ کارت‌ها و پیام‌ها."""
    try:
        widget.opacity = 0
        Animation(opacity=1, d=d, t='out_quad').start(widget)
    except Exception:
        pass


def empty_state(text, hint=None, icon='۝', height=dp(160)):
    """حالتِ خالیِ دلنشین: یک نشانهٔ محو + پیامِ اصلی + یک راهنماییِ تشویقی."""
    box = BoxLayout(orientation='vertical', size_hint_y=None, height=height,
                    padding=dp(16), spacing=dp(8))
    box.add_widget(Widget(size_hint_y=1))
    box.add_widget(RLabel(icon, arabic=True, font_size='52sp', halign='center',
                          color=(C_MUTED[0], C_MUTED[1], C_MUTED[2], 0.45),
                          size_hint_y=None, height=dp(64)))
    box.add_widget(RLabel(text, font_size='16sp', bold=True, halign='center',
                          color=C_TEXT, size_hint_y=None, height=dp(30)))
    if hint:
        box.add_widget(RLabel(hint, font_size='13sp', halign='center',
                              color=C_MUTED, size_hint_y=None, height=dp(44)))
    box.add_widget(Widget(size_hint_y=1))
    _fade_in(box, 0.4)
    return box


# ==================================================================
# صفحهٔ پایه (پس‌زمینه + هدر)
# ==================================================================
class BaseScreen(Screen):
    def __init__(self, title='', show_back=True, **kw):
        super().__init__(**kw)
        self.root_layout = FloatLayout()
        add = self.root_layout.add_widget
        # پس‌زمینه
        with self.root_layout.canvas.before:
            Color(1, 1, 1, 1)
            self._bgrect = Rectangle(pos=(0, 0), size=Window.size, texture=_bg_gradient())
        Window.bind(size=lambda *a: setattr(self._bgrect, 'size', Window.size))
        try:
            self.bg_image = Image(source=asset('bg.jpg'), allow_stretch=True,
                                  keep_ratio=False, color=(1, 1, 1, 0.5),
                                  size_hint=(1, 1), pos_hint={'x': 0, 'y': 0})
            add(self.bg_image)
        except Exception:
            pass
        # ستون اصلی
        self.container = BoxLayout(orientation='vertical', size_hint=(1, 1),
                                   padding=dp(10), spacing=dp(8))
        add(self.container)
        # هدر
        header = RoundBox(bg=(1, 1, 1, 0.06), orientation='horizontal',
                          size_hint_y=None, height=dp(56), padding=dp(8), spacing=dp(6))
        if show_back:
            back = PillButton('بازگشت', bg=(1, 1, 1, 0.14), size_hint_x=None, width=dp(110),
                              font_size='14sp')
            back.bind(on_release=self.go_back)
            header.add_widget(back)
        else:
            header.add_widget(Widget(size_hint_x=None, width=dp(4)))
        self.title_label = RLabel(title, font_name='ui', bold=True, font_size='19sp',
                                  halign='center', color=C_GOLD)
        header.add_widget(self.title_label)
        header.add_widget(Widget(size_hint_x=None, width=dp(110) if show_back else dp(4)))
        self.header = header
        self.container.add_widget(header)
        self.add_widget(self.root_layout)

    def go_back(self, *a):
        self.manager.transition = SlideTransition(direction='right')
        self.manager.current = 'home'

    def body(self, widget):
        self.container.add_widget(widget)


# ==================================================================
# خانه
# ==================================================================
class HomeScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='', show_back=False, **kw)
        app = App.get_running_app()
        scroll = ScrollView(size_hint=(1, 1), do_scroll_x=False, bar_width=dp(4))
        content = BoxLayout(orientation='vertical', size_hint_y=None, padding=dp(6), spacing=dp(12))
        content.bind(minimum_height=content.setter('height'))

        # آیهٔ محوری متحرک در نوار بالای صفحه (خاموش‌روشن)
        try:
            self.header.height = dp(106)
            self.title_label.font_name = 'besmele'
            self.title_label.color = C_GOLD
            self.title_label.font_size = '45sp'
            self.title_label.set_text('به نام الله برای الله')
            # کادر شیشه‌ای بماند (بدونِ قابِ زرد)؛ فقط شهابِ نورانی دورش بچرخد
            orbit_dot(self.header, C_GOLD)
            # جملهٔ «به نام الله برای الله» نرم خاموش/روشن می‌شود
            _va = (Animation(opacity=0.4, duration=1.6) + Animation(opacity=1, duration=1.6))
            _va.repeat = True
            _va.start(self.title_label)
        except Exception:
            pass

        title = RLabel('قطب‌نمای قرآنی', bold=True, font_size='33sp', halign='center',
                       color=C_TEXT, size_hint_y=None, height=dp(46))
        subtitle = RLabel('پردازش آینه‌ای (هولوگرافیک)', font_size='14sp', halign='center',
                          color=C_MUTED, size_hint_y=None, height=dp(28))
        content.add_widget(title)
        content.add_widget(subtitle)
        content.add_widget(Widget(size_hint_y=None, height=dp(10)))
        # نشانهٔ نسخه/بیلد — برای اطمینان از اینکه دقیقاً همین کد روی گوشی اجرا می‌شود
        build_tag = RLabel('﴿ إِنَّا نَحْنُ نَزَّلْنَا الذِّکْرَ وَإِنَّا لَهُ لَحَافِظُونَ ﴾',
                           arabic=True, bold=True, font_size='18sp', halign='center', color=C_GOLD,
                           size_hint_y=1)
        # --- kateebe: glass banner framing the ayah (bigger, ornate, softly glowing) ---
        ayah_box = RoundBox(bg=(0.05, 0.08, 0.14, 0.45), orientation='vertical',
                            size_hint=(None, None), height=dp(54), radius=16,
                            padding=[dp(14), dp(6)], pos_hint={'center_x': 0.5})
        ayah_box.add_widget(build_tag)
        content.add_widget(ayah_box)
        pulse_aura(ayah_box, C_GOLD, alpha=0.6, thickness=0.8)
        def _fit_ayah(*_a):
            try:
                _tw = _text_width(rtl(build_tag._raw), build_tag.font_name, build_tag.font_size)
            except Exception:
                _tw = 0
            if _tw and _tw > 8:
                ayah_box.width = _tw + dp(52)
            else:
                Clock.schedule_once(_fit_ayah, 0.05)
        Clock.schedule_once(_fit_ayah, 0)

        # پنل شیشه‌ایِ جستجوی متن آیه یا ترجمه (بدون نیاز به اعراب دقیق)
        vsbox = RoundBox(bg=(0.05, 0.08, 0.14, 0.62), orientation='vertical',
                         size_hint_y=None, height=dp(172), padding=dp(12), spacing=dp(8))
        vsbox.add_widget(RLabel('جستجوی متن آیه یا ترجمه', bold=True, font_size='16sp',
                                halign='center', color=C_TEXT, size_hint_y=None, height=dp(26)))
        self.vs_in = PersianTextInput(hint_text=P('جستجوی متن یا شمارهٔ آیه'),
                                      font_name='arabic', halign='right', font_size='16sp',
                                      multiline=False, size_hint_y=None, height=dp(48),
                                      background_color=(1, 1, 1, 0.92),
                                      foreground_color=(0.05, 0.08, 0.14, 1))
        self.vs_in.bind(on_text_validate=lambda *a: self.search_verse())
        vsbox.add_widget(self.vs_in)
        vbtnrow = BoxLayout(size_hint_y=None, height=dp(50), spacing=dp(8))
        vbtn = PillButton('انتخاب خودکار', bg=(1, 1, 1, 0.10), fg=C_TEXT)
        vbtn.bind(on_release=lambda *a: self.search_verse())
        vbtn_all = PillButton('نمایش لیست جستجو', bg=(1, 1, 1, 0.10), fg=C_TEXT)
        vbtn_all.bind(on_release=lambda *a: self.show_all_results())
        vbtnrow.add_widget(vbtn)
        vbtnrow.add_widget(vbtn_all)
        vsbox.add_widget(vbtnrow)
        content.add_widget(vsbox)
        # شهابِ نورانی مثلِ کادرِ بالا، اما کوچک‌تر، کم‌رنگ‌تر و سبک‌تر (دنبالهٔ کوتاه‌تر)
        self._vs_orbit = pulse_aura(vsbox, C_TEAL, alpha=0.7)
        # هنگامِ تایپ در کادرِ جستجو، شهابِ آن پررنگ می‌شود (المانِ فعال = روشن‌تر)
        self.vs_in.bind(focus=lambda _i, f: self._vs_orbit['set_alpha'](1.0 if f else 0.6))

        # کارت شیشه‌ایِ نتیجهٔ جستجو (زیر کادر جستجو؛ متن طلایی/نارنجی)
        self.verse_box = RoundBox(bg=(0.05, 0.08, 0.14, 0.62), orientation='vertical',
                                  size_hint_y=None, padding=dp(14), spacing=dp(6))
        self.verse_box.bind(minimum_height=self.verse_box.setter('height'))
        self.verse_meta = RLabel('', font_size='13sp', halign='center', color=C_ORANGE,
                                 size_hint_y=None, height=dp(0))
        self.verse = RLabel('نتیجهٔ جستجوی آیه در اینجا نمایش داده می‌شود',
                            arabic=True, font_size='18sp', halign='center', color=C_MUTED,
                            size_hint_y=None, height=dp(40))
        self.verse.bind(texture_size=lambda i, v: setattr(i, 'height', max(dp(40), v[1] + dp(10))))
        self.verse_box.add_widget(self.verse_meta)
        self.verse_box.add_widget(self.verse)
        content.add_widget(self.verse_box)
        pulse_aura(self.verse_box, C_GOLD)

        # پنل ورودی بذر
        seedbox = RoundBox(bg=(1, 1, 1, 0.09), orientation='vertical', size_hint_y=None,
                           height=dp(230), padding=dp(14), spacing=dp(10))
        seedbox.add_widget(RLabel('انتخاب بذر', bold=True, font_size='17sp',
                                  halign='center', color=C_GOLD, size_hint_y=None, height=dp(28)))
        inrow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(10))
        self.in_s = SeedInput(hint_text=P('سوره'), multiline=False, font_name='ui',
                              halign='center', font_size='18sp', input_filter='int',
                              background_color=(1, 1, 1, 0.92), foreground_color=(0.05, 0.08, 0.14, 1),
                              padding=[dp(8), dp(12)])
        self.in_a = SeedInput(hint_text=P('آیه'), multiline=False, font_name='ui',
                              halign='center', font_size='18sp', input_filter='int',
                              background_color=(1, 1, 1, 0.92), foreground_color=(0.05, 0.08, 0.14, 1),
                              padding=[dp(8), dp(12)])
        inrow.add_widget(self.in_s)
        inrow.add_widget(self.in_a)
        seedbox.add_widget(inrow)

        brow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        b_matrix = PillButton('پردازش ماتریس', bg=C_GOLD, fg=(0.05, 0.08, 0.14, 1), font_size='16sp')
        b_matrix.bind(on_release=lambda *a: self.run('matrix'))
        brow.add_widget(b_matrix)
        seedbox.add_widget(brow)

        prow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        b_sem = PillButton('پیش‌بینی (معنا)', bg=C_GRAPHITE, fg=HOME_FG)
        b_sem.bind(on_release=lambda *a: self.run('semantic'))
        b_num = PillButton('پیش‌بینی (اعداد)', bg=C_GRAPHITE, fg=HOME_FG)
        b_num.bind(on_release=lambda *a: self.run('numeric'))
        prow.add_widget(b_sem)
        prow.add_widget(b_num)
        seedbox.add_widget(prow)
        content.add_widget(seedbox)
        pulse_aura(seedbox, C_GREEN)

        # پنل دستیار هوش مصنوعی: ورود به گفتگو + تنظیمات کلید API
        aibox = RoundBox(bg=(0.05, 0.08, 0.14, 0.62), orientation='vertical', size_hint_y=None,
                         height=dp(118), padding=dp(12), spacing=dp(8))
        aibox.add_widget(RLabel('دستیار هوش مصنوعی', bold=True, font_size='16sp',
                                halign='center', color=C_GOLD, size_hint_y=None, height=dp(26)))
        airow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        b_chat = PillButton('گفتگو با هوش مصنوعی', bg=C_GRAPHITE, fg=HOME_FG, font_size='15sp')
        b_chat.bind(on_release=lambda *a: self.nav('chat'))
        b_set = PillButton('تنظیمات', bg=(0.20, 0.24, 0.32, 1), fg=C_TEXT, size_hint_x=None,
                           width=dp(110), font_size='14sp')
        b_set.bind(on_release=lambda *a: open_ai_settings())
        airow.add_widget(b_chat)
        airow.add_widget(b_set)
        aibox.add_widget(airow)
        content.add_widget(aibox)
        pulse_aura(aibox, C_INDIGO)

        # کاشی‌های ناوبری
        grid = GridLayout(cols=2, size_hint_y=None, spacing=dp(10))
        grid.bind(minimum_height=grid.setter('height'))
        nav = [
            ('لابراتوار کشفیات', C_GREEN, 'lab'),
            ('گلچین برگزیده', C_GOLD, 'featured'),
            ('جستجوی کشفیات', C_BLUE, 'search'),
            ('مدیریت برچسب‌ها', C_PURPLE, 'tags'),
            ('رسانه و معرفی', C_ORANGE, 'media'),
            ('راهنما', (0.3, 0.4, 0.55, 1), 'guide'),
            ('پشتیبان و بازیابی', (0.25, 0.5, 0.6, 1), 'backup'),
            ('درباره', (0.4, 0.35, 0.5, 1), 'about'),
        ]
        def _tint(col, f=0.30):
            return [c * f for c in col[:3]] + [1]
        for label, color, target in nav:
            b = PillButton(label, bg=_tint(color), fg=HOME_FG, size_hint_y=None,
                           height=dp(72), font_size='15sp', radius=16)
            # شهابِ کوچک‌ترِ کم‌رنگ‌تر با دنبالهٔ کوتاه دورِ هر دکمه (سبک‌تر برای CPU چون تعدادشان زیاد است)
            pulse_aura(b, color, alpha=0.7, thickness=0.85)
            b.bind(on_release=lambda inst, t=target: self.nav(t))
            grid.add_widget(b)
        content.add_widget(grid)
        content.add_widget(Widget(size_hint_y=None, height=dp(20)))

        scroll.add_widget(content)
        self.body(scroll)

    def _seed(self):
        try:
            s = int(core.conv(self.in_s.text.strip()))
            a = int(core.conv(self.in_a.text.strip()))
            return s, a
        except Exception:
            toast('لطفاً شمارهٔ سوره و آیهٔ معتبر وارد کنید.', 'خطا')
            return None

    def run(self, kind):
        seed = self._seed()
        if not seed:
            return
        app = App.get_running_app()
        s, a = seed
        if kind == 'matrix':
            fs, fa, is_fb, msg = app.data.apply_circular(s, a)
            if fs is None:
                toast('آیهٔ معتبر یافت نشد.', 'خطا')
                return
            scr = self.manager.get_screen('matrix')
            scr.show(fs, fa)
            self.manager.transition = SlideTransition(direction='left')
            self.manager.current = 'matrix'
        else:
            found = app.data.find_seed(s, a)
            if not found:
                toast('آیهٔ مورد نظر در دیتابیس یافت نشد.', 'خطا')
                return
            fs, fa = found
            scr = self.manager.get_screen('predict')
            scr.show(fs, fa, kind)
            self.manager.transition = SlideTransition(direction='left')
            self.manager.current = 'predict'

    def nav(self, target):
        scr = self.manager.get_screen(target)
        if hasattr(scr, 'refresh'):
            scr.refresh()
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = target

    def _verse_arb(self, s, a):
        try:
            v = App.get_running_app().data.get(s, a)
        except Exception:
            v = None
        if isinstance(v, dict):
            return v.get('arb', '') or ''
        return ''

    def _show_verse_result(self, s, a):
        arb = self._verse_arb(s, a)
        self.verse.color = C_GOLD
        self.verse.set_text(('« %s »' % arb) if arb else '« آیه یافت شد »')
        self.verse_meta.color = C_ORANGE
        self.verse_meta.set_text('سوره %s ، آیه %s' % (s, a))
        self.verse_meta.height = dp(24)
        try:
            self.in_s.text = str(s)
            self.in_a.text = str(a)
        except Exception:
            pass

    def search_verse(self):
        app = App.get_running_app()
        q = (self.vs_in.query or '').strip()
        if not q:
            toast('لطفاً بخشی از متن آیه یا ترجمه را وارد کنید.', 'جستجو')
            return
        if getattr(self, '_searching', False):
            return
        self._searching = True

        def _work():
            try:
                res = app.data.find_by_text(q)
                err = None
            except Exception as ex:
                res, err = None, str(ex)

            def _done(dt):
                self._searching = False
                if err:
                    toast('خطا در جستجو: %s' % err, 'خطا')
                    return
                if not res:
                    self.verse_meta.set_text('')
                    self.verse_meta.height = dp(0)
                    self.verse.color = C_RED
                    self.verse.set_text('آیه‌ای مطابق این متن یافت نشد')
                    toast('آیه‌ای با این متن پیدا نشد؛ دکمهٔ «نمایش همه» را امتحان کنید.', 'یافت نشد')
                    return
                s, a = res
                self._show_verse_result(s, a)

            Clock.schedule_once(_done, 0)

        threading.Thread(target=_work, daemon=True).start()

    def _pick_result(self, s, a, popup=None):
        self._show_verse_result(s, a)
        if popup is not None:
            try:
                popup.dismiss()
            except Exception:
                pass

    def show_all_results(self):
        app = App.get_running_app()
        q = (self.vs_in.query or '').strip()
        if not q:
            toast('ابتدا واژه یا بخشی از متن را وارد کنید.', 'جستجو')
            return
        if getattr(self, '_searching', False):
            return
        self._searching = True
        # پیامِ «لطفاً صبر کنید» فقط تا زمانِ آماده‌شدنِ نتایج نمایش داده می‌شود و سپس بسته می‌شود
        self._wait_popup = toast('در حال جستجو…', 'لطفاً صبر کنید')

        def _work():
            try:
                results = app.data.search_all(q, limit=2000)
                err = None
            except Exception as ex:
                results, err = None, str(ex)

            def _done(dt):
                self._searching = False
                _wp = getattr(self, '_wait_popup', None)
                if _wp is not None:
                    try:
                        _wp.dismiss()
                    except Exception:
                        pass
                    self._wait_popup = None
                if err:
                    toast('خطا در جستجو: %s' % err, 'خطا')
                    return
                if not results:
                    toast('هیچ آیه‌ای برای این عبارت پیدا نشد.', 'یافت نشد')
                    return
                self._build_results_popup(q, results)

            Clock.schedule_once(_done, 0)

        threading.Thread(target=_work, daemon=True).start()

    def _build_results_popup(self, q, results):
        # --- اسکرول بی‌پایان: دستهٔ نخست فوری نمایش، بقیه هنگام اسکرول بارگذاری ---
        is_num = q and core.conv(q).strip().isdigit()
        head_txt = ('%d آیه با شمارهٔ «%s»' % (len(results), q)) if is_num \
            else ('%d نتیجه برای «%s» (از بیشترین تطابق)' % (len(results), q))

        popup = Popup(title=P('نتایج جستجو'), size_hint=(0.96, 0.92))
        root = BoxLayout(orientation='vertical', spacing=dp(8), padding=dp(8))
        header = RLabel(head_txt, bold=True, font_size='15sp', halign='center', color=C_GOLD,
                        size_hint_y=None, height=dp(30))
        root.add_widget(header)

        sv = ScrollView(do_scroll_x=False, bar_width=dp(6), scroll_type=['bars', 'content'])
        grid = GridLayout(cols=1, size_hint_y=None, spacing=dp(6), padding=dp(2))
        grid.bind(minimum_height=grid.setter('height'))
        sv.add_widget(grid)
        root.add_widget(sv)

        BATCH = 40
        state = {'shown': 0, 'loading': False}

        def _make_row(r):
            s, a = r['s'], r['a']
            arb = r.get('arb', '') or ''
            pers = r.get('pers', '') or ''
            arb_s = (arb[:75] + '…') if len(arb) > 75 else arb
            pers_s = (pers[:80] + '…') if len(pers) > 80 else pers
            card = ClickCard(bg=(0.11, 0.14, 0.22, 0.98), border=C_GOLD, orientation='vertical',
                             size_hint_y=None, height=dp(124), padding=dp(8), spacing=dp(2))
            card.add_widget(RLabel('سوره %s ، آیه %s' % (s, a), font_size='12sp',
                                   halign='right', color=C_ORANGE, size_hint_y=None, height=dp(20)))
            card.add_widget(RLabel(arb_s, arabic=True, font_size='15sp', halign='right',
                                   color=C_TEXT, size_hint_y=None, height=dp(46)))
            card.add_widget(RLabel(pers_s, font_size='13sp', halign='right',
                                   color=C_MUTED, size_hint_y=None, height=dp(36)))
            card.bind(on_release=lambda inst, ss=s, aa=a: self._pick_result(ss, aa, popup))
            _fade_in(card)
            return card

        def _load_batch(*a):
            start = state['shown']
            chunk = results[start:start + BATCH]

            # رندرِ تکه‌تکهٔ کارت‌ها (هر فریم ۳ کارت) تا پاپ‌آپ هنگام باز شدن گیر نکند
            def _add_incremental_rows(items):
                if not items:
                    state['shown'] = min(len(results), start + BATCH)
                    header.set_text('%s — %d از %d' % (head_txt, state['shown'], len(results)))
                    state['loading'] = False
                    return
                for r in items[:3]:
                    grid.add_widget(_make_row(r))
                Clock.schedule_once(lambda dt: _add_incremental_rows(items[3:]), 0.01)

            _add_incremental_rows(chunk)

        def _on_scroll(inst, val):
            # scroll_y=0 یعنی انتهای فهرست؛ نزدیک انتها که رسیدیم دستهٔ بعدی را بیاور
            if state['loading'] or state['shown'] >= len(results):
                return
            if val <= 0.15:
                state['loading'] = True
                Clock.schedule_once(_load_batch, 0)

        sv.bind(scroll_y=_on_scroll)

        close = PillButton('بستن', bg=(1, 1, 1, 0.12), fg=C_TEXT, size_hint_y=None, height=dp(46))
        close.bind(on_release=lambda *a: popup.dismiss())
        root.add_widget(close)
        popup.content = root
        popup.open()
        _load_batch()   # دستهٔ نخست بلافاصله نمایش داده می‌شود


# ==================================================================
# کارت آیه (مشترک)
# ==================================================================
# ------------------------------------------------------------------
# مدیریت چرخهٔ عمر انیمیشن‌های نوری (رفع نشت انیمیشن و آزادسازی ترد اصلی)
# ------------------------------------------------------------------
_LIVE_GLOWS = []   # [(anim, target_color, widget)]


def _register_glow(anim, target, widget):
    """هر انیمیشن تکرارشونده را ثبت می‌کند تا بعداً بتوان متوقف‌ش‌ کرد."""
    _LIVE_GLOWS.append((anim, target, widget))
    if len(_LIVE_GLOWS) > 80:
        _prune_glows()


def _prune_glows():
    """انیمیشن ویجت‌هایی که دیگر روی صفحه نیستند را لغو می‌کند (ج��وگیری از اشباع ترد UI)."""
    alive = []
    for anim, target, widget in _LIVE_GLOWS:
        try:
            if widget is not None and widget.get_root_window() is None:
                anim.cancel(target)
            else:
                alive.append((anim, target, widget))
        except Exception:
            pass
    _LIVE_GLOWS[:] = alive


def apply_glow(widget, color=None, speed=1.4, width=1.6, hi=0.85, lo=0.12):
    """نور افکتی نرم که دور یک کارت می‌تپد/می‌چرخد.
    width = ضخامتِ خط، hi/lo = بیشینه/کمینهٔ شفافیتِ چشمک."""
    from kivy.graphics import Color as _Color, Line as _Line
    col = color or C_GOLD
    with widget.canvas.after:
        gc = _Color(col[0], col[1], col[2], lo)
        gl = _Line(width=width)

    def _upd(*a):
        try:
            gl.rounded_rectangle = (widget.x + dp(1), widget.y + dp(1),
                                    widget.width - dp(2), widget.height - dp(2), dp(14))
        except Exception:
            pass
    widget.bind(pos=_upd, size=_upd)
    Clock.schedule_once(_upd, 0)
    anim = Animation(a=hi, duration=speed) + Animation(a=lo, duration=speed)
    anim.repeat = True
    anim.start(gc)
    _register_glow(anim, gc, widget)
    return gl


def _make_glow_texture(size=64):
    """یک بافتِ نرمِ نورانی (محو به سمتِ لبه‌ها) برای شهاب می‌سازد."""
    from kivy.graphics.texture import Texture
    import math as _m
    tex = Texture.create(size=(size, size), colorfmt='rgba')
    c = (size - 1) / 2.0
    buf = bytearray(size * size * 4)
    for y in range(size):
        for x in range(size):
            dx = (x - c) / c
            dy = (y - c) / c
            d2 = dx * dx + dy * dy
            a = _m.exp(-4.5 * d2) if d2 <= 1.0 else 0.0
            j = (y * size + x) * 4
            buf[j] = 255
            buf[j + 1] = 255
            buf[j + 2] = 255
            buf[j + 3] = int(255 * a)
    tex.blit_buffer(bytes(buf), colorfmt='rgba', bufferfmt='ubyte')
    tex.mag_filter = 'linear'
    tex.min_filter = 'linear'
    return tex


_GLOW_TEX = None


def _get_glow_tex():
    global _GLOW_TEX
    if _GLOW_TEX is None:
        try:
            _GLOW_TEX = _make_glow_texture()
        except Exception:
            _GLOW_TEX = None
    return _GLOW_TEX


def orbit_dot(widget, color=None, period=5.5, scale=1.0, fps=60, alpha=1.0, trail=30):
    """شهابِ نورانیِ نرم: همهٔ نورها با بافتِ محوشونده (fade) کشیده می���شوند تا لبهٔ تیز نداشته باشند؛ سرِ درخشان + دنبالهٔ گرادیانی که دورِ لبهٔ کادر می‌چرخد.
    alpha = شدتِ روشناییِ کلی (برای کم‌رنگ‌کردنِ المان‌های غیرفعال)، trail = تعدادِ قطعاتِ دنباله (کمتر = سبک‌تر برای CPU)."""
    import math
    from kivy.graphics import Color as _Color, Ellipse as _Ellipse
    col = color or C_GOLD
    r, g, b = col[0], col[1], col[2]
    base = dp(14) * scale
    TRAIL = max(4, int(trail))
    tex = _get_glow_tex()
    segs = []
    cols = []   # (Color, base_alpha) — برای تنظیمِ زندهٔ شدتِ روشنایی
    with widget.canvas.after:
        # دنبالهٔ گرادیانی: frac=0 نوکِ دم (طلایی/محو)، frac=1 نزدیکِ سر (سفیدِ داغ)
        for i in range(TRAIL):
            frac = i / float(TRAIL - 1)
            mix = frac ** 2
            cr = r + (1.0 - r) * mix
            cg = g + (1.0 - g) * mix
            cb = b + (1.0 - b) * mix
            a = 0.9 * (frac ** 1.5)
            _ci = _Color(cr, cg, cb, a * alpha)
            cols.append((_ci, a))
            e = _Ellipse(texture=tex)
            segs.append((e, frac))
        # بلومِ بزرگِ نرمِ دورِ سر
        _cbloom = _Color(r, g, b, 0.5 * alpha)
        cols.append((_cbloom, 0.5))
        bloom = _Ellipse(texture=tex)
        # مغزِ درخشانِ سفیدتاب
        _ccore = _Color(1.0, 0.97, 0.85, 0.95 * alpha)
        cols.append((_ccore, 0.95))
        core = _Ellipse(texture=tex)
        # جرقهٔ سفیدِ داغِ نوکِ سر
        _cspark = _Color(1, 1, 1, 1 * alpha)
        cols.append((_cspark, 1.0))
        spark = _Ellipse(texture=tex)
    state = {'t': 0.0, 'mul': alpha}

    def _pos(tt, w, h, x0, y0):
        per = 2.0 * (w + h)
        d = (tt % 1.0) * per
        if d < w:
            return x0 + d, y0
        d -= w
        if d < h:
            return x0 + w, y0 + d
        d -= h
        if d < w:
            return x0 + w - d, y0 + h
        d -= w
        return x0, y0 + h - d

    def _put(e, cx, cy, sz):
        e.pos = (cx - sz / 2.0, cy - sz / 2.0)
        e.size = (sz, sz)

    def _tick(dt):
        try:
            w, h = widget.width, widget.height
            x0, y0 = widget.x, widget.y
            if w <= 0 or h <= 0:
                return
            state['t'] = (state['t'] + dt / max(0.1, period)) % 1.0
            t = state['t']
            per = 2.0 * (w + h)
            tail_len = (base * 12.0) / per
            for e, frac in segs:
                tt = t - tail_len * (1.0 - frac)
                px, py = _pos(tt, w, h, x0, y0)
                sz = base * (0.5 + 1.7 * (frac ** 1.3))
                _put(e, px, py, sz)
            hx, hy = _pos(t, w, h, x0, y0)
            twinkle = 1.0 + 0.12 * math.sin(t * 6.2832 * 5)
            _put(bloom, hx, hy, base * 4.2 * twinkle)
            _put(core, hx, hy, base * 1.7)
            _put(spark, hx, hy, base * 0.9)
        except Exception:
            pass

    Clock.schedule_interval(_tick, 1.0 / max(1, fps))

    def _set_alpha(m):
        state['mul'] = m
        for _ci, _ba in cols:
            try:
                _ci.a = _ba * m
            except Exception:
                pass

    return {'tick': _tick, 'set_alpha': _set_alpha}



# ==================================================================
# طراحیِ تازهٔ نورها: هالهٔ تشعشعیِ تپنده + گرادیانِ دکمه‌ها + نورِ چرخانِ دایره‌ای
# (نرم، تپنده، و سبک برای CPU/GPU؛ جنسِ نور مثلِ دنبالهٔ نرمِ شهاب)
# ==================================================================

_GRAD_TEX_CACHE = {}


def _make_grad_texture(c_top, c_bottom, size=64):
    # بافتِ گرادیانِ عمودی؛ ردیفِ پایین = c_bottom و ردیفِ بالا = c_top
    from kivy.graphics.texture import Texture
    h = size
    buf = bytearray()
    for i in range(h):
        t = i / float(h - 1)
        r = int(255 * (c_bottom[0] + (c_top[0] - c_bottom[0]) * t))
        g = int(255 * (c_bottom[1] + (c_top[1] - c_bottom[1]) * t))
        b = int(255 * (c_bottom[2] + (c_top[2] - c_bottom[2]) * t))
        a0 = c_bottom[3] if len(c_bottom) > 3 else 1.0
        a1 = c_top[3] if len(c_top) > 3 else 1.0
        a = int(255 * (a0 + (a1 - a0) * t))
        buf += bytes((max(0, min(255, r)), max(0, min(255, g)),
                      max(0, min(255, b)), max(0, min(255, a))))
    tex = Texture.create(size=(1, h), colorfmt='rgba')
    tex.blit_buffer(bytes(buf), colorfmt='rgba', bufferfmt='ubyte')
    tex.wrap = 'clamp_to_edge'
    return tex


def _grad_pair(bg):
    # از یک رنگِ پایه، جفت‌رنگِ گرادیانِ چشم‌نواز می‌سازد (بالا روشن‌تر و کمی چرخشِ فام)
    import colorsys
    r, g, b = bg[0], bg[1], bg[2]
    a = bg[3] if len(bg) > 3 else 1.0
    hh, ll, ss = colorsys.rgb_to_hls(r, g, b)
    l_top = max(0.0, min(1.0, ll * 0.82 + 0.02))
    s_top = min(1.0, ss * 1.10 + 0.05)
    h_top = (hh + 0.02) % 1.0
    top = colorsys.hls_to_rgb(h_top, l_top, s_top)
    l_bot = max(0.0, ll * 0.42)
    s_bot = min(1.0, ss * 1.14 + 0.05)
    h_bot = (hh - 0.015) % 1.0
    bot = colorsys.hls_to_rgb(h_bot, l_bot, s_bot)
    return (top[0], top[1], top[2], a), (bot[0], bot[1], bot[2], a)


def _get_grad_tex(top, bot):
    key = (tuple(round(c, 3) for c in top), tuple(round(c, 3) for c in bot))
    tex = _GRAD_TEX_CACHE.get(key)
    if tex is None:
        try:
            tex = _make_grad_texture(top, bot)
        except Exception:
            tex = None
        _GRAD_TEX_CACHE[key] = tex
    return tex


def _aura_perim_point(f, w, h, x0, y0):
    # نقطه‌ای روی محیطِ مستطیل، متناسب با f از 0 تا 1
    per = 2.0 * (w + h)
    d = (f % 1.0) * per
    if d < w:
        return x0 + d, y0
    d -= w
    if d < h:
        return x0 + w, y0 + d
    d -= h
    if d < w:
        return x0 + w - d, y0 + h
    d -= w
    return x0, y0 + h - d


def pulse_aura(widget, color=None, speed=2.2, thickness=1.0, hi=0.12, lo=0.03,
               alpha=1.0, radius=None, thin=None):
    # حاشیهٔ نورانیِ «بهم‌پیوسته» که دورِ کلِ کادر خاموش/روشن می‌تپد.
    # همان جنسِ نرمِ دنبالهٔ نور است: چند خطِ گوشه‌گردِ هم‌مرکز با شفافیتِ کم
    # روی هم می‌نشینند و یک نوارِ نرمِ پیوسته می‌سازند (نه نقطه‌های جدا).
    # رنگ‌ها تیره‌تر و کم‌نورترند و فقط شفافیت انیمیت می‌شود (بسیار سبک).
    from kivy.graphics import Color as _Color, Line as _Line
    col = color or C_GOLD
    r, g, b = col[0] * 0.78, col[1] * 0.78, col[2] * 0.78
    LAYERS = 5
    st = {'built': False, 'lines': [], 'toplines': [], 'col': None, 'core': None, 'top': None,
          'anim': None, 'anim2': None, 'animt': None, 'rad': radius}

    def _rad():
        if st['rad'] is not None:
            return st['rad']
        return getattr(widget, '_radius', None) or dp(16)

    def _start(mult):
        c = st.get('col')
        cc = st.get('core')
        _h = max(0.0, min(1.0, hi * mult))
        _l = max(0.0, min(1.0, lo * mult))
        if c is not None:
            if st.get('anim') is not None:
                try:
                    st['anim'].cancel(c)
                except Exception:
                    pass
            an = (Animation(a=_h, duration=speed, t='in_out_sine')
                  + Animation(a=_l, duration=speed, t='in_out_sine'))
            an.repeat = True
            st['anim'] = an
            an.start(c)
            _register_glow(an, c, widget)
        if cc is not None:
            if st.get('anim2') is not None:
                try:
                    st['anim2'].cancel(cc)
                except Exception:
                    pass
            _hc = max(0.0, min(1.0, _h * 0.9))
            _lc = max(0.0, min(1.0, _l * 0.9))
            an2 = (Animation(a=_hc, duration=speed, t='in_out_sine')
                   + Animation(a=_lc, duration=speed, t='in_out_sine'))
            an2.repeat = True
            st['anim2'] = an2
            an2.start(cc)
            _register_glow(an2, cc, widget)
        ct = st.get('top')
        if ct is not None:
            if st.get('animt') is not None:
                try:
                    st['animt'].cancel(ct)
                except Exception:
                    pass
            _ht = max(0.0, min(1.0, hi * mult * 1.9))
            _lt = max(0.0, min(1.0, lo * mult * 1.9))
            ant = (Animation(a=_ht, duration=speed, t='in_out_sine')
                   + Animation(a=_lt, duration=speed, t='in_out_sine'))
            ant.repeat = True
            st['animt'] = ant
            ant.start(ct)
            _register_glow(ant, ct, widget)

    def _reposition(*a):
        if not st['built']:
            return
        w, h = widget.width, widget.height
        if w <= 1 or h <= 1:
            return
        rad = _rad()
        for ln in st['lines']:
            ln.rounded_rectangle = (widget.x, widget.y, w, h, rad)
        for ln in st['toplines']:
            ln.points = [widget.x + rad, widget.y + h,
                         widget.x + w - rad, widget.y + h]

    def _build(*a):
        if st['built']:
            return
        w, h = widget.width, widget.height
        if w <= 1 or h <= 1:
            return
        rad = _rad()
        base = dp(0.8)
        step = dp(1.6) * thickness
        with widget.canvas.after:
            st['col'] = _Color(r, g, b, lo * alpha)
            for _i in range(LAYERS):
                lw = base + (LAYERS - 1 - _i) * step
                st['lines'].append(_Line(
                    rounded_rectangle=(widget.x, widget.y, w, h, rad), width=lw))
            # neon core line removed: keep the aura soft, dim and thin
            st['top'] = _Color(min(1.0, r + 0.06), min(1.0, g + 0.06),
                               min(1.0, b + 0.06), lo * alpha)
            for _i in range(3):
                lw = base + (3 - 1 - _i) * step * 1.2
                st['toplines'].append(_Line(
                    points=[widget.x + rad, widget.y + h,
                            widget.x + w - rad, widget.y + h], width=lw))
        st['built'] = True
        _reposition()
        _start(alpha)

    widget.bind(pos=_reposition, size=lambda *a: (_build(), _reposition()))
    Clock.schedule_once(lambda *a: (_build(), _reposition()), 0)

    return {'set_alpha': lambda m: _start(m), 'state': st}


def orbit_ring(widget, color=None, period=3.8, fps=45, trail=20):
    # نورِ شهابیِ نرم که روی مسیرِ «دایره‌ای» دورِ یک دکمه می‌چرخد،
    # همراه با هالهٔ مشکیِ نرم دورِ محیطِ دکمه (به‌جای کادرِ مستطیلی).
    import math
    from kivy.graphics import Color as _Color, Ellipse as _Ellipse
    col = color or C_GOLD
    r, g, b = col[0], col[1], col[2]
    tex = _get_glow_tex()
    base = dp(12)
    TRAIL = max(4, int(trail))
    segs = []
    with widget.canvas.after:
        _Color(0, 0, 0, 0.55)
        halo = _Ellipse(texture=tex)
        for i in range(TRAIL):
            frac = i / float(TRAIL - 1)
            mix = frac ** 2
            cr = r + (1.0 - r) * mix
            cg = g + (1.0 - g) * mix
            cb = b + (1.0 - b) * mix
            _Color(cr, cg, cb, 0.9 * (frac ** 1.5))
            segs.append((_Ellipse(texture=tex), frac))
        _Color(r, g, b, 0.5)
        bloom = _Ellipse(texture=tex)
        _Color(1.0, 0.97, 0.85, 0.95)
        core = _Ellipse(texture=tex)
        _Color(1, 1, 1, 1)
        spark = _Ellipse(texture=tex)
    state = {'t': 0.0}

    def _put(e, cx, cy, sz):
        e.pos = (cx - sz / 2.0, cy - sz / 2.0)
        e.size = (sz, sz)

    def _cpos(f, cx, cy, rad):
        ang = 2.0 * math.pi * (f % 1.0) - math.pi / 2.0
        return cx + rad * math.cos(ang), cy + rad * math.sin(ang)

    def _tick(dt):
        try:
            w, h = widget.width, widget.height
            if w <= 0 or h <= 0:
                return
            cx, cy = widget.center_x, widget.center_y
            rad = max(w, h) / 2.0 + dp(7)
            _put(halo, cx, cy, max(w, h) + dp(34))
            state['t'] = (state['t'] + dt / max(0.1, period)) % 1.0
            t = state['t']
            tail = 0.16
            for e, frac in segs:
                tt = t - tail * (1.0 - frac)
                px, py = _cpos(tt, cx, cy, rad)
                sz = base * (0.5 + 1.5 * (frac ** 1.3))
                _put(e, px, py, sz)
            hx, hy = _cpos(t, cx, cy, rad)
            _put(bloom, hx, hy, base * 3.4)
            _put(core, hx, hy, base * 1.6)
            _put(spark, hx, hy, base * 0.9)
        except Exception:
            pass

    Clock.schedule_interval(_tick, 1.0 / max(1, fps))
    return {'tick': _tick}


def verse_card(mode, s, a, arb, pers, is_seed=False, is_fallback=False, reason='',
               on_save=None, score_text='', on_select=None, selected=False):
    bg = (0.16, 0.13, 0.05, 1) if is_seed else ((0.22, 0.08, 0.08, 1) if is_fallback else (0.10, 0.14, 0.22, 1))
    border = C_GOLD if is_seed else (C_RED if is_fallback else None)
    card = RoundBox(bg=bg, border=border, orientation='vertical', size_hint_y=None,
                    padding=dp(12), spacing=dp(6))
    head = RLabel(f'{mode}   سوره {s} ، آیه {a}', bold=True, font_size='15sp',
                  color=(C_GOLD if is_seed else C_ORANGE), halign='right', size_hint_y=None)
    head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
    card.add_widget(head)
    if score_text:
        sc = RLabel(score_text, font_size='13sp', color=C_MUTED, halign='right', size_hint_y=None)
        sc.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(sc)
    arb_l = RLabel(f'« {arb} »', arabic=True, font_size='20sp', halign='center',
                   color=C_TEXT, size_hint_y=None)
    arb_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
    card.add_widget(arb_l)
    pers_l = RLabel(pers, font_size='14sp', halign='center', color=C_MUTED, size_hint_y=None)
    pers_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
    card.add_widget(pers_l)
    if is_fallback and reason:
        warn = RLabel('' + reason, font_size='12sp', halign='right', color=C_RED, size_hint_y=None)
        warn.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(warn)
    if on_select is not None:
        _sb = PillButton('انتخاب شد ✓' if selected else 'انتخاب این مقصد',
                         bg=C_GREEN if selected else C_BLUE, size_hint_y=None, height=dp(42), font_size='14sp')
        _sb.bind(on_release=lambda *x: on_select())
        card.add_widget(_sb)
    elif on_save:
        btn = PillButton('ثبت این کشف', bg=C_GREEN, size_hint_y=None, height=dp(42), font_size='14sp')
        btn.bind(on_release=lambda *x: on_save())
        card.add_widget(btn)

    apply_glow(card, C_GREEN if selected else (C_GOLD if is_seed else (C_RED if is_fallback else C_BLUE)))
    _fade_in(card)

    def _h(*a):
        total = sum(c.height for c in card.children) + dp(24) + dp(6) * (len(card.children) - 1)
        card.height = total
    Clock.schedule_once(_h, 0)
    card.bind(minimum_height=lambda i, v: setattr(card, 'height', v))
    return card


# ==================================================================
# صفحهٔ ماتریس
# ==================================================================
class MatrixScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پردازش ماتریس آینه‌ای', **kw)
        self.mode = 'normal'          # normal = هفت‌عملگر | rotation = دورانی
        self._seed = (1, 1)
        self._seed_card = None
        self._cards = []
        self._view_a = {}             # idx -> آیهٔ نمایش‌داده‌شده (ناوبری مستقل کارت)
        self._hidden = set()          # idx کارت‌های موقتاً حذف‌شده (ضربدر)
        self._select_mode = None      # None | 'pair' | 'group'
        self._selected = []           # idx کارت‌های انتخاب‌شده (جفت)

        # هدر پردازش در یک خط و با اندازهٔ پویا نمایش داده شود (بدون بهم‌ریختگی)
        self.title_label._fit_single = True

        top = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(8))
        self.mode_btn = PillButton('حالت: ماتریس هفت‌عملگر', bg=C_PURPLE, font_size='13sp')
        self.mode_btn.bind(on_release=lambda *a: self.toggle_mode())
        top.add_widget(self.mode_btn)
        self.body(top)

        selrow = BoxLayout(size_hint_y=None, height=dp(44), spacing=dp(8))
        self.group_btn = PillButton('انتخاب گروهی', bg=C_BLUE, font_size='13sp')
        self.group_btn.bind(on_release=lambda *a: self.toggle_select('group'))
        self.pair_btn = PillButton('انتخاب جفتی', bg=C_BLUE, font_size='13sp')
        self.pair_btn.bind(on_release=lambda *a: self.toggle_select('pair'))
        selrow.add_widget(self.group_btn)
        selrow.add_widget(self.pair_btn)
        self.body(selrow)

        self.reg_bar = BoxLayout(size_hint_y=None, height=dp(0), spacing=dp(8))
        self.body(self.reg_bar)

        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(12), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

        # دکمهٔ شناورِ هوش مصنوعی (شیشه‌ای/نئونی) — گوشهٔ پایین‌راست؛ پس از ساختِ کارت‌ها ظاهر می‌شود
        self.ai_fab = PillButton('AI', bg=(0.10, 0.16, 0.28, 0.75), fg=(1, 1, 1, 1),
                                 radius=dp(30), font_size='18sp',
                                 size_hint=(None, None), size=(dp(60), dp(60)),
                                 pos_hint={'right': 0.95, 'y': 0.06})
        self.ai_fab.bind(on_release=lambda *a: self._ai_analyze())
        self.ai_fab.opacity = 0
        self.ai_fab.disabled = True
        self.root_layout.add_widget(self.ai_fab)
        orbit_ring(self.ai_fab, C_BLUE, period=3.6, fps=45, trail=20)

    # ---------- حالت پردازش ----------
    def toggle_mode(self):
        self.mode = 'rotation' if self.mode == 'normal' else 'normal'
        self.mode_btn.set_text('حالت: پردازش دورانی ارقام بذر' if self.mode == 'rotation'
                               else 'حالت: ماتریس هفت‌عملگر')
        self.show(*self._seed)
        _n = len(self._cards) if getattr(self, '_cards', None) else 0
        toast('حالت پردازش: %s\n(%d کارت ساخته شد)' % (('دورانی ارقام بذر' if self.mode == 'rotation'
                                    else 'ماتریس هفت‌عملگر'), _n), 'حالت')

    def toggle_select(self, kind):
        self._select_mode = None if self._select_mode == kind else kind
        self._selected = []
        self.group_btn.set_text('★ انتخاب گروهی' if self._select_mode == 'group' else 'انتخاب گروهی')
        self.pair_btn.set_text('★ انتخاب جفتی' if self._select_mode == 'pair' else 'انتخاب جفتی')
        self._render()
        _names = {'group': 'انتخاب گروهی', 'pair': 'انتخاب جفتی'}
        _nt = len(self._visible_target_indices())
        if self._select_mode:
            toast('حالت «%s» فعال شد؛ روی %d کارتِ مقصد بزنید.' % (_names[kind], _nt), 'انتخاب')
        else:
            toast('حالت انتخاب خاموش شد.', 'انتخاب')

    def _reset_state(self):
        self._view_a = {}
        self._hidden = set()
        self._selected = []
        self._select_mode = None
        self.group_btn.set_text('انتخاب گروهی')
        self.pair_btn.set_text('انتخاب جفتی')

    # ---------- ناوبری ----------
    def _seed_nav(self, delta):
        app = App.get_running_app()
        s, a = self._seed
        na = a + delta
        if na < 1:
            toast('به ابتدای ��وره رسیدید.', 'ناوبری')
            return
        maxa = app.data.max_a.get(s)
        if maxa and na > maxa:
            toast('به انتهای سوره رسیدید.', 'ناوبری')
            return
        if app.data.get(s, na) is None:
            toast('آیه یافت نشد.', 'ناوبری')
            return
        self.show(s, na)

    def _card_nav(self, idx, delta):
        app = App.get_running_app()
        c = self._cards[idx]
        cur = self._view_a.get(idx, c['a'])
        na = cur + delta
        if na < 1:
            toast('به ابتدای سوره رسیدید.', 'ناوبری')
            return
        maxa = app.data.max_a.get(c['s'])
        if maxa and na > maxa:
            toast('به انتهای سوره رسیدید.', 'ناوبری')
            return
        if app.data.get(c['s'], na) is None:
            toast('آیه یافت نشد.', 'ناوبری')
            return
        self._view_a[idx] = na
        self._render()

    # ---------- ساخت/نمایش ----------
    def show(self, s, a):
        app = App.get_running_app()
        self._reset_state()
        self._seed = (s, a)
        self.title_label.set_text('پردازش : سوره %s و آیه %s' % (s, a))
        if self.mode == 'rotation':
            self._cards = features.generate_rotation_cards(app.data, s, a)
        else:
            self._cards = core.process_matrix(app.data, s, a)
        self._render()

    def _resolved(self, idx):
        """کارت با اعمال ناوبری مستقل (آیهٔ جایگزین) برگردانده می‌شود."""
        app = App.get_running_app()
        c = dict(self._cards[idx])
        if idx in self._view_a:
            na = self._view_a[idx]
            v = app.data.get(c['s'], na) or {}
            c['a'] = na
            c['arb'] = v.get('arb', '')
            c['pers'] = v.get('pers', '')
        return c

    def _visible_target_indices(self):
        return [i for i, c in enumerate(self._cards)
                if c.get('kind') == 'target' and i not in self._hidden]

    def _render(self):
        self.list.clear_widgets()
        _prune_glows()   # لغو انیمیشن‌های کارت‌های حذف‌شده تا ترد UI آزاد بماند
        self._update_reg_bar()
        if not self._cards:
            self.list.add_widget(empty_state('داده‌ای برای این بذر یافت نشد',
                                             hint='شمارهٔ سوره و آیه را بررسی کن و دوباره تلاش کن'))
            self._set_fab(False)
            return
        self._seed_card = self._cards[0]
        self.list.add_widget(self._make_card(0, is_seed=True))
        for i in range(1, len(self._cards)):
            if self._cards[i].get('kind') != 'target':
                continue
            if i in self._hidden:
                continue
            self.list.add_widget(self._make_card(i, is_seed=False))
        self._set_fab(len(self._visible_target_indices()) > 0)

    def _set_fab(self, show):
        fab = getattr(self, 'ai_fab', None)
        if fab is None:
            return
        fab.disabled = not show
        Animation(opacity=1 if show else 0, d=0.25).start(fab)

    def _ai_analyze(self):
        if not self._cards:
            toast('اول یک بذر را پردازش کن.', 'هوش مصنوعی')
            return
        seed = self._resolved(0)
        targets = []
        for i in range(1, len(self._cards)):
            if self._cards[i].get('kind') != 'target':
                continue
            if i in self._hidden:
                continue
            targets.append(self._resolved(i))
        if not targets:
            toast('کارتی برای تحلیل نیست.', 'هوش مصنوعی')
            return
        msgs = ai_manager.build_matrix_messages(seed, targets)
        show_ai_result_popup('تحلیل یکپارچهٔ ماتریس', msgs,
                             subtitle='بذر: سوره %s ، آیه %s  —  %d مقصد' % (seed['s'], seed['a'], len(targets)))

    def _make_card(self, idx, is_seed=False):
        c = self._resolved(idx)
        is_fb = c.get('is_fallback', False)
        bg = (0.16, 0.13, 0.05, 1) if is_seed else ((0.22, 0.08, 0.08, 1) if is_fb else (0.10, 0.14, 0.22, 1))
        border = C_GOLD if is_seed else (C_RED if is_fb else None)
        card = RoundBox(bg=bg, border=border, orientation='vertical', size_hint_y=None,
                        padding=dp(12), spacing=dp(6))

        # ردیف بالای کارت: نشان عملگر (چپ) + ضربدر حذف (راست)
        toprow = BoxLayout(size_hint_y=None, height=dp(28), spacing=dp(6))
        if is_seed:
            badge = _neon_badge('بذر', C_GOLD, size=(dp(48), dp(26)))
        elif self.mode == 'rotation':
            _rot = (c.get('mode', '').split('←')[0].split('(')[0].strip() or 'چرخش')
            badge = _neon_badge(_rot, C_PURPLE, size=(dp(104), dp(26)))
        else:
            badge = _neon_badge(op_of({'mode': c.get('mode', '')}), C_BLUE, size=(dp(42), dp(26)))
        toprow.add_widget(badge)
        toprow.add_widget(Widget())
        if not is_seed:
            xb = PillButton('حذف موقت', bg=C_RED, fg=(1, 1, 1, 1), size_hint_x=None, width=dp(100), font_size='12sp')
            xb.bind(on_release=lambda *a, i=idx: self._hide_card(i))
            toprow.add_widget(xb)
        card.add_widget(toprow)

        head = RLabel('%s   سوره %s ، آیه %s' % (c.get('mode', ''), c['s'], c['a']), bold=True,
                      font_size='15sp', color=(C_GOLD if is_seed else C_ORANGE), halign='right', size_hint_y=None)
        head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(head)
        arb_l = RLabel('« %s »' % c.get('arb', ''), arabic=True, font_size='20sp', halign='center',
                       color=C_TEXT, size_hint_y=None)
        arb_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
        card.add_widget(arb_l)
        pers_l = RLabel(c.get('pers', ''), font_size='14sp', halign='center', color=C_MUTED, size_hint_y=None)
        pers_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        card.add_widget(pers_l)
        if is_fb and c.get('reason'):
            warn = RLabel(c.get('reason', ''), font_size='12sp', halign='right', color=C_RED, size_hint_y=None)
            warn.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
            card.add_widget(warn)

        # ردیف آیهٔ قبل/بعد (برای همهٔ کارت‌ها)
        nav = BoxLayout(size_hint_y=None, height=dp(40), spacing=dp(8))
        b_next = PillButton('آیهٔ بعد ▶', bg=(1, 1, 1, 0.14), font_size='13sp')
        b_prev = PillButton('◀ آیهٔ قبل', bg=(1, 1, 1, 0.14), font_size='13sp')
        if is_seed:
            b_next.bind(on_release=lambda *a: self._seed_nav(+1))
            b_prev.bind(on_release=lambda *a: self._seed_nav(-1))
        else:
            b_next.bind(on_release=lambda *a, i=idx: self._card_nav(i, +1))
            b_prev.bind(on_release=lambda *a, i=idx: self._card_nav(i, -1))
        nav.add_widget(b_next)
        nav.add_widget(b_prev)
        card.add_widget(nav)

        # دکمهٔ کنش (فقط کارت‌های مقصد)
        if not is_seed:
            if self._select_mode == 'pair':
                picked = idx in self._selected
                order = (self._selected.index(idx) + 1) if picked else None
                lbl = ('انتخاب‌شده (%d)' % order) if picked else 'انتخاب برای جفت'
                sb = PillButton(lbl, bg=C_GREEN if picked else C_BLUE, size_hint_y=None, height=dp(42), font_size='14sp')
                sb.bind(on_release=lambda *a, i=idx: self._pick_pair(i))
                card.add_widget(sb)
            elif self._select_mode == 'group':
                info = RLabel('در انتخاب گروهی ★ (با «حذف موقت» بردارید)', font_size='12sp',
                              halign='center', color=C_GREEN, size_hint_y=None)
                info.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
                card.add_widget(info)
            else:
                rb = PillButton('ثبت این کشف با بذر', bg=C_GREEN, size_hint_y=None, height=dp(42), font_size='14sp')
                rb.bind(on_release=lambda *a, i=idx: self._register_single(i))
                card.add_widget(rb)

        # نور دور کارت
        gcol = C_GOLD if is_seed else (C_GREEN if (self.mode != 'rotation' and self._select_mode == 'group')
                                       else (C_RED if is_fb else C_BLUE))
        apply_glow(card, gcol)

        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v))
        return card

    def _hide_card(self, idx):
        self._hidden.add(idx)
        self._selected = [i for i in self._selected if i != idx]
        self._render()

    def _restore_hidden(self):
        self._hidden = set()
        self._render()

    def _pick_pair(self, idx):
        if idx in self._selected:
            self._selected = [i for i in self._selected if i != idx]
        else:
            if len(self._selected) >= 2:
                toast('برای جفت فقط دو مقصد انتخاب کنید.', 'جفت')
                return
            self._selected.append(idx)
        self._render()

    def _update_reg_bar(self):
        self.reg_bar.clear_widgets()
        widgets = []
        if self._select_mode == 'group':
            n = len(self._visible_target_indices())
            b = PillButton('ثبت کشف گروهی (%d مقصد)' % n, bg=C_GREEN, font_size='14sp')
            b.bind(on_release=lambda *a: self._register_group())
            widgets.append(b)
        elif self._select_mode == 'pair':
            n = len(self._selected)
            b = PillButton('ثبت جفت عملگری (%d از ۲)' % n, bg=C_GREEN, font_size='14sp')
            b.bind(on_release=lambda *a: self._register_pair())
            widgets.append(b)
        if self._hidden:
            rb = PillButton('بازگردانی حذف‌شده‌ها (%d)' % len(self._hidden), bg=C_ORANGE, font_size='13sp')
            rb.bind(on_release=lambda *a: self._restore_hidden())
            widgets.append(rb)
        if not widgets:
            self.reg_bar.height = dp(0)
            return
        self.reg_bar.height = dp(48)
        for w in widgets:
            self.reg_bar.add_widget(w)

    # ---------- ثبت کشف ----------
    def _register_single(self, idx):
        app = App.get_running_app()
        app.add_discovery(self._seed_card, self._resolved(idx))

    def _register_group(self):
        app = App.get_running_app()
        idxs = self._visible_target_indices()
        if not idxs:
            toast('هیچ کارتی برای ثبت گروهی نمانده است.', 'گروهی')
            return
        targets = [self._resolved(i) for i in idxs]
        app.add_group_discovery(self._seed_card, targets)
        self._reset_state()
        self._render()

    def _register_pair(self):
        app = App.get_running_app()
        if len(self._selected) != 2:
            toast('برای جفت دقیقاً دو مقصد انتخاب کنید.', 'جفت')
            return
        ta = self._resolved(self._selected[0])
        tb = self._resolved(self._selected[1])
        app.add_pair_discovery(self._seed_card, ta, tb)
        self._reset_state()
        self._render()


# ==================================================================
# صفحهٔ پیش‌بینی (معنا / اعداد)
# ==================================================================
class PredictScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پیش‌بینی آینه', **kw)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def show(self, s, a, kind):
        app = App.get_running_app()
        self.list.clear_widgets()
        seed = app.data.get(s, a)
        self.title_label.set_text('پیش‌بینی ' + ('معنایی' if kind == 'semantic' else 'عددی'))
        # کارت بذر
        self.list.add_widget(verse_card('بذر ساختاری', s, a, seed['arb'], seed['pers'], is_seed=True))
        if kind == 'semantic':
            preds = core.predict_mirror(app.data, s, a, seed['arb'])
            ctx = {'s': s, 'a': a, 'seed_arb': seed['arb'], 'seed_pers': seed['pers'], 'preds': []}
            for _r, (op, ts, ta, score, status, is_fb, msg) in enumerate(preds, 1):
                d = app.data.get(ts, ta) or {}
                ctx['preds'].append({'rank': _r, 'op_code': op_of({'mode': op}), 's': ts, 'a': ta,
                                     'arb': d.get('arb', ''), 'pers': d.get('pers', ''), 'is_fallback': is_fb})
            self._pred_ctx = ctx
            actions = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
            pbtn = PillButton('کپی پرامپت هوش مصنوعی (۷ قانونی)', bg=C_PURPLE, font_size='12sp')
            pbtn.bind(on_release=lambda *a: self._copy_pred_prompt())
            hbtn = PillButton('نمایش گزارش HTML', bg=C_GREEN, font_size='12sp')
            hbtn.bind(on_release=lambda *a: self._open_html())
            actions.add_widget(pbtn)
            actions.add_widget(hbtn)
            self.list.add_widget(actions)
            for rank, (op, ts, ta, score, status, is_fb, msg) in enumerate(preds, 1):
                d = app.data.get(ts, ta) or {}
                sc = 'رتبه %d | امتیاز: %.0f٪ | %s' % (rank, score, status)
                if rank <= 3:
                    sc += '  •  نیازمند تحلیل هوش مصنوعی یا انسانی'
                seed_card = {'mode': op, 's': s, 'a': a, 'arb': seed['arb'], 'pers': seed['pers']}
                tgt = {'mode': op, 's': ts, 'a': ta, 'arb': d.get('arb', ''), 'pers': d.get('pers', ''),
                       'is_fallback': is_fb, 'reason': msg}
                w = verse_card(op, ts, ta, d.get('arb', ''), d.get('pers', ''),
                               is_fallback=is_fb, reason=(msg if is_fb else ''),
                               score_text=sc,
                               on_save=(lambda sd=seed_card, cc=tgt: app.add_discovery(sd, cc)))
                w.opacity = 0
                self.list.add_widget(w)
                Animation(opacity=1, duration=0.3).start(w)
        else:
            preds = core.predict_mirror_numeric(app.data, s, a)
            if not preds:
                self.list.add_widget(empty_state('مقصد معتبری با الگوریتمِ عددی یافت نشد',
                                                 hint='این بذر با روشِ عددی نتیجه نداد؛ روشِ معنایی را امتحان کن'))
                return
            for (op, ts, ta, prio, detail, is_fb, msg) in preds:
                d = app.data.get(ts, ta) or {}
                sc = f'اولویت {prio} | {detail}'
                seed_card = {'mode': op, 's': s, 'a': a, 'arb': seed['arb'], 'pers': seed['pers']}
                tgt = {'mode': op, 's': ts, 'a': ta, 'arb': d.get('arb', ''), 'pers': d.get('pers', ''),
                       'is_fallback': is_fb, 'reason': msg}
                w = verse_card(op, ts, ta, d.get('arb', ''), d.get('pers', ''),
                               is_fallback=is_fb, reason=(msg if is_fb else ''), score_text=sc,
                               on_save=(lambda sd=seed_card, cc=tgt: app.add_discovery(sd, cc)))
                w.opacity = 0
                self.list.add_widget(w)
                Animation(opacity=1, duration=0.3).start(w)

    def _copy_pred_prompt(self):
        ctx = getattr(self, '_pred_ctx', None)
        if not ctx:
            toast('ابتدا یک پیش‌بینی معنایی اجرا کنید.', 'هوش مصنوعی')
            return
        try:
            from kivy.core.clipboard import Clipboard
            Clipboard.copy(features.build_prediction_prompt(
                ctx['s'], ctx['a'], ctx['seed_arb'], ctx['seed_pers'], ctx['preds']))
            toast('پرامپت پیش‌بینی معنایی کپی شد؛ آن را در چت هوش مصنوعی بچسبانید.', 'هوش مصنوعی')
        except Exception:
            toast('کپی ممکن نشد.', 'خطا')

    def _open_html(self):
        app = App.get_running_app()
        content = BoxLayout(orientation='vertical', padding=dp(12), spacing=dp(10))
        content.add_widget(RLabel('کد HTML گزارش هوش مصنوعی را اینجا بچسبانید:',
                                  font_size='14sp', color=C_GOLD, halign='center', size_hint_y=None, height=dp(50)))
        ti = PlainInput(multiline=True, font_size='12sp', size_hint_y=1,
                       background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
        content.add_widget(ti)
        pop = Popup(title=P('نمایش گزارش HTML'), content=content, size_hint=(0.96, 0.85),
                    title_font='ui', title_align='center', separator_color=C_GOLD)
        row = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
        bshow = PillButton('نمایش گزارش', bg=C_GREEN)

        def _show(*a):
            html = ti.text or ''
            if not html.strip():
                toast('متن HTML خالی است.', 'خطا')
                return
            show_html_in_app(html)   # نمایش داخل خودِ برنامه (بدون نیاز به مرورگر)
        bshow.bind(on_release=_show)
        bclose = PillButton('بستن', bg=C_RED)
        bclose.bind(on_release=pop.dismiss)
        row.add_widget(bshow)
        row.add_widget(bclose)
        content.add_widget(row)
        pop.open()


# ==================================================================
# لابراتوار کشفیات
# ==================================================================
OPERATORS = [
    ('T1', 'جابجایی خالص بذر'),
    ('T2', 'تقارن درجا کامل'),
    ('T3', 'تقارن درجا فقط سوره'),
    ('T4', 'تقارن درجا فقط آیه'),
    ('T5', 'جابجایی + تقارن کامل'),
    ('T6', 'جابجایی + تقارن فقط سوره'),
    ('T7', 'جابجایی + تقارن فقط آیه'),
    ('OTHER', 'گروهی و سایر'),
]


def op_of(item):
    if item.get('op_key'):
        return item['op_key']
    m = str(item.get('mode', ''))
    if 'خالص' in m:
        return 'T1'
    if 'ضربدری' in m:
        return 'T5'
    if 'جابجایی' in m and 'کامل' in m:
        return 'T5'
    if 'تقارن درجا کامل' in m:
        return 'T2'
    if 'جابجایی' in m and 'فقط سوره' in m:
        return 'T6'
    if 'جابجایی' in m and 'فقط آیه' in m:
        return 'T7'
    if 'فقط سوره' in m:
        return 'T3'
    if 'فقط آ��ه' in m:
        return 'T4'
    return 'OTHER'


def discovery_key(item):
    """کلید یکتای یک کشف برای تطبیق نئون/آخرین کشف."""
    if 'all_targets' in item:
        return ('G', item.get('seed_s'), item.get('seed_a'), item.get('mode'),
                tuple((t.get('s'), t.get('a')) for t in item.get('all_targets', [])))
    return (item.get('seed_s'), item.get('seed_a'), item.get('target_s'),
            item.get('target_a'), item.get('mode'))


def lab_section_of(item):
    """بخش لابراتوار: تردیدی‌ها جدا، گروهی‌ها جدا، بقیه زیر عملگر خودشان."""
    if item.get('is_doubtful'):
        return 'DOUBT'
    if item.get('mode') == 'گروهی':
        return 'GROUP'
    return op_of(item)


LAB_SECTIONS = [(k, t) for k, t in OPERATORS if k != 'OTHER'] + [
    ('GROUP', 'کشفیات گروهی'),
    ('DOUBT', 'کشفیات تردیدی'),
]


def _neon_badge(text, color=None, size=None):
    """نشان کوچک نئونی چشمک‌زن (مثلاً T1..T7) برای گوشهٔ کارت."""
    from kivy.graphics import Color as _C, RoundedRectangle as _RR, Line as _L
    if color is None:
        color = C_BLUE
    if size is None:
        size = (dp(42), dp(26))
    lbl = Label(text=rtl(text), font_name='ui', bold=True, font_size='13sp',
                size_hint=(None, None), size=size, color=(1, 1, 1, 1))
    with lbl.canvas.before:
        _C(color[0], color[1], color[2], 0.20)
        bg = _RR(radius=[dp(8)])
        lc = _C(color[0], color[1], color[2], 0.95)
        ln = _L(width=1.5)

    def _u(*a):
        bg.pos = lbl.pos
        bg.size = lbl.size
        ln.rounded_rectangle = (lbl.x, lbl.y, lbl.width, lbl.height, dp(8))
    lbl.bind(pos=_u, size=_u)
    anim = Animation(a=1.0, duration=0.6) + Animation(a=0.12, duration=0.6)
    anim.repeat = True
    anim.start(lc)
    _register_glow(anim, lc, lbl)
    return lbl


def _auto_label(text, arabic=False, **kw):
    kw.setdefault('size_hint_y', None)
    lbl = RLabel(text, arabic=arabic, **kw)
    lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
    return lbl


def _verse_block(border, s, a, arb, pers):
    c = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=border, orientation='vertical',
                 size_hint_y=None, padding=dp(10), spacing=dp(4))
    c.add_widget(_auto_label('سوره %s ، آیه %s' % (s, a), font_size='12sp', color=C_MUTED, halign='right'))
    c.add_widget(_auto_label('« %s »' % (arb or ''), arabic=True, font_size='18sp', color=C_TEXT, halign='center'))
    c.add_widget(_auto_label('ترجمه: ' + (pers or ''), font_size='13sp', color=C_MUTED, halign='right'))
    c.bind(minimum_height=lambda i, v: setattr(c, 'height', v + dp(24)))
    return c


def generate_default_analysis(e):
    """متن تحلیل پیش‌فرض یک کشف (هم‌راستا با نسخهٔ ویندوز):
    فقط نام عملگر و یک یادآوری کوتاه؛ بدون متن آیه و ترجمه، تا باکس تحلیل
    فقط چیزی باشد که خودِ کاربر می‌نویسد."""
    return (
        'این کشف بر اساس عملگر «%s» پیشنهاد شده است.\n\n'
        '(لطفاً تحلیل دقیق خود را از رابطهٔ معنایی این دو آیه یادداشت کنید…)'
        % (e.get('mode', '') or '—'))


def _add_analysis_field(content, initial_text=''):
    """فیلدِ ویرایشِ تحلیل همراه با پیش‌نمایشِ زندهٔ درست.

    زیر موتورِ sdl2 (که RTL بومی ندارد)، متنِ داخلِ کادرِ ویرایش هنگامِ تایپ
    به‌صورتِ منطقی نشان داده می‌شود تا مکان‌نما/انتخاب/حذفِ وسطِ متن سالم کار کند؛
    و بالای آن یک «پیش‌نمایشِ درست» می‌گذاریم که همان لحظه شکلِ نهاییِ راست‌به‌چپ را نشان می‌دهد.
    زیر موتورِ Pango نیازی به پیش‌نمایش نیست و فقط کادرِ ویرایش نمایش داده می‌شود.
    """
    content.add_widget(RLabel('تحلیل شما:', font_size='15sp', size_hint_y=None, height=dp(26)))
    ti = PersianEditor(multiline=True, font_size='15sp',
                       size_hint_y=1, background_color=(1, 1, 1, 0.95),
                       foreground_color=(0.05, 0.08, 0.14, 1))
    ti.set_logical(initial_text or '')
    content.add_widget(ti)
    if not _native_text_shaping():
        content.add_widget(RLabel('پیش‌نمایشِ درست (زنده):', font_size='12sp', color=C_GOLD,
                                  halign='right', size_hint_y=None, height=dp(20)))
        pv_box = RoundBox(bg=(1, 1, 1, 0.95), radius=10, size_hint_y=None,
                          height=dp(96), padding=dp(6))
        pv_scroll = ScrollView(bar_width=dp(4), do_scroll_x=False)
        prev = RLabel(initial_text or '', font_size='15sp', halign='right', valign='top',
                      color=(0.05, 0.08, 0.14, 1), size_hint_y=None)
        prev.bind(texture_size=lambda inst, ts: setattr(inst, 'height', ts[1] + dp(6)))
        pv_scroll.add_widget(prev)
        pv_box.add_widget(pv_scroll)
        content.add_widget(pv_box)
        ti.attach_preview(prev)
    return ti


def open_note_editor(item, source='lab', title='ویرایش تحلیل', intro=None, on_saved=None, saved_msg='ذخیره شد ✓'):
    """پنجرهٔ ثبت/ویرایش تحلیل یک کشف (با برچسب و وضعیت تردید)."""
    app = App.get_running_app()
    content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
    if intro:
        content.add_widget(RLabel(intro, font_size='14sp', color=C_GOLD, halign='center',
                                  size_hint_y=None, height=dp(50)))
    # فیلد متن + پیش‌نمایشِ زندهٔ درست (سازگار با موتورِ sdl2)
    ti = _add_analysis_field(content, item.get('note', ''))
    # دکمهٔ دستیارِ هوش مصنوعی: تولیدِ متنِ تحلیلیِ کوتاه و تزریق در فیلد (کاربر بعداً ویرایش می‌کند)
    ai_row = BoxLayout(size_hint_y=None, height=dp(44), spacing=dp(8))
    ai_btn = PillButton('✦ کمکِ هوش مصنوعی در تحلیل', bg=C_PURPLE, fg=HOME_FG, font_size='14sp')

    def _ai_fill(*a):
        seed = {'s': item.get('seed_s'), 'a': item.get('seed_a'),
                'arb': item.get('seed_arb', ''), 'pers': item.get('seed_pers', '')}
        target = {'s': item.get('target_s'), 'a': item.get('target_a'),
                  'arb': item.get('target_arb', ''), 'pers': item.get('target_pers', '')}
        operator = item.get('mode', '') or 'نامشخص'
        msgs = ai_manager.build_note_messages(seed, target, operator)
        ai_btn.disabled = True
        ai_btn.set_text('در حال نوشتن تحلیل...')

        def _done(full):
            ai_btn.disabled = False
            ai_btn.set_text('✦ بازنویسی با هوش مصنوعی')
            txt = (full or '').strip()
            if txt:
                ti.set_logical(txt)
            else:
                toast('پاسخی دریافت نشد.', 'هوش مصنوعی', kind='warn')

        def _err(msg):
            ai_btn.disabled = False
            ai_btn.set_text('✦ کمکِ هوش مصنوعی در تحلیل')
            toast(msg, 'هوش م��نوع��', kind='error')

        ai_manager.chat(msgs, on_done=_done, on_error=_err,
                        stream=False, temperature=0.5, max_tokens=600)

    ai_btn.bind(on_release=_ai_fill)
    ai_row.add_widget(ai_btn)
    content.add_widget(ai_row)
    tags = list(app.get_all_tags())
    _get_tags = _tag_multiselect(content, list(app.get_all_tags()), item.get('relation_type'), 'برچسب تحلیلی (رفتار شبکه) — هر تعداد که خواستی انتخاب کن:')
    state = {'d': bool(item.get('is_doubtful', False))}
    btog = PillButton('وضعیت: تردیدی' if state['d'] else 'وضعیت: مطمئن',
                      bg=C_ORANGE if state['d'] else C_GREEN, size_hint_y=None, height=dp(44), font_size='14sp')
    def _tog(*a):
        state['d'] = not state['d']
        btog.set_text('وضعیت: تردیدی' if state['d'] else 'وضعیت: مطمئن')
        # رنگ پس‌زمینه هم همگام شود تا تغییر وضعیت کاملاً دیده شود
        btog._bg = list(C_ORANGE if state['d'] else C_GREEN)
        btog._state()
    btog.bind(on_release=_tog)
    content.add_widget(btog)
    ep = Popup(title=P(title), content=content, size_hint=(0.96, 0.94),
               title_font='ui', title_align='center', separator_color=C_GOLD)
    row2 = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
    sv = PillButton('ذخیره', bg=C_GREEN)
    def _sv(*a):
        item['note'] = ti.query
        item['relation_type'] = _get_tags()
        item['is_doubtful'] = state['d']
        if source == 'lab':
            app.save_favs()
        else:
            app.save_featured()
        ep.dismiss()
        if on_saved:
            on_saved()
        toast(saved_msg, 'ذخیره')
    sv.bind(on_release=_sv)
    cn = PillButton('انصراف', bg=C_RED)
    cn.bind(on_release=ep.dismiss)
    row2.add_widget(sv)
    row2.add_widget(cn)
    content.add_widget(row2)
    ep.open()


def show_group_discovery(item, source='lab', screen=None):
    app = App.get_running_app()
    root = BoxLayout(orientation='vertical', padding=dp(8), spacing=dp(8))
    scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
    box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
    box.bind(minimum_height=box.setter('height'))
    p = Popup(title=P('جزئیات کشف گروهی'), content=root, size_hint=(0.96, 0.92),
              title_font='ui', title_align='center', separator_color=C_GOLD)

    def _refresh_parent():
        if screen is not None and hasattr(screen, 'refresh'):
            screen.refresh()

    def _render():
        box.clear_widgets()
        box.add_widget(_auto_label('نوع: ' + str(item.get('mode', '')), bold=True,
                                   font_size='15sp', color=C_GOLD, halign='right'))
        box.add_widget(_auto_label('آیهٔ مبدأ (بذر)', bold=True, font_size='14sp', color=C_GOLD, halign='right'))
        box.add_widget(_verse_block(C_PURPLE, item.get('seed_s'), item.get('seed_a'),
                                    item.get('seed_arb', ''), item.get('seed_pers', '')))
        targets = item.get('all_targets', [])
        box.add_widget(_auto_label('مقصدها (%d):' % len(targets), bold=True,
                                   font_size='14sp', color=C_ORANGE, halign='right'))
        for idx, t in enumerate(targets):
            box.add_widget(_verse_block(C_GOLD, t.get('s'), t.get('a'), t.get('arb', ''), t.get('pers', '')))
            if t.get('operator'):
                box.add_widget(_auto_label('عملگر: ' + str(t.get('operator')), font_size='12sp',
                                           color=C_MUTED, halign='right'))
            if source == 'lab':
                rmb = PillButton('حذف این مقصد از گروه', bg=C_RED, size_hint_y=None, height=dp(38), font_size='12sp')

                def _rm(*a, i=idx):
                    def _do():
                        ok, msg = app.remove_target_from_group(item, i)
                        toast(msg, 'گروهی' if ok else 'خطا')
                        if ok:
                            if not item.get('all_targets'):
                                p.dismiss()
                                _refresh_parent()
                            else:
                                _render()
                    confirm('این مقصد از گروه حذف شود؟', _do, 'حذف مقصد')
                rmb.bind(on_release=_rm)
                box.add_widget(rmb)
        box.add_widget(_auto_label('رفتار شبکه: ' + str(item.get('relation_type', 'نامشخص')),
                                   font_size='14sp', color=C_ORANGE, halign='right'))
        _note = item.get('note', '')
        box.add_widget(_auto_label('تحلیل شما: ' + (_note if _note else '—'),
                                   font_size='14sp', color=C_TEXT, halign='right'))

    _render()
    scroll.add_widget(box)
    root.add_widget(scroll)

    def _edit(*a):
        open_note_editor(item, source, title='ویرایش تحلیل', on_saved=_render)

    def _delete(*a):
        def _do():
            lst = app.favs if source == 'lab' else app.featured
            if item in lst:
                lst.remove(item)
            if source == 'lab':
                app.save_favs()
            else:
                app.save_featured()
            p.dismiss()
            _refresh_parent()
        confirm('کل این کشف گروهی حذف شود؟', _do, 'حذف کشف')

    def _to_featured(*a):
        app.featured.append(dict(item))
        app.save_featured()
        toast('به گلچین افزوده شد.', 'گلچین')

    grid = GridLayout(cols=2, size_hint_y=None, height=dp(104), spacing=dp(8))
    if source == 'lab':
        bb = PillButton('افزودن به گلچین', bg=C_GOLD, font_size='14sp')
        bb.bind(on_release=_to_featured)
        grid.add_widget(bb)
    be = PillButton('ویرایش تحلیل', bg=C_BLUE, font_size='14sp')
    be.bind(on_release=_edit)
    grid.add_widget(be)
    bd = PillButton('حذف کشف' if source == 'lab' else 'حذف از گلچین', bg=C_RED, font_size='14sp')
    bd.bind(on_release=_delete)
    grid.add_widget(bd)
    root.add_widget(grid)
    close = PillButton('بستن', bg=(1, 1, 1, 0.14), size_hint_y=None, height=dp(46))
    close.bind(on_release=p.dismiss)
    root.add_widget(close)
    p.open()


def show_discovery(item, source='lab', screen=None):
    if 'all_targets' in item:
        return show_group_discovery(item, source, screen)
    app = App.get_running_app()
    root = BoxLayout(orientation='vertical', padding=dp(8), spacing=dp(8))
    scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
    box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
    box.bind(minimum_height=box.setter('height'))
    box.add_widget(_auto_label('آیهٔ مبدأ (بذر)', bold=True, font_size='15sp', color=C_GOLD, halign='right'))
    box.add_widget(_verse_block(C_PURPLE, item.get('seed_s'), item.get('seed_a'),
                                item.get('seed_arb', ''), item.get('seed_pers', '')))
    box.add_widget(_auto_label('گرهٔ کشف‌شده: ' + str(item.get('mode', '')), bold=True,
                               font_size='14sp', color=C_GOLD, halign='right'))
    box.add_widget(_verse_block(C_GOLD, item.get('target_s'), item.get('target_a'),
                                item.get('target_arb', ''), item.get('target_pers', '')))
    box.add_widget(_auto_label('رفتار شبکه: ' + str(item.get('relation_type', 'نامشخص')),
                               font_size='14sp', color=C_ORANGE, halign='right'))
    _note = item.get('note', '')
    box.add_widget(_auto_label('تحلیل شما: ' + (_note if _note else '—'),
                               font_size='14sp', color=C_TEXT, halign='right'))
    scroll.add_widget(box)
    root.add_widget(scroll)
    p = Popup(title=P('جزئیات کشف'), content=root, size_hint=(0.96, 0.9),
              title_font='ui', title_align='center', separator_color=C_GOLD)

    def _refresh_parent():
        if screen is not None and hasattr(screen, 'refresh'):
            screen.refresh()

    def _copy(*a):
        try:
            from kivy.core.clipboard import Clipboard
            txt = ('[%s] رفتار: %s\n' % (item.get('mode', ''), item.get('relation_type', ''))
                   + 'مبدأ (سوره %s آیه %s): %s\n%s\n' % (item.get('seed_s'), item.get('seed_a'),
                     item.get('seed_arb', ''), item.get('seed_pers', ''))
                   + 'مقصد (سوره %s آیه %s): %s\n%s\n' % (item.get('target_s'), item.get('target_a'),
                     item.get('target_arb', ''), item.get('target_pers', ''))
                   + ('تحلیل: ' + _note if _note else ''))
            Clipboard.copy(txt)
            toast('اطلاعات کشف کپی شد.', 'کپی')
        except Exception:
            toast('کپی ممکن نشد.', 'خطا')

    def _delete(*a):
        def _do():
            lst = app.favs if source == 'lab' else app.featured
            key = (item.get('seed_s'), item.get('seed_a'), item.get('target_s'),
                   item.get('target_a'), item.get('mode'))
            for i, it in enumerate(lst):
                if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'),
                        it.get('target_a'), it.get('mode')) == key:
                    del lst[i]
                    break
            if source == 'lab':
                app.save_favs()
            else:
                app.save_featured()
            p.dismiss()
            _refresh_parent()
        confirm('این کشف حذف شود؟', _do, 'حذف کشف')

    def _to_featured(*a):
        app.add_featured(item)

    def _edit(*a):
        content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
        # فیلد متن + پیش‌نمایشِ زندهٔ درست (سازگار با موتورِ sdl2)
        ti = _add_analysis_field(content, item.get('note', ''))
        _ai_row = BoxLayout(size_hint_y=None, height=dp(44), spacing=dp(8))
        _ai_btn = PillButton('✦ کمکِ هوش مصنوعی در تحلیل', bg=C_PURPLE, fg=HOME_FG, font_size='14sp')

        def _ai_fill(*a):
            seed = {'s': item.get('seed_s'), 'a': item.get('seed_a'),
                    'arb': item.get('seed_arb', ''), 'pers': item.get('seed_pers', '')}
            target = {'s': item.get('target_s'), 'a': item.get('target_a'),
                      'arb': item.get('target_arb', ''), 'pers': item.get('target_pers', '')}
            operator = item.get('mode', '') or 'نامشخص'
            msgs = ai_manager.build_note_messages(seed, target, operator)
            _ai_btn.disabled = True
            _ai_btn.set_text('در حال نوشتن تحلیل...')

            def _done(full):
                _ai_btn.disabled = False
                _ai_btn.set_text('✦ بازنویسی با هوش مصنوعی')
                txt = (full or '').strip()
                if txt:
                    ti.set_logical(txt)
                else:
                    toast('پاسخی دریافت نشد.', 'هوش مصنوعی', kind='warn')

            def _err(msg):
                _ai_btn.disabled = False
                _ai_btn.set_text('✦ کمکِ هوش مصنوعی در تحلیل')
                toast(msg, 'هوش مصنوعی', kind='error')

            ai_manager.chat(msgs, on_done=_done, on_error=_err,
                            stream=False, temperature=0.5, max_tokens=600)

        _ai_btn.bind(on_release=_ai_fill)
        _ai_row.add_widget(_ai_btn)
        content.add_widget(_ai_row)
        _get_tags2 = _tag_multiselect(content, list(app.get_all_tags()), item.get('relation_type'), 'برچسب (رفتار شبکه) — چند انتخابی:')
        # امکان تغییر وضعیت (تردیدی ↔ مطمئن) برای همین کشف — مخصوصاً برای خارج‌کردنِ کشف از بخشِ تردیدی‌ها
        _st = {'d': bool(item.get('is_doubtful', False))}
        b_status = PillButton('وضعیت: ت��دیدی' if _st['d'] else 'وضعیت: مطمئن',
                              bg=C_ORANGE if _st['d'] else C_GREEN, size_hint_y=None, height=dp(44), font_size='14sp')

        def _tog_status(*a):
            _st['d'] = not _st['d']
            b_status.set_text('وضعیت: تردیدی' if _st['d'] else 'وضعیت: مطمئن')
            b_status._bg = list(C_ORANGE if _st['d'] else C_GREEN)
            b_status._state()
        b_status.bind(on_release=_tog_status)
        content.add_widget(b_status)
        ep = Popup(title=P('ویرایش تحلیل'), content=content, size_hint=(0.96, 0.94),
                   title_font='ui', title_align='center', separator_color=C_GOLD)
        row2 = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
        sv = PillButton('ذخیره', bg=C_GREEN)

        def _sv(*a):
            item['note'] = ti.query
            item['relation_type'] = _get_tags2()
            item['is_doubtful'] = _st['d']
            if source == 'lab':
                app.save_favs()
            else:
                app.save_featured()
            ep.dismiss()
            p.dismiss()
            _refresh_parent()
            toast('تغییرات ذخیره شد ✓', 'ذخیره')
        sv.bind(on_release=_sv)
        cn = PillButton('انصراف', bg=C_RED)
        cn.bind(on_release=ep.dismiss)
        row2.add_widget(sv)
        row2.add_widget(cn)
        content.add_widget(row2)
        ep.open()

    grid = GridLayout(cols=2, size_hint_y=None, height=dp(104), spacing=dp(8))
    if source == 'lab':
        bb = PillButton('افزودن به گلچین', bg=C_GOLD, font_size='14sp')
        bb.bind(on_release=_to_featured)
        grid.add_widget(bb)
    be = PillButton('ویرایش تحلیل', bg=C_BLUE, font_size='14sp')
    be.bind(on_release=_edit)
    grid.add_widget(be)
    bd = PillButton('حذف کشف' if source == 'lab' else 'حذف از ��ل��ین', bg=C_RED, font_size='14sp')
    bd.bind(on_release=_delete)
    grid.add_widget(bd)
    bc = PillButton('کپی اطلاعات', bg=C_GREEN, font_size='14sp')
    bc.bind(on_release=_copy)
    grid.add_widget(bc)
    root.add_widget(grid)
    close = PillButton('بستن', bg=(1, 1, 1, 0.14), size_hint_y=None, height=dp(46))
    close.bind(on_release=p.dismiss)
    root.add_widget(close)
    p.open()


class LabScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='لابراتوار کشفیات', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
        b_all = PillButton('افزودن همه به گلچین', bg=C_GOLD, font_size='14sp')
        b_all.bind(on_release=lambda *a: self.add_all_featured())
        top.add_widget(b_all)
        self.count_lbl = RLabel('', font_size='14sp', halign='center', color=C_MUTED)
        top.add_widget(self.count_lbl)
        self.body(top)
        row2 = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(8))
        b_ai = PillButton('کپی پرامپت هوش مصنوعی', bg=C_PURPLE, font_size='13sp')
        b_ai.bind(on_release=lambda *a: self.copy_ai_prompt())
        b_html = PillButton('نمایش گزارش HTML', bg=C_BLUE, font_size='13sp')
        b_html.bind(on_release=lambda *a: self.open_html_viewer())
        row2.add_widget(b_ai)
        row2.add_widget(b_html)
        self.body(row2)
        row3 = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(8))
        b_clear = PillButton('پاک کردن کل کشفیات', bg=C_RED, font_size='13sp')
        b_clear.bind(on_release=lambda *a: self.clear_all())
        row3.add_widget(b_clear)
        self.body(row3)
        self.body(RLabel('برای دیدن کشف‌های هر عملگر، روی آن بزنید.', font_size='13sp',
                         halign='center', color=C_MUTED, size_hint_y=None, height=dp(26)))
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.grid = GridLayout(cols=1, size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.grid.bind(minimum_height=self.grid.setter('height'))
        self.scroll.add_widget(self.grid)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.grid.clear_widgets()
        self.count_lbl.set_text('%d کشف' % len(app.favs))
        if not app.favs:
            self.grid.add_widget(empty_state('هنوز کشفی ثبت نشده است',
                                             hint='از صفحهٔ اصلی یک بذر انتخاب کن و اولین کشفت را ثبت کن'))
            return
        counts = {}
        for it in app.favs:
            k = lab_section_of(it)
            counts[k] = counts.get(k, 0) + 1
        last_sec = getattr(app, 'last_discovery_section', None)
        ordered = list(LAB_SECTIONS)
        if last_sec:
            ordered.sort(key=lambda kt: 0 if kt[0] == last_sec else 1)
        for key, title in ordered:
            n = counts.get(key, 0)
            if n == 0:
                continue
            if key == last_sec:
                label = '● %s  (%d کشف)  — جدید' % (title, n)
            else:
                label = '%s  (%d کشف)' % (title, n)
            b = PillButton(label, bg=(0.16, 0.13, 0.05, 1),
                           size_hint_y=None, height=dp(64), font_size='16sp')
            b.bind(on_release=lambda inst, k=key, t=title: self.open_op(k, t))
            self.grid.add_widget(b)
            if key == last_sec:
                apply_glow(b, C_GREEN, speed=0.6, width=5.0, hi=1.0, lo=0.25)

    def open_op(self, key, title):
        scr = self.manager.get_screen('operator')
        scr.load('lab', key, title)
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = 'operator'

    def add_all_featured(self):
        app = App.get_running_app()
        n = app.add_all_featured()
        toast('%d کشف به گلچین اضافه شد.' % n, 'گلچین')

    def clear_all(self):
        app = App.get_running_app()

        def _do():
            app.favs = []
            app.save_favs()
            self.refresh()
        confirm('کل کشفیات لابراتوار پاک شود؟', _do, 'پاک کردن کل')

    def copy_ai_prompt(self):
        app = App.get_running_app()
        if not app.favs:
            toast('ابتدا چند کشف ثبت کنید.', 'هوش مصنوعی')
            return
        try:
            from kivy.core.clipboard import Clipboard
            Clipboard.copy(features.build_semantic_prompt(app.favs))
            toast('پرامپت تحلیل معنایی کپی شد؛ آن را در چت هوش مصنوعی بچسبانید.', 'هوش مصنوعی')
        except Exception:
            toast('کپی ممکن نشد.', 'خطا')

    def open_html_viewer(self):
        app = App.get_running_app()
        content = BoxLayout(orientation='vertical', padding=dp(12), spacing=dp(10))
        content.add_widget(RLabel('کد HTML گزارش هوش مصنوعی ��ا اینجا ��چسبانید:',
                                  font_size='14sp', color=C_GOLD, halign='center', size_hint_y=None, height=dp(50)))
        ti = PlainInput(multiline=True, font_size='12sp', size_hint_y=1,
                       background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
        content.add_widget(ti)
        pop = Popup(title=P('نمایش گزارش HTML'), content=content, size_hint=(0.96, 0.85),
                    title_font='ui', title_align='center', separator_color=C_GOLD)
        row = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
        bshow = PillButton('نمایش گزارش', bg=C_GREEN)

        def _show(*a):
            html = ti.text or ''
            if not html.strip():
                toast('متن HTML خالی است.', 'خطا')
                return
            show_html_in_app(html)   # نمایش داخل خودِ برنامه (بدون نیاز به مرورگر)
        bshow.bind(on_release=_show)
        bclose = PillButton('بستن', bg=C_RED)
        bclose.bind(on_release=pop.dismiss)
        row.add_widget(bshow)
        row.add_widget(bclose)
        content.add_widget(row)
        pop.open()


class OperatorScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='کشف‌های عملگر', **kw)
        self.title_label._fit_single = True
        self.source = 'lab'
        self.op_key = 'T1'
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def load(self, source, op_key, title):
        self.source = source
        self.op_key = op_key
        src_name = 'لابراتوار' if source == 'lab' else 'گلچین'
        self.title_label.set_text('%s — %s' % (src_name, title))
        self.refresh()

    def refresh(self):
        app = App.get_running_app()
        self.list.clear_widgets()
        # هر بار که این بخش دوباره باز می‌شود، رندرِ قبلی (اگر در جریان بود) لغو شود
        self._op_gen = getattr(self, '_op_gen', 0) + 1
        gen = self._op_gen
        if getattr(self, '_op_render_ev', None) is not None:
            self._op_render_ev.cancel()
            self._op_render_ev = None
        items = app.favs if self.source == 'lab' else app.featured
        key_fn = lab_section_of if self.source == 'lab' else op_of
        matched = [it for it in items if key_fn(it) == self.op_key]
        # آخرین کشفِ ثبت‌شده باید اولِ لیست بیاید (favs به‌ترتیبِ افزوده‌شدن است)
        if self.source == 'lab':
            matched.reverse()
        if not matched:
            self.list.add_widget(empty_state('در این بخش هنوز کشفی نیست',
                                             hint='کشف‌های ثبت‌شدهٔ این عملگر اینجا نمایش داده می‌شوند'))
            return
        last_key = getattr(app, 'last_discovery_key', None)
        # رندرِ تکه‌تکه: هر فریم چند کارت، تا صفحه فوری باز شود و رابط کاربری قفل نشود
        queue = list(matched)

        def _add_batch(_dt):
            if gen != getattr(self, '_op_gen', 0):
                return  # این بخش دوباره باز شده؛ رندرِ قدیمی را رها کن
            for _ in range(6):
                if not queue:
                    self._op_render_ev = None
                    return
                it = queue.pop(0)
                card = self._group_card(it) if 'all_targets' in it else self._card(it)
                self.list.add_widget(card)
                if last_key is not None and discovery_key(it) == last_key:
                    apply_glow(card, C_GREEN, speed=0.6, width=5.0, hi=1.0, lo=0.25)
            if queue:
                self._op_render_ev = Clock.schedule_once(_add_batch, 0)
            else:
                self._op_render_ev = None

        self._op_render_ev = Clock.schedule_once(_add_batch, 0)

    def _card(self, item):
        border = C_GOLD if self.source == 'featured' else C_BLUE
        card = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=border, orientation='vertical',
                        size_hint_y=None, padding=dp(10), spacing=dp(4))
        pair = RLabel('سوره %s:%s     سوره %s:%s' % (item.get('seed_s'), item.get('seed_a'),
                      item.get('target_s'), item.get('target_a')),
                      font_size='13sp', color=C_MUTED, halign='center', size_hint_y=None)
        pair.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(pair)
        a2 = RLabel('« %s »' % (item.get('target_arb', '')), arabic=True, font_size='16sp',
                    halign='center', color=C_TEXT, size_hint_y=None)
        a2.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        card.add_widget(a2)
        _tp = (item.get('target_pers', '') or '').strip()
        if _tp:
            tr = RLabel(_tp, font_size='13sp', halign='center', color=C_MUTED, size_hint_y=None)
            tr.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
            card.add_widget(tr)
        extra = '  (تردیدی)' if item.get('is_doubtful') else ''
        rel = RLabel('رفتار: %s%s' % (item.get('relation_type', 'نامشخص'), extra), font_size='12sp',
                     color=C_ORANGE, halign='right', size_hint_y=None)
        rel.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(rel)
        btn = PillButton('مشاهده و ویرایش', bg=C_BLUE, size_hint_y=None, height=dp(40), font_size='13sp')
        btn.bind(on_release=lambda *a, it=item: show_discovery(it, self.source, self))
        card.add_widget(btn)
        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v + dp(20)))
        _fade_in(card)
        return card

    def _group_card(self, item):
        border = C_GOLD if self.source == 'featured' else C_PURPLE
        card = RoundBox(bg=(0.12, 0.10, 0.20, 1), border=border, orientation='vertical',
                        size_hint_y=None, padding=dp(10), spacing=dp(4))
        head = RLabel('%s — سوره %s:%s ← %d مقصد' % (item.get('mode', ''), item.get('seed_s'),
                      item.get('seed_a'), len(item.get('all_targets', []))),
                      font_size='13sp', color=C_ORANGE, halign='center', size_hint_y=None)
        head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(head)
        a1 = RLabel('« %s »' % item.get('seed_arb', ''), arabic=True, font_size='15sp',
                    halign='center', color=C_TEXT, size_hint_y=None)
        a1.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        card.add_widget(a1)
        _sp = (item.get('seed_pers', '') or '').strip()
        if _sp:
            tr = RLabel(_sp, font_size='13sp', halign='center', color=C_MUTED, size_hint_y=None)
            tr.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
            card.add_widget(tr)
        extra = '  (تردیدی)' if item.get('is_doubtful') else ''
        rel = RLabel('رفتار: %s%s' % (item.get('relation_type', 'نامشخص'), extra), font_size='12sp',
                     color=C_ORANGE, halign='right', size_hint_y=None)
        rel.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(rel)
        btn = PillButton('مشاهده و ویرایش', bg=C_BLUE, size_hint_y=None, height=dp(40), font_size='13sp')
        btn.bind(on_release=lambda *a, it=item: show_discovery(it, self.source, self))
        card.add_widget(btn)
        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v + dp(20)))
        _fade_in(card)
        return card

    def go_back(self, *a):
        self.manager.transition = SlideTransition(direction='right')
        self.manager.current = 'lab' if self.source == 'lab' else 'featured'


# ==================================================================
# گلچین برگزیده
# ==================================================================
class FeaturedScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='گلچین برگزیده', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
        b_word = PillButton('اشتراک‌گذاری فایل JSON', bg=C_BLUE, font_size='12sp')
        # متن داخلِ خودِ دکمه بشکند و جا بگیرد (به‌جای بیرون‌زدن)
        b_word.halign = 'center'
        b_word.valign = 'middle'
        b_word.bind(size=lambda i, v: setattr(i, 'text_size', (i.width - dp(10), v[1])))
        b_word.bind(on_release=lambda *a: self.share_json())
        b_save = PillButton('ذخیره در گوشی', bg=C_GOLD, fg=(0.05, 0.08, 0.14, 1), font_size='13sp')
        b_save.bind(on_release=lambda *a: self.save_json_device())
        b_clear = PillButton('پاک کردن کل', bg=C_RED, font_size='13sp')
        b_clear.bind(on_release=lambda *a: self.clear_all())
        top.add_widget(b_word)
        top.add_widget(b_save)
        top.add_widget(b_clear)
        self.body(top)
        b_ai = PillButton('گفتگو با هوش مصنوعی', bg=C_PURPLE, fg=HOME_FG, font_size='14sp',
                          size_hint_y=None, height=dp(46))
        b_ai.bind(on_release=lambda *a: self._go_chat())
        self.body(b_ai)
        self.body(RLabel('برای دیدن نمونه‌های هر عملگر، روی آن بزنید.', font_size='13sp',
                         halign='center', color=C_MUTED, size_hint_y=None, height=dp(26)))
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.grid = GridLayout(cols=1, size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.grid.bind(minimum_height=self.grid.setter('height'))
        self.scroll.add_widget(self.grid)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.grid.clear_widgets()
        self.title_label.set_text('گلچین برگزیده (%d)' % len(app.featured))
        if not app.featured:
            self.grid.add_widget(empty_state('گلچین برگزیده‌ات خالی است',
                                             hint='از لابراتوار، کشف‌های برگزیده را به گلچین اضافه کن'))
            return
        counts = {}
        for it in app.featured:
            k = op_of(it)
            counts[k] = counts.get(k, 0) + 1
        for key, title in OPERATORS:
            n = counts.get(key, 0)
            if n == 0:
                continue
            b = PillButton('%s  (%d نمونه)' % (title, n), bg=(0.16, 0.13, 0.05, 1),
                           size_hint_y=None, height=dp(64), font_size='16sp')
            b.bind(on_release=lambda inst, k=key, t=title: self.open_op(k, t))
            self.grid.add_widget(b)

    def _go_chat(self):
        scr = self.manager.get_screen('chat')
        if hasattr(scr, 'refresh'):
            scr.refresh()
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = 'chat'

    def open_op(self, key, title):
        scr = self.manager.get_screen('operator')
        scr.load('featured', key, title)
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = 'operator'

    def clear_all(self):
        app = App.get_running_app()

        def _do():
            app.featured = []
            app.save_featured()
            self.refresh()
        confirm('کل لیست گلچین پاک شود؟', _do, 'پاک کردن کل')

    def share_json(self):
        app = App.get_running_app()
        path = app.export_featured_json()
        if not path:
            toast('گلچین خالی است یا خطایی رخ داد.', 'خطا')
            return
        import share_util
        share_util.save_file_to_device(path, on_done=lambda ok, msg: toast(msg, 'گلچین' if ok else 'خطا'),
                                       mime='application/json', then_share=True)

    def save_json_device(self):
        app = App.get_running_app()
        path = app.export_featured_json()
        if not path:
            toast('گلچین خالی است یا خطایی رخ داد.', 'خطا')
            return
        import share_util
        share_util.save_file_to_device(path, on_done=lambda ok, msg: toast(msg, 'گلچین' if ok else 'خطا'),
                                       mime='application/json', then_share=False)


# ==================================================================
# جستجو
# ==================================================================
class SearchResultCard(RecycleDataViewBehavior, ClickCard):
    """ردیفِ فشرده و سبکِ نتیجهٔ جستجو. برای دیدنِ جزئیاتِ کامل روی ردیف لمس کنید."""
    def __init__(self, **kw):
        kw.setdefault('orientation', 'vertical')
        kw.setdefault('size_hint_y', None)
        kw.setdefault('height', dp(62))
        kw.setdefault('padding', [dp(10), dp(6)])
        kw.setdefault('spacing', dp(2))
        kw.setdefault('bg', (0.10, 0.14, 0.22, 1))
        kw.setdefault('radius', 12)
        super().__init__(**kw)
        self._src = ''
        self._item = None
        self._top = RLabel('', font_size='12sp', halign='right',
                           size_hint_y=None, height=dp(20))
        self._arb = RLabel('', arabic=True, font_size='14sp', halign='right',
                           color=C_TEXT, size_hint_y=None, height=dp(28))
        self.add_widget(self._top)
        self.add_widget(self._arb)
        self.bind(on_release=self._open)

    def refresh_view_attrs(self, rv, index, data):
        # فقط متنِ دو لیبلِ موجود به‌روز می‌شود (بدون ساختِ ویدجتِ جدید)
        self._src = data.get('source', '')
        self._item = data.get('item')
        it = self._item or {}
        tag = 'گلچین' if self._src == 'featured' else 'لابراتوار'
        self._top.color = C_GOLD if self._src == 'featured' else C_MUTED
        self._top.set_text('[%s]  سوره %s:%s ↔ سوره %s:%s  %s' % (
            tag, it.get('seed_s'), it.get('seed_a'),
            it.get('target_s'), it.get('target_a'), it.get('mode', '') or ''))
        arb = (it.get('target_arb', '') or it.get('seed_arb', '') or '')
        arb_s = (arb[:45] + '…') if len(arb) > 45 else arb
        self._arb.set_text('« %s »' % arb_s)
        return super().refresh_view_attrs(rv, index, data)

    def _open(self, *a):
        if self._item is None:
            return
        try:
            screen = App.get_running_app().sm.get_screen('search')
        except Exception:
            screen = None
        show_discovery(self._item, self._src, screen)


class SearchScreen(BaseScreen):
    # حداکثر تعداد نتیجهٔ نمایشی (RecycleView فقط کارت‌های قابل‌مشاهده را می‌سازد)
    MAX_RESULTS = 200

    def __init__(self, **kw):
        super().__init__(title='جستجو در کشفیات', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        # با هر تغییرِ متن، جستجوی زنده (با کمی تأخیر) اجرا می‌شود؛ دکمه هم کار می‌کند
        self.q = PersianTextInput(hint_text=P('جستجو در لابراتوار و گلچین...'),
                                  on_change=lambda *a: self._schedule_search(),
                                  font_size='15sp',
                                  background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1),
                                  padding=[dp(8), dp(14)])
        self.q.bind(on_text_validate=lambda *a: self._run_search())
        b = PillButton('جستجو', bg=C_BLUE, size_hint_x=None, width=dp(96), font_size='14sp')
        b.bind(on_release=lambda *a: self._run_search())
        top.add_widget(self.q)
        top.add_widget(b)
        self.body(top)
        self.info = RLabel('', font_size='13sp', halign='center', color=C_MUTED,
                           size_hint_y=None, height=dp(24))
        self.body(self.info)
        # لیستِ نتایج: همان روشِ مطمئنِ لابراتوار (ScrollView + BoxLayout)؛ کارت‌ها فشرده و سبک‌اند
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(6), scroll_type=['bars', 'content'])
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)
        self._index = []          # ایندکسِ نرمال‌شدهٔ کشفیات: [(source, item, haystack)]
        self._index_sig = None    # امضای دیتاست برای کش: (len(favs), len(featured))
        self._search_ev = None    # نوبتِ زمان‌بندی‌شدهٔ جستجو (برای تأخیر و لغو)

    def refresh(self):
        # فهرست را اینجا (روی ترد اصلی) نمی‌سازیم؛ ساختِ سنگین در اولین جستجو و در ترد پس‌زمینه انجام می‌شود
        self.list.clear_widgets()
        self.info.set_text('واژهٔ عربی/فارسی یا شمارهٔ سوره/آیه را بنویسید؛ نتایج خودکار می‌آید.')

    def _build_index(self):
        app = App.get_running_app()
        idx = []
        for source, coll in (('lab', app.favs), ('featured', app.featured)):
            for it in coll:
                idx.append((source, it, self._make_hay(it)))
        self._index = idx

    def _make_hay(self, it):
        # متنِ جست‌وجوپذیرِ هر کشف: متنِ آیات، ترجمه‌ها، ��ملگر، برچسب، تحلیل و شماره‌ها
        parts = [it.get('seed_arb', ''), it.get('target_arb', ''),
                 it.get('seed_pers', ''), it.get('target_pers', ''),
                 it.get('mode', ''), it.get('relation_type', ''), it.get('note', '')]
        nums = 'سوره %s آیه %s سوره %s آیه %s %s:%s %s:%s' % (
            it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'),
            it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'))
        combined = ' '.join(str(x) for x in parts) + ' ' + nums
        # همان نرمال‌سازیِ جستجوی آیه: حذف اعراب + یکسان‌سازیِ حروف (ا/ی/ک/ه) + یکسان‌سازیِ ارقام
        return core.normalize_text(core.conv(combined))

    def _schedule_search(self, *a):
        # هم‌سبکِ جستجوی آیه: کارِ جستجو را بی‌درنگ اجرا نمی‌کنیم؛ با کمی تأخیر و لغوِ نوبتِ قبلی،
        # تا با تایپِ هر حرف، اپ سنگین/هنگ نشود.
        if self._search_ev is not None:
            self._search_ev.cancel()
        self._search_ev = Clock.schedule_once(self._run_search, 0.35)

    def _run_search(self, *a):
        # لغو زمان‌بندی‌های قبلی هنگام تایپِ سریع
        if self._search_ev is not None:
            self._search_ev.cancel()
            self._search_ev = None

        raw = (self.q.query or '').strip()
        term = core.normalize_text(core.conv(raw))

        if len(term) < 2:
            self.list.clear_widgets()
            self.info.set_text('حداقل ۲ حرف یا یک عدد وارد کنید.')
            return

        app = App.get_running_app()
        sig = (len(app.favs), len(app.featured))
        need_build = (getattr(self, '_index_sig', None) != sig) or (not self._index)
        self.info.set_text('در حال آماده‌سازی و جستجو...' if need_build else 'در حال جستجو...')

        # هر جستجو یک شمارهٔ نوبت می‌گیرد؛ نتیجهٔ نوبت‌های قدیمی (تایپِ سریع) نادیده گرفته می‌شود
        self._search_gen = getattr(self, '_search_gen', 0) + 1
        gen = self._search_gen

        # همهٔ کارِ سنگین (ساختِ فهرست + فیلتر) در ترد پس‌زمینه؛ نمایش با RecycleView آنی است
        def _do_search():
            try:
                if need_build:
                    idx = []
                    for source, coll in (('lab', app.favs), ('featured', app.featured)):
                        for it in coll:
                            idx.append((source, it, self._make_hay(it)))
                    self._index = idx
                    self._index_sig = sig
                matches = [(src, it) for src, it, hay in self._index if term in hay]
            except Exception:
                matches = []
            Clock.schedule_once(lambda dt: _on_search_done(matches), 0)

        def _on_search_done(matches):
            if gen != getattr(self, '_search_gen', 0):
                return  # نتیجهٔ یک جستجوی قدیمی‌ست؛ رهایش کن
            self.list.clear_widgets()
            if not matches:
                self.info.set_text('کشفی همخوان با «%s» یافت نشد.' % raw)
                return
            shown = matches[:self.MAX_RESULTS]
            if len(matches) > self.MAX_RESULTS:
                self.info.set_text('%d نتیجه — نمایشِ %d موردِ اول (برای کمتر شدن، عبارتِ دقیق‌تری بنویسید).'
                                   % (len(matches), self.MAX_RESULTS))
            else:
                self.info.set_text('%d نتیجه در لابراتوار و گلچین' % len(matches))

            # رندرِ تکه‌تکه: هر فریم چند کارت تا رابط کاربری روان بماند (بدون هنگ حتی با نتایج زیاد)
            queue = list(shown)

            def _add_batch(_dt):
                if gen != getattr(self, '_search_gen', 0):
                    return
                for _ in range(6):
                    if not queue:
                        return
                    src, it = queue.pop(0)
                    self.list.add_widget(self._result_row(src, it))
                if queue:
                    self._render_ev = Clock.schedule_once(_add_batch, 0)

            if getattr(self, '_render_ev', None) is not None:
                self._render_ev.cancel()
            self._render_ev = Clock.schedule_once(_add_batch, 0)

        threading.Thread(target=_do_search, daemon=True).start()

    def _result_row(self, src, it):
        # کارتِ فشرده و سبکِ نتیجه؛ با کلیک، همان پنجرهٔ ویرایشِ لابراتوار باز می‌شود
        card = ClickCard(bg=(0.10, 0.14, 0.22, 1),
                         border=(C_GOLD if src == 'featured' else C_BLUE),
                         orientation='vertical', size_hint_y=None, height=dp(66),
                         padding=[dp(10), dp(6)], spacing=dp(2), radius=12)
        tag = 'گلچین' if src == 'featured' else 'لابراتوار'
        top = RLabel('[%s]  سوره %s:%s ↔ سوره %s:%s  %s' % (
            tag, it.get('seed_s'), it.get('seed_a'),
            it.get('target_s'), it.get('target_a'), it.get('mode', '') or ''),
            font_size='12sp', halign='right',
            color=(C_GOLD if src == 'featured' else C_MUTED),
            size_hint_y=None, height=dp(20))
        card.add_widget(top)
        arb = (it.get('target_arb', '') or it.get('seed_arb', '') or '')
        arb_s = (arb[:45] + '…') if len(arb) > 45 else arb
        a = RLabel('« %s »' % arb_s, arabic=True, font_size='14sp', halign='right',
                   color=C_TEXT, size_hint_y=None, height=dp(30))
        card.add_widget(a)
        card.bind(on_release=lambda *_a, s=src, i=it: show_discovery(i, s, self))
        return card


# ==================================================================
# مدیریت برچسب‌ها
# ==================================================================
class TagsScreen(BaseScreen):
    DEFAULT = ["پژواک واژگانی", "تقارن ساختاری/نحوی", "هم‌آوایی صوتی", "تقابل کامل", "مکمل و بسط‌دهنده", "علت و معلول", "پرسش و پاسخ", "دیالوگ متقاطع", "گفت و گو", "وعده و تحقق وعده", "تمثیل موازی", "دادگاه و اعتراف", "تسبیح کائنات", "تکمیل پازل داستانی", "زاویه ��ید متفاوت"]

    def __init__(self, **kw):
        super().__init__(title='مدیریت برچسب‌ها', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        self.q = PersianTextInput(hint_text=P('برچسب جدید...'), multiline=False, font_name='ui',
                           font_size='15sp', background_color=(1, 1, 1, 0.95),
                           foreground_color=(0.05, 0.08, 0.14, 1), padding=[dp(8), dp(14)])
        b = PillButton('افزودن', bg=C_GREEN, size_hint_x=None, width=dp(110), font_size='14sp')
        b.bind(on_release=lambda *a: self.add_tag())
        top.add_widget(self.q)
        top.add_widget(b)
        self.body(top)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.list.clear_widgets()
        for tag in app.get_all_tags():
            row = RoundBox(bg=(0.10, 0.14, 0.22, 1), orientation='horizontal', size_hint_y=None,
                           height=dp(50), padding=dp(10), spacing=dp(8))
            row.add_widget(RLabel(tag, font_size='15sp', halign='right'))
            if tag not in self.DEFAULT and tag != 'نامشخص':
                b = PillButton('', bg=C_RED, size_hint_x=None, width=dp(56), font_size='14sp')
                b.bind(on_release=lambda x, t=tag: self.del_tag(t))
                row.add_widget(b)
            self.list.add_widget(row)

    def add_tag(self):
        app = App.get_running_app()
        t = self.q.query.strip()
        if not t:
            return
        if t in app.get_all_tags():
            toast('این برچسب قبلاً وجود دارد.', 'تک��ار')
            return
        app.user_tags.append(t)
        app.save_user_tags()
        self.q.clear_logical()
        self.refresh()

    def del_tag(self, tag):
        app = App.get_running_app()
        def _do():
            if tag in app.user_tags:
                app.user_tags.remove(tag)
            for it in app.favs:
                if it.get('relation_type') == tag:
                    it['relation_type'] = 'نامشخص'
            app.save_user_tags()
            app.save_favs()
            self.refresh()
        confirm(f'برچسب «{tag}» حذف شود؟ (کشفیات مربوطه به نامشخص تغییر می‌کند)', _do, 'حذف برچسب')


# ==================================================================
# رسانه و معرفی
# ==================================================================
class MediaScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='رسانه و معرفی', **kw)
        self.sound = None
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(12), padding=dp(6))
        box.bind(minimum_height=box.setter('height'))

        box.add_widget(RLabel('چند کلام از طراح', bold=True, font_size='18sp', color=C_GOLD,
                              halign='center', size_hint_y=None, height=dp(34)))
        b_designer = PillButton('چند کلام از طراح در مورد اپلیکیشن', bg=C_ORANGE, size_hint_y=None, height=dp(52), font_size='14sp')
        b_designer.bind(on_release=lambda *a: self.play('designer.mp3'))
        box.add_widget(b_designer)
        b_stop = PillButton('توقف', bg=C_RED, size_hint_y=None, height=dp(52))
        b_stop.bind(on_release=lambda *a: self.stop())
        box.add_widget(b_stop)
        box.add_widget(Widget(size_hint_y=None, height=dp(20)))

        scroll.add_widget(box)
        self.body(scroll)

    def refresh(self):
        pass

    def play(self, name):
        try:
            from kivy.core.audio import SoundLoader
        except Exception as e:
            print('audio provider error:', e)
            toast('پخش صدا در این دستگاه در دسترس نیست.', 'خطا')
            return
        self.stop()
        path = asset(name)
        if not os.path.exists(path):
            toast('فایل صوتی یافت نشد.', 'خطا')
            return
        try:
            self.sound = SoundLoader.load(path)
        except Exception as e:
            print('audio load error:', e)
            self.sound = None
            toast('پخش این فایل صوتی ممکن نشد.', 'خطا')
            return
        if not self.sound:
            toast('پخش این فایل صوتی ممکن نشد.', 'خطا')
            return
        try:
            self.sound.play()
        except Exception as e:
            print('audio play error:', e)
            toast('پخش این فایل صوتی ممکن نشد.', 'خطا')

    def stop(self):
        if self.sound:
            self.sound.stop()
            self.sound = None

    def go_back(self, *a):
        self.stop()
        super().go_back()


# ==================================================================
# پشتیبان و بازیابی
# ==================================================================
class BackupScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پشتیبان و بازیابی', **kw)
        box = BoxLayout(orientation='vertical', spacing=dp(14), padding=dp(16))
        box.add_widget(RLabel('از کشفیات، گلچین و برچسب‌های خود نسخهٔ پشتیبان بگیرید یا آن را بازیابی کنید.',
                              font_size='15sp', halign='center', color=C_MUTED, size_hint_y=None, height=dp(70)))
        b_backup = PillButton('پشتیبان‌گیری و اشتراک‌گذاری فایل JSON', bg=C_GREEN, size_hint_y=None, height=dp(56))
        b_backup.bind(on_release=lambda *a: self.backup(share=True))
        box.add_widget(b_backup)
        b_save = PillButton('ذخیرهٔ فایل پشتیبان در حافظهٔ گوشی', bg=C_GOLD, fg=(0.05, 0.08, 0.14, 1), size_hint_y=None, height=dp(56))
        b_save.bind(on_release=lambda *a: self.backup(share=False))
        box.add_widget(b_save)
        b_import = PillButton('بارگذاری/بازیابی از فایل JSON', bg=C_BLUE, size_hint_y=None, height=dp(56))
        b_import.bind(on_release=lambda *a: self.open_import())
        box.add_widget(b_import)
        self.info = RLabel('', font_size='13sp', halign='center', color=C_GOLD, size_hint_y=None, height=dp(60))
        box.add_widget(self.info)
        box.add_widget(Widget())
        self.body(box)

    def refresh(self):
        self.info.set_text('')

    def backup(self, share=True):
        app = App.get_running_app()
        try:
            path = app.build_backup_zip()
        except Exception as e:
            self.info.set_text('خطا در ساخت پشتیبان:\n' + str(e))
            return
        if not path:
            self.info.set_text('هیچ داده‌ای برای پشتیبان‌گیری وجود ندارد.')
            toast('هنوز کشفی ثبت نشده است.', 'پشتیبان')
            return
        import share_util

        def _cb(ok, msg):
            self.info.set_text(msg)
            toast(msg, 'پشتیبان' if ok else 'خطا')
        share_util.save_file_to_device(path, on_done=_cb, mime='application/zip', then_share=share)

    def open_import(self):
        app = App.get_running_app()
        content = BoxLayout(orientation='vertical', padding=dp(12), spacing=dp(8))
        content.add_widget(RLabel('فایل پشتیبان ویندوز (ZIP) یا فایل JSON (favorites.json / featured.json) را انتخاب کنید، یا متن JSON را بچسبانید:',
                                  font_size='13sp', color=C_GOLD, halign='center', size_hint_y=None, height=dp(64)))
        bpick = PillButton('انتخاب فایل ZIP یا JSON از دستگاه', bg=C_BLUE, size_hint_y=None, height=dp(50), font_size='13sp')
        content.add_widget(bpick)
        picked = {'text': None, 'path': None}
        status = RLabel('', font_size='13sp', halign='center', color=C_GOLD,
                        size_hint_y=None, height=dp(30))
        content.add_widget(status)
        ti = PlainInput(multiline=True, font_size='12sp', size_hint_y=1,
                       hint_text=P('(اختیاری) در صورت تمایل، متن JSON را اینجا بچسبانید'),
                       background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
        content.add_widget(ti)

        def _on_file(path, msg):
            if not path:
                if msg:
                    toast(msg, 'انتخاب فایل')
                return
            picked['path'] = path
            picked['text'] = None
            name = ''
            try:
                import os as _os
                name = _os.path.basename(path)
            except Exception:
                pass
            status.set_text('فایل انتخاب شد%s. روی «بارگذاری» بزنید.' % ((' (%s)' % name) if name else ''))
            toast('فایل انتخاب شد؛ حالا روی «بارگذاری» بزنید.', 'بارگذاری')

        def _pick(*a):
            import share_util
            # هم فایل ZIP خروجی ویندوز و هم فایل JSON پذیرفته می‌شود
            share_util.pick_file(_on_file, mime='*/*')
        bpick.bind(on_release=_pick)

        opts = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(8))
        sp = Spinner(text=P('کشفیات لابراتوار'), values=[P('کشفیات لابراتوار'), P('گلچین')],
                     font_name='ui', size_hint_y=None, height=dp(44))
        state = {'mode': 'merge'}
        btog = PillButton('حالت: ادغام (افزودن جدیدها)', bg=C_GREEN, font_size='12sp')

        def _tog(*a):
            state['mode'] = 'replace' if state['mode'] == 'merge' else 'merge'
            btog.set_text('حالت: جایگزینی کامل' if state['mode'] == 'replace' else 'حالت: ادغام (افزودن جدیدها)')
            btog._bg = list(C_ORANGE if state['mode'] == 'replace' else C_GREEN)
            btog._state()
        btog.bind(on_release=_tog)
        opts.add_widget(sp)
        opts.add_widget(btog)
        content.add_widget(opts)

        pop = Popup(title=P('بارگذاری از فایل JSON'), content=content, size_hint=(0.96, 0.92),
                    title_font='ui', title_align='center', separator_color=C_GOLD)
        row = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
        bload = PillButton('بارگذاری', bg=C_GREEN)

        def _load(*a):
            target = 'lab' if sp.text == P('کشفیات لابراتوار') else 'featured'
            try:
                if picked['path']:
                    added, total, err = app.import_from_path(picked['path'], target, state['mode'])
                else:
                    txt = (ti.text or '')
                    if not txt.strip():
                        toast('ابتدا فایلی انتخاب کنید یا متن JSON را بچسبانید.', 'خطا')
                        return
                    added, total, err = app.import_items_json(txt, target, state['mode'])
            except Exception as e:
                added, total, err = 0, 0, 'خطای غیرمنتظره: %s' % str(e)[:100]
            if err:
                self.info.set_text('خطا: ' + err)
                toast('خطا: ' + err, 'خطا')
                return
            where = 'لابراتوار' if target == 'lab' else 'گلچین'
            msg = '%d مورد جدید به %s افزوده شد (مجموع: %d).' % (added, where, total)
            self.info.set_text(msg)
            toast(msg, 'بارگذاری')
            try:
                sm = self.manager
                sm.get_screen('lab' if target == 'lab' else 'featured').refresh()
            except Exception:
                pass
            pop.dismiss()
        bload.bind(on_release=_load)
        bcancel = PillButton('انصراف', bg=C_RED)
        bcancel.bind(on_release=pop.dismiss)
        row.add_widget(bload)
        row.add_widget(bcancel)
        content.add_widget(row)
        pop.open()


# ==================================================================
# درباره
# ==================================================================
class AboutScreen(BaseScreen):
    WEBSITE = 'https://6a304b9599e34.mywebzi.ir/'
    BALE_URL = 'https://ble.ir/dr_parsa114'

    def __init__(self, **kw):
        super().__init__(title='درباره', **kw)
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', spacing=dp(12), padding=dp(16), size_hint_y=None)
        box.bind(minimum_height=box.setter('height'))
        box.add_widget(self._lbl('قطب‌نمای قرآنی', bold=True, font_size='24sp', color=C_GOLD, halign='center'))
        box.add_widget(self._lbl('پردازش آینه‌ای (هولوگرافیک) — نسخهٔ موبایل', font_size='15sp', color=C_TEXT, halign='center'))
        box.add_widget(self._lbl('کاوش الگوهای عددی و معنایی میان آیات قرآن کریم. تمام ۶۲۳۶ آیه به صورت آفلاین در اپ گنجانده شده است.', font_size='13sp', color=C_MUTED, halign='center'))
        box.add_widget(Widget(size_hint_y=None, height=dp(8)))
        box.add_widget(self._lbl('راه ارتباطی با مؤلف:', bold=True, font_size='17sp', color=C_GOLD, halign='right'))
        b_site = PillButton('سایت مرجع قرآن ابر ماتریس', bg=C_BLUE, size_hint_y=None, height=dp(56), font_size='15sp')
        b_site.bind(on_release=lambda *a: self.open_url(self.WEBSITE))
        box.add_widget(b_site)
        box.add_widget(self._lbl(self.WEBSITE, font_size='12sp', color=C_MUTED, halign='center'))
        b_bale = PillButton('ارتباط در پیام‌رسان بله:  dr_parsa114', bg=C_GREEN, size_hint_y=None, height=dp(56), font_size='15sp')
        b_bale.bind(on_release=lambda *a: self.open_url(self.BALE_URL))
        box.add_widget(b_bale)
        box.add_widget(Widget(size_hint_y=None, height=dp(8)))
        box.add_widget(self._card_text('لطفاً کشفیات ویژهٔ خود را با مؤلف در میان بگذارید و به اشتراک بگذارید تا در نسخه‌های بعدی گنجانده شود.'))
        box.add_widget(self._card_text('این اپلیکیشن و سامانهٔ پردازش آن در حال توسعه و تکامل است؛ ان‌شاءالله به لطف خالق هستی و با کمک یکدیگر، با بزرگ‌تر کردن فهرست آیات آینه‌ای به این هدف مهم دست خواهیم یافت.'))
        box.add_widget(Widget(size_hint_y=None, height=dp(16)))
        scroll.add_widget(box)
        self.body(scroll)

    def _lbl(self, text, **kw):
        kw.setdefault('size_hint_y', None)
        l = RLabel(text, **kw)
        l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(12)))
        return l

    def _card_text(self, text):
        c = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=C_GOLD, orientation='vertical', size_hint_y=None, padding=dp(12))
        l = RLabel(text, font_size='14sp', color=C_TEXT, halign='right', size_hint_y=None)
        l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        c.add_widget(l)
        c.bind(minimum_height=lambda i, v: setattr(c, 'height', v + dp(24)))
        return c

    def open_url(self, url):
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception:
            toast('نشانی: ' + url, 'لینک')

    def refresh(self):
        pass


# ==================================================================
# راهنما
# ==================================================================
class GuideScreen(BaseScreen):
    # هر بخش: (عنوان، رنگ، توضیحِ کامل و خودآموز)
    SECTIONS = [
        ('۱) انتخاب بذر (سوره و آیه)', C_GOLD,
         'نقطهٔ شروعِ همه‌چیز اینجاست. در صفحهٔ اصلی شمارهٔ سوره و شمارهٔ آیه را وارد کنید؛ به این جفت «بذر» می‌گوییم و مبنای تمامِ پردازش‌هاست.\nاگر عددی خارج از محدوده وارد کنید، برنامه خودش آن را به نزدیک‌ترین آیهٔ معتبر اصلاح می‌کند تا خطا نگیرید.\nبرای تایپ کافی است روی کادرِ عدد یک‌بار لمس کنید تا کیبورد بالا بیاید.'),
        ('۲) جستجوی متنِ آیه (در صفحهٔ اصلی)', C_BLUE,
         'اگر شمارهٔ آیه را نمی‌دانید، از کادرِ «جستجوی متن آیه یا ترجمه» در همان صفحهٔ اصلی استفاده کنید.\nبخشی از متنِ عربی یا ترجمهٔ فارسی را بنویسید (نیازی به اعرابِ دقیق نیست) و «انتخاب خودکار» را بزنید تا نزدیک‌ترین آیه پیدا و به‌عنوان بذر انتخاب شود.\nبا دکمهٔ «نمایش لیست جستجو» همهٔ آیه‌های همخوان را در یک فهرست می‌بینید و یکی را برمی‌گزینید؛ اگر عدد بزنید، بر اساسِ شمارهٔ آیه می‌گردد.\nتوجه: این جستجو در متنِ قرآن است، نه در کشفیاتِ شما.'),
        ('۳) پردازش ماتریس', C_PURPLE,
         'قلبِ برنامه. هفت عملگرِ آینه‌ای (ج��بجا��ی و تقارنِ شماره‌های سوره و آیه) را روی بذر اجرا می‌کند و هفت آیهٔ «مقصد» به‌دست می‌آید.\nمتنِ کاملِ عربی و ترجمهٔ هر مقصد نمایش داده می‌شود.\nهر مقصدی که برایتان معنادار بود، با دکمهٔ «ثبت این کشف» در لابراتوار ذخیره کنید تا بعداً بررسی‌اش کنید.'),
        ('۴) پیش‌بینی (معنا)', C_GREEN,
         'کمک می‌کند از میانِ مقصدها، محتمل‌ترین‌ها را زودتر ببینید.\nمقصدهای آینه‌ای را بر اساسِ اشتراکِ واژه‌ها و نزدیکیِ معناییِ آن‌ها با بذر رتبه‌بندی می‌کند؛ هر چه بالاتر، ارتباطِ معناییِ قوی‌تر.'),
        ('۵) پیش‌بینی (اعداد)', C_ORANGE,
         'همان رتبه‌بندی، اما با معیارهای عددی.\nبا فیلترهایی مثلِ نیم‌کرهٔ سوره، اثرِانگشتِ رقمی و میزانِ تلورانس، نامزدهای عددی را غربال و اولویت‌بندی می‌کند تا الگوهای عددی راحت‌تر دیده شوند.'),
        ('۶) لابراتوار کشفیات', C_BLUE,
         'انبارِ همهٔ کشف‌های ثبت‌شدهٔ شما.\nکشف‌ها خودکار در این دسته‌ها مرتب می‌شوند: هفت عملگرِ آینه‌ای، به‌همراهِ «کشفیات گروهی» (ثبتِ گروهی) و «کشفیات تردیدی».\nروی هر دسته بزنید تا کشف‌های آن باز شود؛ فهرست به‌صورتِ تدریجی و روان بارگذاری می‌شود، پس حتی با هزاران کشف هم صفحه فوری باز می‌شود و معطل نمی‌مانید.\nروی هر کشف بزنید تا پنجرهٔ کامل باز شود: متنِ عربی و ترجمهٔ مبدأ و مقصد، همراه با گلچین‌کردن، ویرایشِ تحلیل و برچسب، حذف و کپی.'),
        ('۷) گلچین برگزیده', C_GOLD,
         'ویترینِ بهترین کشف‌های شما.\nکشف‌های مهم را از لابراتوار به گلچین می‌آورید. اینجا هم مثلِ لابراتوار بر اساسِ همان عملگرها و دسته‌ها (گروهی و تردیدی) مرتب شده و به‌صورتِ روان بارگذاری می‌شود.\nمی‌توانید از کلِ گلچین یک خروجیِ تمیزِ JSON بگیرید.'),
        ('۸) جستجوی کشفیات', C_PURPLE,
         'این جستجو با «جستجوی متن آیه» فرق دارد: اینجا فقط داخلِ کشف‌های خودتان (لابراتوار و گلچین) می‌گردد، نه کلِ قرآن.\nهر چیزی را می‌توانید بگردید: متنِ عربی، ترجمه، برچسب، تحلیلِ خودتان و شماره‌ها.\nنتیجه‌ها زنده و همزمان با تایپ نشان داده می‌شوند و حتی با هزاران کشف هم سریع می‌مانند.\nروی هر نتیجه بزنید تا همان پنجرهٔ کاملِ ویرایش باز شود.'),
        ('۹) مدیریت برچسب‌ها', C_GREEN,
         'برچسب‌های «رفتارِ آیه» را اینجا می‌سازید یا حذف می‌کنید؛ مثلاً «تقابلِ کامل»، «گفت‌وگو»، «علت و معلول».\nبعداً هنگامِ ثبتِ تحلیلِ یک کشف، این برچسب‌ها را به آن نسبت می‌دهید تا کشف‌هایتان منظم و قابلِ جستجو شوند.'),
        ('۱۰) رسانه و معرفی', C_ORANGE,
         'در این بخش «چند کلام از طراح» را می‌شنوید.\nیک معرفیِ صوتیِ کوتاه هنگامِ باز شدنِ برنامه یک‌بار پخش می‌شود و از اینجا هم می‌توانید دوباره آن را بشنوید.'),
        ('۱۱) پشتیبان و بازیابی', C_BLUE,
         'برای اینکه دادهٔ شما هیچ‌وقت گم نشود.\nاز کشفیات، گلچین و برچسب‌ها یک فایلِ پشتیبان بگیرید و هر وقت خواستید بازیابی کنید.\nهنگامِ بازیابی، مقصد را انتخاب می‌کنید (لابراتوار یا گلچین) و حالتِ «جایگزینی» یا «ادغام» را برمی‌گزینید؛ کشف‌های تکراری خودکار حذف می‌شوند.'),
        ('۱۲) درباره', C_PURPLE,
         'معرفیِ برنامه و راه‌های ارتباط با مؤلف (سایتِ مرجع و شناسهٔ پیام‌رسانِ بله).'),
    ]

    def __init__(self, **kw):
        super().__init__(title='راهنما', **kw)
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(8))
        box.bind(minimum_height=box.setter('height'))
        box.add_widget(RLabel('روی هر بخش بزنید تا توضیحِ کاملِ آن باز شود.', font_size='14sp',
                              color=C_MUTED, halign='center', size_hint_y=None, height=dp(34)))

        def _tint(col, f=0.32):
            return [c * f for c in col[:3]] + [1]

        for title, color, body in self.SECTIONS:
            b = PillButton(title, bg=_tint(color), size_hint_y=None, height=dp(58), font_size='15sp')
            _neon_border(b, color, width=1.4, alpha=0.85)
            b.bind(on_release=lambda inst, t=title, d=body, c=color: self.show_help(t, d, c))
            box.add_widget(b)
        gt = asset('guide_table.png')
        if os.path.exists(gt):
            img = Image(source=gt, size_hint_y=None, height=dp(200), allow_stretch=True, keep_ratio=True)
            box.add_widget(img)
        scroll.add_widget(box)
        self.body(scroll)

    def show_help(self, title, body, color=None):
        if color is None:
            color = C_GOLD
        content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
        sc = ScrollView(do_scroll_x=False, bar_width=dp(4))
        inner = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
        inner.bind(minimum_height=inner.setter('height'))
        head = RLabel(title, bold=True, font_size='18sp', color=color,
                      halign='center', size_hint_y=None)
        head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(10)))
        inner.add_widget(head)
        lbl = RLabel(body, font_size='15sp', color=C_TEXT,
                     halign='center', size_hint_y=None)
        lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(12)))
        inner.add_widget(lbl)
        sc.add_widget(inner)
        content.add_widget(sc)
        p = Popup(title=P(title), content=content, size_hint=(0.92, 0.72),
                  title_font='ui', title_align='center', separator_color=color)
        btn = PillButton('بستن', bg=color, size_hint_y=None, height=dp(46))
        btn.bind(on_release=p.dismiss)
        content.add_widget(btn)
        p.open()

    def refresh(self):
        pass


def _open_url(url):
    """باز کردنِ یک لینک در مرورگرِ سیستم (اندروید و دسکتاپ)."""
    try:
        if _kivy_platform == 'android':
            from jnius import autoclass
            Intent = autoclass('android.content.Intent')
            Uri = autoclass('android.net.Uri')
            PythonActivity = autoclass('org.kivy.android.PythonActivity')
            intent = Intent(Intent.ACTION_VIEW)
            intent.setData(Uri.parse(url))
            PythonActivity.mActivity.startActivity(intent)
        else:
            import webbrowser
            webbrowser.open(url)
    except Exception:
        toast('نشد سایت باز شود؛ این آدرس را دستی باز کن:\n' + url, 'راهنما')


def open_ai_settings():
    """پنجرهٔ تنظیماتِ هوش مصنوعی: کلید API و آدرس سرور + تست اتصال."""
    app = App.get_running_app()
    cfg = getattr(app, 'ai_settings', None) or ai_manager.default_settings()
    content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(8))

    sv = ScrollView(size_hint=(1, 1), do_scroll_x=False, bar_width=dp(4))
    box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(6), padding=dp(2))
    box.bind(minimum_height=box.setter('height'))

    box.add_widget(RLabel('کلید API', bold=True, font_size='15sp', color=C_GOLD,
                          halign='right', size_hint_y=None, height=dp(26)))
    key_in = PlainInput(text=cfg.get('api_key', ''), multiline=False, font_name='ui',
                        font_size='14sp', size_hint_y=None, height=dp(46), hint_text='aa-...',
                        background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
    box.add_widget(key_in)

    box.add_widget(RLabel('آدرس سرور (Base URL)', bold=True, font_size='15sp', color=C_GOLD,
                          halign='right', size_hint_y=None, height=dp(26)))
    url_in = PlainInput(text=cfg.get('base_url', ai_manager.DEFAULT_BASE_URL), multiline=False,
                        font_name='ui', font_size='13sp', size_hint_y=None, height=dp(46),
                        background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
    box.add_widget(url_in)

    box.add_widget(RLabel('نام مدل (پیشرفته)', bold=True, font_size='15sp', color=C_GOLD,
                          halign='right', size_hint_y=None, height=dp(26)))
    model_in = PlainInput(text=cfg.get('model', ai_manager.DEFAULT_MODEL), multiline=False,
                          font_name='ui', font_size='13sp', size_hint_y=None, height=dp(46),
                          background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
    box.add_widget(model_in)
    guide = RLabel('''«دستیارِ هوشمند» چگونه فعال می‌شود؟

۱) این بخش برای تحلیلِ هوشمند به یک «کلیدِ API» نیاز دارد؛ کلید را در کادرِ بالا وارد و ذخیره کن.

۲) برای «شروع»، مؤلف یک «کدِ هدیه» به تو می‌دهد تا رایگان امتحان کنی.

۳) برای «ادامهٔ» استفاده، خودت از سایتِ اَوال اعتبار تهیه کن. آدرسِ سایت:
https://avalai.ir

۴) پیش‌فرض روی سرویسِ «اَوال» تنظیم است. برای اتصال به سرویسِ دیگر (مثلِ NVIDIA) فقط «آدرس سرور» و «نام مدل» را عوض کن. نمونهٔ نامِ مدل:
nvidia/llama-3.1-nemotron-70b-instruct

۵) برای دریافتِ کدِ هدیه یا هر پرسشی، از پیام‌رسانِ «بله» با مؤلف در تماس باش.''',
                   font_size='13sp', color=C_MUTED, halign='center', size_hint_y=None)
    guide.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(12)))
    box.add_widget(guide)
    b_link = PillButton('باز کردن سایتِ اَوال', bg=(1, 1, 1, 0.12), fg=C_TEXT,
                        size_hint_y=None, height=dp(44), font_size='13sp')
    b_link.bind(on_release=lambda *a: _open_url(ai_manager.MODELS_URL))
    box.add_widget(b_link)
    _note = RLabel('''توجه: برخی سرورها (مثلِ NVIDIA) ممکن است از داخلِ ایران بدونِ فیلترشکن پاسخ ندهند؛ در این صورت از همان سرویسِ پیش‌فرضِ اَوال استفاده کن.''', font_size='12sp', color=C_MUTED,
                   halign='center', size_hint_y=None)
    _note.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(10)))
    box.add_widget(_note)

    sv.add_widget(box)
    content.add_widget(sv)
    status = RLabel('', font_size='13sp', halign='center', color=C_MUTED,
                    size_hint_y=None, height=dp(26))
    content.add_widget(status)

    p = Popup(title=P('تنظیمات هوش مصنوعی'), content=content, size_hint=(0.94, 0.9),
              title_font='ui', title_align='center', separator_color=C_GOLD)

    def _collect():
        return (key_in.text.strip(), url_in.text.strip(),
                model_in.text.strip() or ai_manager.DEFAULT_MODEL)

    def _save():
        k, u, m = _collect()
        app.save_ai_settings(api_key=k, base_url=u, model=m)

    def _test(*a):
        k, u, m = _collect()
        if not k:
            status.set_text('× ابتدا کلید API را وارد کن')
            toast('کلید API را وارد کن.', 'هوش مصنوعی', kind='warn')
            return
        app.save_ai_settings(api_key=k, base_url=u, model=m)
        status.set_text('در حال آزمایش اتصال...')
        b_test.disabled = True

        def _ok(msg):
            b_test.disabled = False
            status.set_text('✓ ' + msg)
            toast(msg, 'هوش مصنوعی', kind='success')

        def _fail(msg):
            b_test.disabled = False
            status.set_text('× ' + msg)
            toast(msg, 'هوش مصنوعی', kind='error')

        ai_manager.test_connection(on_ok=_ok, on_fail=_fail)

    btnrow = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
    b_test = PillButton('تست اتصال', bg=C_BLUE, font_size='14sp')
    b_test.bind(on_release=_test)
    b_save = PillButton('ذخیره و اتصال', bg=C_GREEN, font_size='14sp')
    b_save.bind(on_release=lambda *a: (_save(), toast('ت��ظیمات ذخیره ش�� ✓', 'هوش مصنوعی'), p.dismiss()))
    btnrow.add_widget(b_test)
    btnrow.add_widget(b_save)
    content.add_widget(btnrow)
    close = PillButton('بستن', bg=C_RED, size_hint_y=None, height=dp(44))
    close.bind(on_release=lambda *a: p.dismiss())
    content.add_widget(close)
    p.open()


def show_ai_result_popup(title, messages, subtitle=None, temperature=0.5, max_tokens=None):
    """پاپ‌آپِ گفتگویی: نخست تحلیل را زنده نشان می‌دهد و سپس می‌توان دربارهٔ همان تحلیل گفتگو کرد."""
    convo = list(messages)
    state = {'busy': True}

    root = BoxLayout(orientation='vertical', spacing=dp(8), padding=dp(8))
    if subtitle:
        root.add_widget(RLabel(subtitle, font_size='13sp', color=C_GOLD, halign='center',
                               size_hint_y=None, height=dp(28)))

    sv = ScrollView(size_hint=(1, 1), do_scroll_x=False, bar_width=dp(8))
    log = BoxLayout(orientation='vertical', size_hint_y=None, padding=dp(4), spacing=dp(10))
    log.bind(minimum_height=log.setter('height'))
    sv.add_widget(log)
    root.add_widget(sv)

    p = Popup(title=P(title), content=root, size_hint=(0.96, 0.94),
              title_font='ui', title_align='center', separator_color=C_GOLD)

    def _scroll_bottom(*a):
        sv.scroll_y = 0

    def _bubble(text, role='ai'):
        is_user = (role == 'user')
        bg = (0.10, 0.30, 0.18, 1) if is_user else (0.10, 0.14, 0.22, 1)
        card = RoundBox(bg=bg, orientation='vertical', size_hint_y=None, padding=dp(10), spacing=dp(4))
        card.bind(minimum_height=card.setter('height'))
        who = RLabel('تو' if is_user else 'هوش مصنوعی', font_size='12sp', bold=True,
                     color=(C_GREEN if is_user else C_GOLD), halign='right',
                     size_hint_y=None, height=dp(20))
        card.add_widget(who)
        lbl = RLabel(text, font_size='15sp', color=C_TEXT, halign='right', size_hint_y=None)
        lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
        card.add_widget(lbl)
        log.add_widget(card)
        _fade_in(card)
        Clock.schedule_once(_scroll_bottom, 0.05)
        return lbl

    def _stream_into(target_lbl):
        acc = {'t': ''}

        def _delta(piece):
            acc['t'] += piece
            target_lbl.set_text(acc['t'])
            Clock.schedule_once(_scroll_bottom, 0)

        def _done(full):
            text = (full or acc['t'] or '').strip() or '(پاسخی دریافت نشد)'
            target_lbl.set_text(text)
            convo.append({'role': 'assistant', 'content': text})
            state['busy'] = False

        def _err(msg):
            target_lbl.set_text('⚠ ' + msg)
            toast(msg, 'هوش مصنوعی', kind='error')
            state['busy'] = False

        ai_manager.chat(convo, on_delta=_delta, on_done=_done, on_error=_err,
                        stream=True, temperature=temperature, max_tokens=max_tokens)

    def _send(*a):
        if state['busy']:
            toast('صبر کن تا پاسخِ قبلی کامل شود.', 'هوش مصنوعی')
            return
        q = (inp.query or '').strip()
        if not q:
            return
        inp.clear_logical()
        convo.append({'role': 'user', 'content': q})
        _bubble(q, role='user')
        state['busy'] = True
        _stream_into(_bubble('...', role='ai'))

    inrow = RoundBox(bg=(0.05, 0.08, 0.14, 0.55), orientation='horizontal', size_hint_y=None,
                     height=dp(60), padding=dp(6), spacing=dp(6))
    inp = PersianTextInput(hint_text=P('دربارهٔ این تحلیل بپرس...'), multiline=False,
                           font_size='15sp', size_hint_y=None, height=dp(48),
                           background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
    inp.bind(on_text_validate=_send)
    b_send = PillButton('ارسال', bg=C_GREEN, size_hint_x=None, width=dp(88), font_size='14sp')
    b_send.bind(on_release=_send)
    inrow.add_widget(inp)
    inrow.add_widget(b_send)
    root.add_widget(inrow)

    close = PillButton('بستن', bg=C_RED, size_hint_y=None, height=dp(44))
    close.bind(on_release=lambda *a: p.dismiss())
    root.add_widget(close)

    p.open()
    _stream_into(_bubble('در حال دریافت تحلیل...', role='ai'))
    return p


# ==================================================================
# صفحهٔ گفتگو با هوش مصنوعی
# ==================================================================
class ChatScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='گفت��و با هوش ��صنوعی', **kw)
        self._messages = []      # تاریخچهٔ گفتگو (بدون system)
        self._busy = False
        self._cur_label = None
        self._acc = ''
        self._attach = None      # پیوستِ در انتظارِ ارسال
        try:
            self.title_label._fit_single = True
        except Exception:
            pass

        try:
            gear = PillButton('تنظیمات', bg=(1, 1, 1, 0.14), size_hint_x=None,
                              width=dp(96), font_size='13sp')
            gear.bind(on_release=lambda *a: open_ai_settings())
            self.header.add_widget(gear)
        except Exception:
            pass

        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.log = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(6))
        self.log.bind(minimum_height=self.log.setter('height'))
        self.scroll.add_widget(self.log)
        self.body(self.scroll)

        # نوارِ نمایشِ پیوستِ انتخاب‌شده (پیش‌فرض پنهان)
        self.attach_bar = BoxLayout(size_hint_y=None, height=0, spacing=dp(6), opacity=0)
        self.attach_lbl = RLabel('', font_size='12sp', color=C_GOLD, halign='right')
        b_clr = PillButton('حذف پیوست', bg=C_RED, size_hint_x=None, width=dp(104), font_size='12sp')
        b_clr.bind(on_release=lambda *a: self._clear_attach())
        self.attach_bar.add_widget(self.attach_lbl)
        self.attach_bar.add_widget(b_clr)
        self.body(self.attach_bar)

        inrow = RoundBox(bg=(0.05, 0.08, 0.14, 0.62), orientation='horizontal', size_hint_y=None,
                         height=dp(60), padding=dp(6), spacing=dp(6))
        b_attach = PillButton('پیوست', bg=(1, 1, 1, 0.12), fg=C_TEXT, size_hint_x=None,
                              width=dp(76), font_size='13sp')
        b_attach.bind(on_release=lambda *a: self._pick_file())
        self.inp = PersianTextInput(hint_text=P('پیامت را بنویس...'), multiline=False,
                                    font_size='15sp', size_hint_y=None, height=dp(48),
                                    background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
        self.inp.bind(on_text_validate=lambda *a: self._send())
        send = PillButton('ارسال', bg=C_GREEN, size_hint_x=None, width=dp(84), font_size='14sp')
        send.bind(on_release=lambda *a: self._send())
        inrow.add_widget(b_attach)
        inrow.add_widget(self.inp)
        inrow.add_widget(send)
        self.body(inrow)
        _neon_border(inrow, C_BLUE, width=1.4, alpha=0.9)

    def refresh(self):
        if not self.log.children:
            self._add_bubble('سلام! من به همهٔ کشفیاتِ لابراتوار و گلچینِ تو دسترسی دارم. '
                             'دربارهٔ الگوها، شباهت‌ها یا هر تحلیلی که بخواهی بپرس. '
                             'با دکمهٔ «پیوست» هم می‌توانی عکس، فایلِ متنی یا PDF بفرستی.', role='ai')

    # ---------- پیوستِ فایل / تصویر ----------
    def _pick_file(self):
        try:
            from plyer import filechooser
            filechooser.open_file(on_selection=self._on_file, multiple=False)
        except Exception:
            toast('انتخاب فایل روی این دستگاه ممکن نشد.', 'پیوست', kind='error')

    def _on_file(self, selection):
        if not selection:
            return
        path = selection[0]
        Clock.schedule_once(lambda *a: self._process_file(path), 0)

    def _process_file(self, path):
        try:
            ext = os.path.splitext(path)[1].lower()
            name = os.path.basename(path)
            if ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp'):
                with open(path, 'rb') as f:
                    raw = f.read()
                if len(raw) > 6 * 1024 * 1024:
                    toast('حجمِ تصویر خیلی زیاد است (بیش از ۶ مگابایت).', 'پیوست', kind='warn')
                    return
                mime = 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/' + ext.lstrip('.')
                b64 = base64.b64encode(raw).decode('ascii')
                self._attach = {'kind': 'image', 'name': name,
                                'data_url': 'data:%s;base64,%s' % (mime, b64)}
                self._show_attach('��صویر: ' + name)
            elif ext == '.pdf':
                txt = self._extract_pdf(path)
                if txt is None:
                    return
                self._attach = {'kind': 'text', 'name': name, 'text': txt}
                self._show_attach('PDF: ' + name)
            elif ext in ('.txt', '.md', '.csv', '.json', '.log'):
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    txt = f.read()
                self._attach = {'kind': 'text', 'name': name, 'text': txt[:8000]}
                self._show_attach('فایل: ' + name)
            else:
                toast('این نوع فایل پشتیبانی نمی‌شود (عکس، PDF یا فایلِ متنی بفرست).', 'پیوست', kind='warn')
        except Exception:
            toast('خواندنِ فایل ممکن نشد.', 'پیوست', kind='error')

    def _extract_pdf(self, path):
        try:
            from pypdf import PdfReader
        except Exception:
            try:
                from PyPDF2 import PdfReader
            except Exception:
                toast('کتابخانهٔ خواندنِ PDF در دسترس نیست.', 'پیوست', kind='error')
                return None
        try:
            reader = PdfReader(path)
            parts = []
            for pg in reader.pages[:30]:
                parts.append(pg.extract_text() or '')
            txt = (chr(10).join(parts)).strip()
            if not txt:
                toast('متنی از این PDF استخراج نشد (شاید اسکن‌شده باشد).', 'پیوست', kind='warn')
                return None
            return txt[:8000]
        except Exception:
            toast('خواندنِ PDF ممکن نشد.', 'پیوست', kind='error')
            return None

    def _show_attach(self, label):
        self.attach_lbl.set_text(label)
        self.attach_bar.height = dp(34)
        self.attach_bar.opacity = 1

    def _clear_attach(self):
        self._attach = None
        self.attach_lbl.set_text('')
        self.attach_bar.height = 0
        self.attach_bar.opacity = 0

    def _add_bubble(self, text, role='ai'):
        is_user = (role == 'user')
        bg = (0.10, 0.28, 0.16, 1) if is_user else (0.10, 0.14, 0.22, 1)
        card = RoundBox(bg=bg, orientation='vertical', size_hint_y=None, padding=dp(10), spacing=dp(4))
        who = RLabel('تو' if is_user else 'هوش مصنوعی', font_size='12sp', bold=True,
                     color=(C_GREEN if is_user else C_GOLD), halign='right',
                     size_hint_y=None, height=dp(20))
        card.add_widget(who)
        lbl = RLabel(text, font_size='15sp', color=C_TEXT, halign='right', size_hint_y=None)
        lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
        card.add_widget(lbl)
        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v))
        self.log.add_widget(card)
        _fade_in(card)
        Clock.schedule_once(lambda *a: setattr(self.scroll, 'scroll_y', 0), 0.05)
        return lbl

    def _send(self):
        if self._busy:
            toast('صبر کن تا پاسخِ قبلی کامل شود.', 'هوش مصنوعی')
            return
        q = self.inp.query.strip()
        att = self._attach
        if not q and not att:
            return
        self.inp.clear_logical()
        nl = chr(10)

        if att and att['kind'] == 'image':
            shown = (q + '   ' if q else '') + '[تصویر: ' + att['name'] + ']'
            hist = q or 'این تصویر را بررسی کن.'
            send_content = [{'type': 'text', 'text': q or 'این تصویر را بررسی کن و توضیح بده.'},
                            {'type': 'image_url', 'image_url': {'url': att['data_url']}}]
        elif att and att['kind'] == 'text':
            shown = (q + '   ' if q else '') + '[فایل: ' + att['name'] + ']'
            hist = q or ('بررسیِ فایل: ' + att['name'])
            send_content = ((q or 'این فایل را بررسی و خلاصه کن.') + nl + nl +
                            '[محتوای فایلِ «' + att['name'] + '»]:' + nl + att['text'])
        else:
            shown = q
            hist = q
            send_content = q

        self._add_bubble(shown, role='user')
        self._messages.append({'role': 'user', 'content': hist})
        self._clear_attach()

        app = App.get_running_app()
        system = ai_manager.build_chat_system(app.favs, app.featured)
        msgs = [{'role': 'system', 'content': system}] + self._messages[-12:]
        msgs[-1] = {'role': 'user', 'content': send_content}

        self._busy = True
        self._acc = ''
        self._cur_label = self._add_bubble('...', role='ai')

        def _delta(piece):
            self._acc += piece
            if self._cur_label:
                self._cur_label.set_text(self._acc)
                Clock.schedule_once(lambda *a: setattr(self.scroll, 'scroll_y', 0), 0)

        def _done(full):
            self._busy = False
            text = (full or self._acc or '').strip() or '(پاسخی دریافت نشد)'
            if self._cur_label:
                self._cur_label.set_text(text)
            self._messages.append({'role': 'assistant', 'content': text})
            self._cur_label = None

        def _err(msg):
            self._busy = False
            if self._cur_label:
                self._cur_label.set_text('⚠ ' + msg)
            self._cur_label = None
            toast(msg, 'هوش مصنوعی', kind='error')

        ai_manager.chat(msgs, on_delta=_delta, on_done=_done, on_error=_err,
                        stream=True, temperature=0.6, max_tokens=1024)


# ==================================================================
# اپلیکیشن
# ==================================================================
class QuranMirrorApp(App):
    def build(self):
        self.title = 'قطب‌نمای قرآنی'
        Window.clearcolor = C_BG
        # رفع مشکل پوشاندنِ باکسِ ورودی توسط کیبورد:
        # حالت 'pan' کلِ صفحه را به‌اندازهٔ لازم بالا می‌برد تا باکسِ در حالِ تایپ
        # همیشه بالای کیبورد و در دیدِ کاربر بماند (در همهٔ پنجره‌ها و پاپ‌آپ‌ها).
        try:
            # فقط روی موبایل لازم است (تا کیبوردِ صفحه‌ای باکس را نپوشاند)؛
            # روی ویندوز/دسکتاپ این حالت می‌تواند هنگامِ فوکوسِ باکس باعثِ هنگ/فریز شود.
            if _kivy_platform in ('android', 'ios'):
                Window.softinput_mode = ''  # جابه‌جایی به عهدهٔ خودِ اندروید (adjustPan)؛ پنل‌کردنِ کیوی باعثِ هنگ/سیاه‌شدن می‌شد
        except Exception:
            pass
        # تور ایمنی: خطاهای پیش‌بینی‌نشده به‌جای بستنِ کامل برنامه نادیده گرفته ��وند
        try:
            from kivy.base import ExceptionManager, ExceptionHandler

            _guard_state = {'last': 0.0}

            class _AppGuard(ExceptionHandler):
                def handle_exception(self, exc):
                    try:
                        import traceback
                        traceback.print_exc()
                    except Exception:
                        pass
                    # نمایشِ خطا روی صفحه (به‌جای نادیده‌گرفتن کامل) تا علت «کارنکردن دکمه‌ها» دیده شود
                    try:
                        now = Clock.get_boottime()
                        if now - _guard_state['last'] > 2.0:   # جلوگیری از اسپم پاپ‌آپ
                            _guard_state['last'] = now
                            msg = '%s: %s' % (type(exc).__name__, exc)
                            Clock.schedule_once(lambda *a: toast(msg[:400], 'خطای داخلی'), 0)
                    except Exception:
                        pass
                    return ExceptionManager.PASS

            ExceptionManager.add_handler(_AppGuard())
        except Exception:
            pass
        # داده
        self.data = core.QuranData(asset('datakavosh.csv'))
        self._init_storage()
        self.load_favs()
        self.load_featured()
        self.load_user_tags()
        self.load_ai_settings()
        self.last_discovery_key = None
        self.last_discovery_section = None
        # صفحات
        sm = ScreenManager(transition=FadeTransition(duration=0.25))
        sm.add_widget(HomeScreen(name='home'))
        sm.add_widget(MatrixScreen(name='matrix'))
        sm.add_widget(PredictScreen(name='predict'))
        sm.add_widget(LabScreen(name='lab'))
        sm.add_widget(OperatorScreen(name='operator'))
        sm.add_widget(FeaturedScreen(name='featured'))
        sm.add_widget(SearchScreen(name='search'))
        sm.add_widget(TagsScreen(name='tags'))
        sm.add_widget(MediaScreen(name='media'))
        sm.add_widget(BackupScreen(name='backup'))
        sm.add_widget(AboutScreen(name='about'))
        sm.add_widget(GuideScreen(name='guide'))
        sm.add_widget(ChatScreen(name='chat'))
        self.sm = sm
        return sm

    def on_start(self):
        try:
            from kivy.core.audio import SoundLoader
            path = asset('voice.mp3')
            if os.path.exists(path):
                snd = SoundLoader.load(path)
                if snd:
                    self._intro_sound = snd
                    Clock.schedule_once(lambda *a: snd.play(), 0.6)
        except Exception:
            pass

    # ---------- ذخیره‌سازی ----------
    def _init_storage(self):
        self.store_dir = self.user_data_dir
        for name in ('favorites.json', 'featured.json', 'user_tags.json'):
            dst = os.path.join(self.store_dir, name)
            if not os.path.exists(dst):
                src = asset(name)
                try:
                    if os.path.exists(src):
                        shutil.copy(src, dst)
                    else:
                        with open(dst, 'w', encoding='utf-8') as f:
                            json.dump([], f, ensure_ascii=False)
                except Exception:
                    pass

    def _p(self, name):
        return os.path.join(self.store_dir, name)

    def load_favs(self):
        try:
            with open(self._p('favorites.json'), encoding='utf-8') as f:
                self.favs = json.load(f)
        except Exception:
            self.favs = []
        if not isinstance(self.favs, list):
            self.favs = []
        for _it in self.favs:
            _normalize_item_modes(_it)

    def save_favs(self):
        _atomic_write_json(self._p('favorites.json'), self.favs, indent=2)

    def load_featured(self):
        try:
            with open(self._p('featured.json'), encoding='utf-8') as f:
                self.featured = json.load(f)
        except Exception:
            self.featured = []
        if not isinstance(self.featured, list):
            self.featured = []
        for _it in self.featured:
            _normalize_item_modes(_it)

    def save_featured(self):
        _atomic_write_json(self._p('featured.json'), self.featured, indent=2)

    def load_user_tags(self):
        try:
            with open(self._p('user_tags.json'), encoding='utf-8') as f:
                self.user_tags = json.load(f)
        except Exception:
            self.user_tags = []

    def save_user_tags(self):
        _atomic_write_json(self._p('user_tags.json'), self.user_tags, indent=2)

    # ---------- تنظیمات هوش مصنوعی ----------
    def load_ai_settings(self):
        self.ai_settings = ai_manager.default_settings()
        try:
            with open(self._p('settings.json'), encoding='utf-8') as f:
                data = json.load(f)
            if isinstance(data, dict):
                for k in ('api_key', 'base_url', 'model'):
                    if data.get(k):
                        self.ai_settings[k] = data[k]
        except Exception:
            pass
        # ماژ��ل هوش مصنوعی تنظیمات را «زنده» از همین‌جا می‌خواند
        ai_manager.set_config_provider(lambda: getattr(self, 'ai_settings', {}))

    def save_ai_settings(self, api_key=None, base_url=None, model=None):
        if not isinstance(getattr(self, 'ai_settings', None), dict):
            self.ai_settings = ai_manager.default_settings()
        if api_key is not None:
            self.ai_settings['api_key'] = api_key.strip()
        if base_url is not None:
            self.ai_settings['base_url'] = base_url.strip() or ai_manager.DEFAULT_BASE_URL
        if model is not None:
            self.ai_settings['model'] = model.strip() or ai_manager.DEFAULT_MODEL
        _atomic_write_json(self._p('settings.json'), self.ai_settings, indent=2)

    def get_all_tags(self):
        # برچسب‌ها فقط از منابعِ صریح ساخته می‌شوند: پیش‌فرض‌ها + برچسب‌هایی که کاربر دستی ساخته.
        # دیگر هیچ برچسبِ جدیدی از روی انتخاب‌های گذشته (relation_type ترکیبی/چنداتیکتی) ساخته نمی‌شود.
        tags = set(TagsScreen.DEFAULT) | {'نامشخص'}
        tags.update(self.user_tags)
        return sorted(tags)

    # ---------- عملیات کشف ----------
    def add_discovery(self, seed, target):
        entry = {
            'mode': normalize_mode(target.get('mode', '')),
            'seed_s': seed['s'], 'seed_a': seed['a'],
            'seed_arb': seed['arb'], 'seed_pers': seed['pers'],
            'target_s': target['s'], 'target_a': target['a'],
            'target_arb': target.get('arb', ''), 'target_pers': target.get('pers', ''),
            'note': '', 'relation_type': 'نامشخص',
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        }
        # جلوگیری از تکرار
        for it in self.favs:
            if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'),
                    it.get('mode')) == (entry['seed_s'], entry['seed_a'], entry['target_s'],
                                        entry['target_a'], entry['mode']):
                toast('این کشف قبلاً ثبت شده است.', 'تکرار')
                return
        if not entry.get('note'):
            entry['note'] = generate_default_analysis(entry)
        self.favs.append(entry)
        self.save_favs()
        self.last_discovery_key = discovery_key(entry)
        self.last_discovery_section = lab_section_of(entry)
        open_note_editor(entry, 'lab', title='ثبت تحلیل کشف',
                         intro='کشف در لابراتوار ثبت شد. تحلیل خود را ثبت کنید:',
                         on_saved=lambda: setattr(self, 'last_discovery_section', lab_section_of(entry)),
                         saved_msg='کشف با موفقیت در لابراتوار ثبت شد ✓')

    def add_featured(self, item, screen=None):
        for it in self.featured:
            if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a')) == \
               (item.get('seed_s'), item.get('seed_a'), item.get('target_s'), item.get('target_a')):
                toast('این مورد در گلچین هست.', 'تکرار')
                return
        self.featured.append(dict(item))
        self.save_featured()
        toast('به گلچین اضافه شد. ', 'گلچین')

    def add_all_featured(self):
        existing = {(it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'))
                    for it in self.featured}
        n = 0
        for it in self.favs:
            key = (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'))
            if key not in existing:
                self.featured.append(dict(it))
                existing.add(key)
                n += 1
        self.save_featured()
        return n

    # ---------- خروجی JSON تمیز ----------
    def export_featured_word(self):
        if not self.featured:
            return None
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            h = doc.add_heading('', level=0)
            run = h.add_run('گلچین آیات آینه‌ای')
            run.font.size = Pt(20)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, it in enumerate(self.featured, 1):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(f"{i}. [{it.get('mode', '')}] سوره {it.get('seed_s')}:{it.get('seed_a')} سوره {it.get('target_s')}:{it.get('target_a')}")
                r.bold = True
                for key in ('seed_arb', 'seed_pers', 'target_arb', 'target_pers'):
                    par = doc.add_paragraph(it.get(key, ''))
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if it.get('note'):
                    par = doc.add_paragraph('یادداشت: ' + it['note'])
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph('—' * 20)
            out = self._p('golchin_%s.docx' % datetime.now().strftime('%Y%m%d_%H%M'))
            doc.save(out)
            return out
        except Exception as e:
            print('word export error:', e)
            return None

    # ---------- خروجی JSON تمیز (برای هر عملگر) ----------
    def _clean_by_operator(self, items):
        op_names = dict(OPERATORS)
        grouped = {}
        for it in items:
            k = op_of(it)
            entry = {
                'mode': it.get('mode', ''),
                'seed': {'sura': it.get('seed_s'), 'ayah': it.get('seed_a'),
                         'arabic': it.get('seed_arb', ''), 'translation': it.get('seed_pers', '')},
                'target': {'sura': it.get('target_s'), 'ayah': it.get('target_a'),
                           'arabic': it.get('target_arb', ''), 'translation': it.get('target_pers', '')},
                'relation_type': it.get('relation_type', 'نامشخص'),
                'is_doubtful': bool(it.get('is_doubtful', False)),
                'note': it.get('note', ''),
                'date': it.get('date', ''),
            }
            grouped.setdefault(k, {'operator': k, 'operator_name': op_names.get(k, k), 'items': []})
            grouped[k]['items'].append(entry)
        return [grouped[k] for k, _ in OPERATORS if k in grouped]

    def export_clean_json(self, filename, payload):
        out = self._p(filename)
        _atomic_write_json(out, payload, indent=2)
        return out

    def build_backup_json(self):
        if not self.favs and not self.featured:
            return None
        payload = {
            'app': 'قطب‌نمای قرآنی',
            'version': '3.0',
            'exported_at': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'lab': self._clean_by_operator(self.favs),
            'featured': self._clean_by_operator(self.featured),
            'tags': self.get_all_tags(),
        }
        return self.export_clean_json('backup_%s.json' % datetime.now().strftime('%Y%m%d_%H%M'), payload)

    def build_backup_zip(self):
        """پشتیبانِ کاملِ سازگار با نسخهٔ ویندوز.
        یک فایل ZIP می‌سازد که شاملِ favorites.json / featured.json / user_tags.json
        (همان قالبِ خامِ داخلی) است؛ نسخهٔ ویندوز این ZIP را مستقیماً «بازیابی» می‌کند
        و خودِ اپ هم آن را کامل می‌خواند. یک نسخهٔ خوانا (backup_readable.json) هم صرفاً
        برای مطالعهٔ انسانی داخلِ ZIP گذاشته می‌شود و برای بازیابی لازم نیست."""
        if not self.favs and not self.featured:
            return None
        # ابتدا آخرین وضعیت را روی دیسک ذخیره کن تا ZIP دقیقاً به‌روز ��اشد
        self.save_favs()
        self.save_featured()
        self.save_user_tags()
        fname = 'QuranCompass_Backup_%s.zip' % datetime.now().strftime('%Y%m%d_%H%M%S')
        out = self._p(fname)
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
            for name in ('favorites.json', 'featured.json', 'user_tags.json'):
                pth = self._p(name)
                if os.path.exists(pth):
                    z.write(pth, name)
            try:
                readable = {
                    'app': 'قطب‌نمای قرآنی',
                    'version': '3.0',
                    'exported_at': datetime.now().strftime('%Y-%m-%d %H:%M'),
                    'lab': self._clean_by_operator(self.favs),
                    'featured': self._clean_by_operator(self.featured),
                    'tags': self.get_all_tags(),
                }
                z.writestr('backup_readable.json', json.dumps(readable, ensure_ascii=False, indent=2))
            except Exception:
                pass
        return out

    def export_featured_json(self):
        if not self.featured:
            return None
        payload = {
            'app': 'قطب‌نمای قرآنی',
            'exported_at': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'featured': self._clean_by_operator(self.featured),
        }
        return self.export_clean_json('golchin_%s.json' % datetime.now().strftime('%Y%m%d_%H%M'), payload)

    # ---------- کشف گروهی و جفتی ----------
    def add_group_discovery(self, seed, targets, note='', relation_type='نامشخص', is_doubtful=False):
        """ثبت یک کشف گروهی با چند مقصد."""
        if not targets:
            toast('حداقل یک مقصد انتخاب کنید.', 'خطا')
            return
        entry = {
            'mode': 'گروهی',
            'seed_s': seed['s'], 'seed_a': seed['a'],
            'seed_arb': seed['arb'], 'seed_pers': seed['pers'],
            'all_targets': [{'s': t['s'], 'a': t['a'], 'arb': t.get('arb', ''),
                             'pers': t.get('pers', ''), 'operator': t.get('mode', '')}
                            for t in targets],
            'note': note, 'relation_type': relation_type, 'is_doubtful': bool(is_doubtful),
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        }
        self.favs.append(entry)
        self.save_favs()
        self.last_discovery_key = discovery_key(entry)
        self.last_discovery_section = lab_section_of(entry)
        open_note_editor(entry, 'lab', title='ثبت تحلیل کشف گروهی',
                         intro='کشف گروهی با %d مقصد ثبت شد. تحلیل و وضعیت تردید را انتخاب کنید:' % len(targets),
                         on_saved=lambda: setattr(self, 'last_discovery_section', lab_section_of(entry)),
                         saved_msg='کشف گروهی با موفقیت ثبت شد ✓')

    def add_pair_discovery(self, seed, ta, tb, note='', relation_type='نامشخص', is_doubtful=False):
        """ثبت یک جفت عملگری (دو مقصد). کارت اول نوع عملگر را تعیین می‌کند."""
        op_key = op_of({'mode': ta.get('mode', '')})
        entry = {
            'mode': 'جفت عملگری', 'pair_type': 'operator_pair', 'op_key': op_key,
            'seed_s': seed['s'], 'seed_a': seed['a'],
            'seed_arb': seed['arb'], 'seed_pers': seed['pers'],
            'all_targets': [
                {'s': ta['s'], 'a': ta['a'], 'arb': ta.get('arb', ''), 'pers': ta.get('pers', ''), 'operator': ta.get('mode', '')},
                {'s': tb['s'], 'a': tb['a'], 'arb': tb.get('arb', ''), 'pers': tb.get('pers', ''), 'operator': tb.get('mode', '')},
            ],
            'note': note, 'relation_type': relation_type, 'is_doubtful': bool(is_doubtful),
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        }
        self.favs.append(entry)
        self.save_favs()
        self.last_discovery_key = discovery_key(entry)
        self.last_discovery_section = lab_section_of(entry)
        open_note_editor(entry, 'lab', title='ثبت تحلیل جفت عملگری',
                         intro='جفت عملگری زیر بخش %s ثبت شد. تحلیل و وضعیت تردید را انتخاب کنید:' % op_key,
                         on_saved=lambda: setattr(self, 'last_discovery_section', lab_section_of(entry)),
                         saved_msg='جفت عملگری با موفقیت ثبت شد ✓')

    def remove_target_from_group(self, group_item, target_index):
        """حذف یک مقصد از یک کشف گروهی/جفتی."""
        found = None
        for f in self.favs:
            if (f.get('mode') == group_item.get('mode') and
                    f.get('seed_s') == group_item.get('seed_s') and
                    f.get('seed_a') == group_item.get('seed_a') and
                    f.get('seed_arb') == group_item.get('seed_arb') and
                    'all_targets' in f):
                found = f
                break
        if found is None:
            return False, 'کشف مورد نظر یافت نشد.'
        if 0 <= target_index < len(found['all_targets']):
            del found['all_targets'][target_index]
            if not found['all_targets']:
                self.favs.remove(found)
                self.save_favs()
                return True, 'آخرین مقصد حذف شد؛ کل کشف گروهی پاک شد.'
            self.save_favs()
            return True, 'مقصد از گروه حذف شد.'
        return False, 'ایندکس نامعتبر است.'

    # ---------- بازیابی ----------
    def restore_backup(self, payload):
        """بازیابی کشفیات و گلچین از محتوای JSON پشتیبان."""
        favs, featured = features.parse_backup(payload)
        self.favs = favs if isinstance(favs, list) else []
        self.featured = featured if isinstance(featured, list) else []
        for _it in self.favs:
            _normalize_item_modes(_it)
        for _it in self.featured:
            _normalize_item_modes(_it)
        self.save_favs()
        self.save_featured()
        return len(favs), len(featured)

    # ---------- بارگذاری از فایل JSON ----------
    def read_text_file(self, path):
        with open(path, encoding='utf-8') as f:
            return f.read()

    def _extract_items(self, payload, target):
        if isinstance(payload, list):
            return payload
        if isinstance(payload, dict):
            if 'lab' in payload or 'featured' in payload:
                favs, featured = features.parse_backup(payload)
                return favs if target == 'lab' else featured
            if target == 'lab' and 'favs' in payload:
                return payload.get('favs') or []
            if 'featured' in payload:
                return payload.get('featured') or []
        return []

    def _parse_json_tolerant(self, text):
        """پارس بردبار JSON: در برابر «Extra data»، دادهٔ تکراری/اضافه،
        BOM و JSON خط‌به‌خط مقاوم است. (payload, error) برمی‌گرداند."""
        import json as _json
        s = (text or '')
        # حذف BOM و نویسه‌های نامرئی ابتدا/انتها
        s = s.replace('\ufeff', '').replace('\u200b', '').strip()
        if not s:
            return None, 'متن خالی است.'
        # ۱) حالت عادی
        try:
            return _json.loads(s), None
        except Exception:
            pass
        # ۲) پارس مقادیر پشت‌سرهم و نادیده‌گرفتن دادهٔ اضافهٔ انتها
        dec = _json.JSONDecoder()
        collected = []
        idx, n = 0, len(s)
        while idx < n:
            while idx < n and s[idx] in ' \t\r\n,;\ufeff':
                idx += 1
            if idx >= n:
                break
            try:
                val, end = dec.raw_decode(s, idx)
            except Exception:
                break
            collected.append(val)
            if end <= idx:
                break
            idx = end
        if collected:
            if len(collected) == 1:
                return collected[0], None
            merged = []
            only_lists = True
            for v in collected:
                if isinstance(v, list):
                    merged.extend(v)
                else:
                    only_lists = False
            if only_lists:
                return merged, None
            return collected[0], None
        # ۳) JSON خط‌به‌خط
        out = []
        for line in s.splitlines():
            line = line.strip().rstrip(',')
            if not line or line in ('[', ']', '{', '}'):
                continue
            try:
                out.append(_json.loads(line))
            except Exception:
                pass
        if out:
            return out, None
        return None, 'قالب JSON قابل خواندن نبود.'

    def _import_item_list(self, items, target='lab', mode='merge'):
        """هستهٔ واردکردنِ فهرستِ کشف‌ها به لابراتوار یا گلچین (با حذفِ تکراری‌ها)."""
        if not isinstance(items, list):
            return (0, 0, 'ساختار فایل قابل‌شناسایی نیست.')
        dest = [] if mode == 'replace' else (self.favs if target == 'lab' else self.featured)
        existing = set()
        for it in dest:
            try:
                existing.add(discovery_key(it))
            except Exception:
                pass
        added = 0
        for it in items:
            if not isinstance(it, dict) or not it.get('seed_s'):
                continue
            it.setdefault('note', '')
            it.setdefault('relation_type', 'نامشخص')
            it.setdefault('is_doubtful', False)
            _normalize_item_modes(it)
            try:
                k = discovery_key(it)
            except Exception:
                k = None
            if k is not None and k in existing:
                continue
            dest.append(it)
            if k is not None:
                existing.add(k)
            added += 1
        if target == 'lab':
            self.favs = dest
            self.save_favs()
        else:
            self.featured = dest
            self.save_featured()
        return (added, len(dest), '')

    def import_items_json(self, text, target='lab', mode='merge'):
        payload, perr = self._parse_json_tolerant(text)
        if perr:
            return (0, 0, 'JSON نامعتبر: %s' % perr)
        items = self._extract_items(payload, target)
        return self._import_item_list(items, target, mode)

    def _merge_user_tags(self, raw_or_list):
        """برچسب‌های شخصی را از یک فهرست/متنِ JSON با برچسب‌های فعلی ادغام می‌کند.
        برچسب‌های پیش‌فرض، «نامشخص» و برچسب‌های ترکیبی (چنداتیکتی) اضافه نمی‌شوند."""
        try:
            tags = raw_or_list
            if isinstance(raw_or_list, str):
                tags = json.loads(raw_or_list)
            if not isinstance(tags, list):
                return
            sep = chr(1548)
            changed = False
            for t in tags:
                if not isinstance(t, str):
                    continue
                t = t.strip()
                if not t or t == 'نامشخص' or t in TagsScreen.DEFAULT or sep in t:
                    continue
                if t not in self.user_tags:
                    self.user_tags.append(t)
                    changed = True
            if changed:
                self.save_user_tags()
        except Exception:
            pass

    def import_from_path(self, path, target='lab', mode='merge'):
        """بارگذاریِ قدرتمند و بی‌خطا از:
          • فایلِ ZIP پشتیبان (نسخهٔ ویندوز یا خودِ اپ): favorites.json + featured.json + user_tags.json
          • فایلِ JSON کاملِ خروجیِ اپ (شاملِ lab/featured/tags) ← هر دو بخش و برچسب‌ها بازیابی می‌شوند
          • فایلِ JSON خام یا تک‌بخشی ← به مقصدِ انتخاب‌شده وارد می‌شود
        هرگز کرش نمی‌کند؛ خروجی: (added, total, err)."""
        import os as _os
        import zipfile as _zip
        try:
            if not path or not _os.path.exists(path):
                return (0, 0, 'فایل یافت نشد.')
            with open(path, 'rb') as _f:
                head = _f.read(4)
            # ---------- فایلِ ZIP ----------
            if head[:2] == b'PK':
                try:
                    with _zip.ZipFile(path) as z:
                        names = z.namelist()

                        def _member(basename):
                            for nm in names:
                                if nm.split('/')[-1].lower() == basename:
                                    return z.read(nm).decode('utf-8-sig', 'replace')
                            return None
                        favs_raw = _member('favorites.json')
                        feat_raw = _member('featured.json')
                        tags_raw = _member('user_tags.json')
                        readable = _member('backup_readable.json')
                except Exception as e:
                    return (0, 0, 'خواندن فایل ZIP ممکن نشد: %s' % str(e)[:80])
                # اگر فقط نسخهٔ خوانا داخلِ ZIP بود
                if favs_raw is None and feat_raw is None and readable is not None:
                    return self._restore_full_text(readable, mode, target)
                if favs_raw is None and feat_raw is None:
                    return (0, 0, 'داخلِ فایل ZIP، favorites.json یا featured.json پیدا نشد.')
                total_added = 0
                grand_total = 0
                if favs_raw is not None:
                    a, t, err = self.import_items_json(favs_raw, 'lab', mode)
                    if err:
                        return (0, 0, err)
                    total_added += a
                    grand_total += t
                if feat_raw is not None:
                    a, t, err = self.import_items_json(feat_raw, 'featured', mode)
                    if err:
                        return (0, 0, err)
                    total_added += a
                    grand_total += t
                if tags_raw is not None:
                    self._merge_user_tags(tags_raw)
                return (total_added, grand_total, '')
            # ---------- فایلِ متنیِ JSON ----------
            try:
                with open(path, encoding='utf-8-sig', errors='replace') as _f:
                    raw = _f.read()
            except Exception as e:
                return (0, 0, 'خواندن فایل ممکن نشد: %s' % str(e)[:80])
            return self._restore_full_text(raw, mode, target)
        except Exception as e:
            return (0, 0, 'خطا در بارگذاری: %s' % str(e)[:80])

    def _restore_full_text(self, raw, mode='merge', target='lab'):
        """اگر متن یک «پشتیبانِ کامل» باشد (شاملِ lab/featured)، هر دو بخش و برچسب‌ها را بازیابی می‌کند؛
        در غیرِ این‌صورت به‌صورتِ تک‌بخشی به مقصدِ انتخاب‌شده وارد می‌کند."""
        payload, perr = self._parse_json_tolerant(raw)
        if perr:
            return (0, 0, 'JSON نامعتبر: %s' % perr)
        if isinstance(payload, dict) and ('lab' in payload or 'featured' in payload):
            favs, featured = features.parse_backup(payload)
            fa, ft, err = self._import_item_list(favs, 'lab', mode)
            if err:
                return (0, 0, err)
            xa, xt, err = self._import_item_list(featured, 'featured', mode)
            if err:
                return (0, 0, err)
            tags = payload.get('tags')
            if isinstance(tags, list):
                self._merge_user_tags(tags)
            return (fa + xa, ft + xt, '')
        # تک‌بخشی (فهرستِ خام یا فایلِ favorites/featured تنها)
        return self.import_items_json(raw, target, mode)

    # ---------- پشتیبان ----------
    def make_backup(self, dest_dir=None):
        fname = 'backup_%s.zip' % datetime.now().strftime('%Y%m%d_%H%M')
        if dest_dir:
            out = os.path.join(dest_dir, fname)
        else:
            out = self._p(fname)
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
            for name in ('favorites.json', 'featured.json', 'user_tags.json'):
                pth = self._p(name)
                if os.path.exists(pth):
                    z.write(pth, name)
        return out


if __name__ == '__main__':
    QuranMirrorApp().run()
